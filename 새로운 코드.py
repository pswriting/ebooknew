import streamlit as st
import re
import json
import html
import base64
import urllib.parse
import uuid
import platform
import hashlib
import requests
from datetime import datetime, timedelta
from pathlib import Path

# Claude API
try:
    import anthropic
    CLAUDE_AVAILABLE = True
except ImportError:
    CLAUDE_AVAILABLE = False

# Gemini (이미지 생성용으로만 사용)
try:
    import google.generativeai as genai
    from google import genai as google_genai
    from google.genai import types as genai_types
    GEMINI_AVAILABLE = True
    IMAGEN_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    IMAGEN_AVAILABLE = False

# YouTube 자막 추출용
try:
    from youtube_transcript_api import YouTubeTranscriptApi
    YOUTUBE_TRANSCRIPT_AVAILABLE = True
except ImportError:
    YOUTUBE_TRANSCRIPT_AVAILABLE = False

# 브라우저 ID용 (클라우드 배포 시 필요)
try:
    from streamlit_javascript import st_javascript
    BROWSER_ID_AVAILABLE = True
except ImportError:
    BROWSER_ID_AVAILABLE = False

# 쿠키 매니저 (데이터 저장용)
try:
    import extra_streamlit_components as stx
    COOKIE_AVAILABLE = True
except ImportError:
    COOKIE_AVAILABLE = False

# Word 문서 생성용
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ==========================================
# 설정
# ==========================================
def get_config_path():
    return Path.home() / ".ebook_app_config.json"

def load_config():
    try:
        if get_config_path().exists():
            with open(get_config_path(), 'r') as f:
                return json.load(f)
    except:
        pass
    return {}

def save_config(data):
    try:
        config = load_config()
        config.update(data)
        with open(get_config_path(), 'w') as f:
            json.dump(config, f)
    except:
        pass

def load_saved_api_key():
    return load_config().get('api_key', '')

def save_api_key(api_key):
    save_config({'api_key': api_key})

def is_authenticated():
    return load_config().get('authenticated', False)

def save_authenticated():
    save_config({'authenticated': True})

# ==========================================
# 기기별 승인 시스템 (JSONBin)
# ==========================================
JSONBIN_API_KEY = "$2a$10$19x6FoPHLekIcgOGgYqyROGXOyC3p4d0Yp.C89yXfE2IFYn7yYy0K"
JSONBIN_BIN_ID = None  # 첫 실행 시 자동 생성됨
ADMIN_PASSWORD = "admin2024"  # 관리자 비밀번호 (변경 가능)

def get_device_id():
    """브라우저별 고유 ID 생성 (클라우드 배포용)"""
    # 캐시된 브라우저 ID가 있으면 사용
    if 'browser_device_id' in st.session_state and st.session_state['browser_device_id']:
        return st.session_state['browser_device_id']
    return None

def get_saved_password():
    """저장된 비밀번호 불러오기"""
    if 'saved_password' in st.session_state:
        return st.session_state['saved_password']
    return None

def get_saved_api_key():
    """저장된 API 키 불러오기"""
    if 'saved_api_key' in st.session_state:
        return st.session_state['saved_api_key']
    return None

def save_device_id_to_browser(device_id):
    """기기 코드 저장"""
    st.session_state['browser_device_id'] = device_id
    st.session_state['pending_save_device'] = device_id

def save_password_to_browser(password):
    """비밀번호 저장"""
    st.session_state['saved_password'] = password
    st.session_state['pending_save_password'] = password

def save_api_key_to_browser(api_key):
    """API 키 저장"""
    st.session_state['saved_api_key'] = api_key
    st.session_state['pending_save_api'] = api_key

def get_bin_id():
    """저장된 Bin ID 반환"""
    return load_config().get('jsonbin_bin_id', None)

def save_bin_id(bin_id):
    """Bin ID 저장"""
    save_config({'jsonbin_bin_id': bin_id})

def create_jsonbin():
    """새 JSONBin 생성 (첫 실행 기기는 자동 승인)"""
    try:
        url = "https://api.jsonbin.io/v3/b"
        headers = {
            "Content-Type": "application/json",
            "X-Master-Key": JSONBIN_API_KEY,
            "X-Bin-Private": "true",
            "X-Bin-Name": "writey_approved_devices"
        }
        # 첫 실행 기기(관리자)는 자동 승인
        current_device = get_device_id()
        data = {"approved_devices": [current_device], "pending_devices": []}
        response = requests.post(url, json=data, headers=headers)
        if response.status_code == 200:
            bin_id = response.json()["metadata"]["id"]
            save_bin_id(bin_id)
            return bin_id
    except:
        pass
    return None

def get_jsonbin_data():
    """JSONBin에서 데이터 가져오기"""
    bin_id = get_bin_id()
    if not bin_id:
        return None
    try:
        url = f"https://api.jsonbin.io/v3/b/{bin_id}/latest"
        headers = {"X-Master-Key": JSONBIN_API_KEY}
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            return response.json()["record"]
    except:
        pass
    return None

def update_jsonbin_data(data):
    """JSONBin 데이터 업데이트"""
    bin_id = get_bin_id()
    if not bin_id:
        return False
    try:
        url = f"https://api.jsonbin.io/v3/b/{bin_id}"
        headers = {
            "Content-Type": "application/json",
            "X-Master-Key": JSONBIN_API_KEY
        }
        response = requests.put(url, json=data, headers=headers)
        return response.status_code == 200
    except:
        return False

def is_device_approved_jsonbin():
    """현재 기기가 JSONBin에서 승인되었는지 확인"""
    device_id = get_device_id()
    if not device_id:
        return False
    data = get_jsonbin_data()
    if data and "approved_devices" in data:
        return device_id in data["approved_devices"]
    return False

def add_device_to_approved(device_id):
    """기기를 승인 목록에 추가"""
    data = get_jsonbin_data()
    if data:
        if "approved_devices" not in data:
            data["approved_devices"] = []
        if device_id not in data["approved_devices"]:
            data["approved_devices"].append(device_id)
        # pending에서 제거
        if "pending_devices" in data and device_id in data["pending_devices"]:
            data["pending_devices"].remove(device_id)
        return update_jsonbin_data(data)
    return False

def remove_device_from_approved(device_id):
    """기기를 승인 목록에서 제거"""
    data = get_jsonbin_data()
    if data and "approved_devices" in data:
        if device_id in data["approved_devices"]:
            data["approved_devices"].remove(device_id)

def save_device_settings(device_id, api_key=None, password=None):
    """기기별 설정을 JSONBin에 저장"""
    data = get_jsonbin_data()
    if data:
        if "device_settings" not in data:
            data["device_settings"] = {}
        if device_id not in data["device_settings"]:
            data["device_settings"][device_id] = {}
        if api_key:
            data["device_settings"][device_id]["api_key"] = api_key
        if password:
            data["device_settings"][device_id]["password"] = password
        return update_jsonbin_data(data)
    return False

def get_device_settings(device_id):
    """기기별 설정을 JSONBin에서 불러오기"""
    data = get_jsonbin_data()
    if data and "device_settings" in data:
        return data["device_settings"].get(device_id, {})
    return {}

def add_device_to_pending(device_id):
    """기기를 대기 목록에 추가"""
    data = get_jsonbin_data()
    if data:
        if "pending_devices" not in data:
            data["pending_devices"] = []
        if device_id not in data["pending_devices"] and device_id not in data.get("approved_devices", []):
            data["pending_devices"].append(device_id)
            return update_jsonbin_data(data)
    return False

def get_approved_devices():
    """승인된 기기 목록 반환"""
    data = get_jsonbin_data()
    if data:
        return data.get("approved_devices", [])
    return []

def get_pending_devices():
    """대기 중인 기기 목록 반환"""
    data = get_jsonbin_data()
    if data:
        return data.get("pending_devices", [])
    return []

# 비디오 배경용 base64 인코딩
@st.cache_data
def get_video_base64(video_path):
    try:
        with open(video_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except:
        return None

st.set_page_config(page_title="Writey", layout="wide", page_icon="✍")

# 쿠키 매니저 초기화 및 데이터 불러오기/저장
if COOKIE_AVAILABLE:
    cookie_manager = stx.CookieManager(key="writey_cookies")
    cookies = cookie_manager.get_all()

    # 쿠키에서 데이터 불러오기
    if cookies:
        if 'writey_device_id' in cookies and cookies['writey_device_id']:
            if 'browser_device_id' not in st.session_state:
                st.session_state['browser_device_id'] = cookies['writey_device_id']
        if 'writey_password' in cookies and cookies['writey_password']:
            if 'saved_password' not in st.session_state:
                st.session_state['saved_password'] = cookies['writey_password']
        if 'writey_api_key' in cookies and cookies['writey_api_key']:
            if 'saved_api_key' not in st.session_state:
                st.session_state['saved_api_key'] = cookies['writey_api_key']

    # pending 값이 있으면 쿠키에 저장
    if 'pending_save_device' in st.session_state:
        cookie_manager.set('writey_device_id', st.session_state['pending_save_device'], expires_at=datetime.now() + timedelta(days=365))
        del st.session_state['pending_save_device']
    if 'pending_save_password' in st.session_state:
        cookie_manager.set('writey_password', st.session_state['pending_save_password'], expires_at=datetime.now() + timedelta(days=365))
        del st.session_state['pending_save_password']
    if 'pending_save_api' in st.session_state:
        cookie_manager.set('writey_api_key', st.session_state['pending_save_api'], expires_at=datetime.now() + timedelta(days=365))
        del st.session_state['pending_save_api']
else:
    cookie_manager = None

# ==========================================
# APPLE STYLE CSS
# ==========================================
st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+KR:wght@400;500;600;700;900&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;500;600;700;800;900&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@400;500;600;700&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Cinzel:wght@400;500;600;700&display=swap');

/* S-Core Dream 폰트 */
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-1Thin.woff') format('woff');
    font-weight: 100;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-2ExtraLight.woff') format('woff');
    font-weight: 200;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-3Light.woff') format('woff');
    font-weight: 300;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-4Regular.woff') format('woff');
    font-weight: 400;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-5Medium.woff') format('woff');
    font-weight: 500;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-6Bold.woff') format('woff');
    font-weight: 600;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-7ExtraBold.woff') format('woff');
    font-weight: 700;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-8Heavy.woff') format('woff');
    font-weight: 800;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-9Black.woff') format('woff');
    font-weight: 900;
}

:root {
    --gold: #d4af37;
    --gold-light: #f4e4bc;
    --gold-dark: #996515;
    --rose-gold: #b76e79;
    --cream: #faf7f2;
    --charcoal: #1a1a1a;
    --dark: #0a0a0a;
    --card: rgba(20,20,20,0.9);
    --card2: rgba(30,30,30,0.9);
    --text: #f5f5f5;
    --text2: #888888;
    --line: rgba(212,175,55,0.2);
    --glow: rgba(212,175,55,0.4);
    --success: #50c878;
    --warning: #ffb347;
    --danger: #ff6b6b;
}

/* 애니메이션 정의 */
@keyframes fadeInUp {
    from { opacity: 0; transform: translateY(30px); }
    to { opacity: 1; transform: translateY(0); }
}
@keyframes shimmer {
    0% { background-position: -200% 0; }
    100% { background-position: 200% 0; }
}
@keyframes pulse {
    0%, 100% { opacity: 1; transform: scale(1); }
    50% { opacity: 0.8; transform: scale(1.02); }
}
@keyframes borderGlow {
    0%, 100% { box-shadow: 0 0 5px var(--glow), inset 0 0 5px rgba(212,175,55,0.1); }
    50% { box-shadow: 0 0 20px var(--glow), inset 0 0 10px rgba(212,175,55,0.2); }
}
@keyframes float {
    0%, 100% { transform: translateY(0); }
    50% { transform: translateY(-5px); }
}
@keyframes goldShine {
    0% { background-position: -100% 0; }
    100% { background-position: 200% 0; }
}

*:not([data-testid*="Icon"]):not(.material-icons):not([class*="icon"]):not(span[aria-hidden="true"]) {
    font-family: 'S-CoreDream', 'Pretendard', -apple-system, sans-serif !important;
}
/* 아이콘 폰트 복원 */
[data-testid*="Icon"], .material-icons, span[aria-hidden="true"], button[kind="header"] span {
    font-family: 'Material Symbols Rounded', 'Material Icons', sans-serif !important;
}
.stDeployButton, footer, #MainMenu { display: none !important; }
/* 헤더 투명하게 (사이드바 버튼은 보임) */
header[data-testid="stHeader"] {
    background: transparent !important;
}

/* 럭셔리 배경 - 미세한 그라데이션 */
.stApp {
    background:
        radial-gradient(ellipse at 20% 0%, rgba(212,175,55,0.03) 0%, transparent 50%),
        radial-gradient(ellipse at 80% 100%, rgba(183,110,121,0.03) 0%, transparent 50%),
        linear-gradient(180deg, #0a0a0a 0%, #050505 50%, #0a0a0a 100%) !important;
    background-attachment: fixed;
}

.main .block-container { max-width: 1000px; padding: 3rem 2rem; }

/* 사이드바 - 미니멀 */
[data-testid="stSidebar"] {
    background: var(--charcoal) !important;
    border-right: 1px solid var(--line);
}
[data-testid="stSidebar"] * { color: var(--text2) !important; }

/* 타이포그래피 - 가독성 향상 */
h1, h2, h3 { color: var(--text) !important; font-weight: 400 !important; letter-spacing: 1px; }
h1 { font-size: 36px !important; color: var(--cream) !important; }
h2 { font-size: 28px !important; margin-bottom: 20px !important; }
h3 { font-size: 22px !important; color: var(--gold) !important; }
p, span, label, div { color: var(--text) !important; font-size: 16px !important; line-height: 1.7 !important; }
li { font-size: 16px !important; line-height: 1.8 !important; }

/* 버튼 - 럭셔리 골드 + 고급 효과 */
.stButton > button {
    background: linear-gradient(135deg, rgba(212,175,55,0.1) 0%, transparent 50%, rgba(212,175,55,0.1) 100%) !important;
    color: var(--gold) !important;
    border: 1px solid var(--gold) !important;
    border-radius: 4px;
    font-weight: 600;
    font-size: 16px !important;
    padding: 18px 40px;
    letter-spacing: 2px;
    text-transform: uppercase;
    transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
    position: relative;
    overflow: hidden;
    box-shadow: 0 2px 10px rgba(212,175,55,0.1);
}
.stButton > button::before {
    content: '';
    position: absolute;
    top: 0;
    left: -100%;
    width: 100%;
    height: 100%;
    background: linear-gradient(90deg, transparent, rgba(212,175,55,0.3), transparent);
    transition: left 0.6s ease;
}
.stButton > button:hover::before {
    left: 100%;
}
.stButton > button::after {
    content: '';
    position: absolute;
    inset: 0;
    background: linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 50%, var(--gold) 100%);
    opacity: 0;
    transition: opacity 0.4s ease;
    z-index: -1;
}
.stButton > button:hover {
    color: var(--dark) !important;
    border-color: var(--gold-light) !important;
    box-shadow: 0 8px 30px rgba(212,175,55,0.4), 0 0 20px rgba(212,175,55,0.2);
    transform: translateY(-3px);
    text-shadow: 0 1px 2px rgba(0,0,0,0.3);
}
.stButton > button:hover::after {
    opacity: 1;
}
.stButton > button:active {
    transform: translateY(-1px);
    box-shadow: 0 4px 15px rgba(212,175,55,0.3);
}

/* 입력 필드 - 밝은 배경 + 검은 글씨 */
.stTextInput input, .stTextArea textarea, .stNumberInput input {
    background: #ffffff !important;
    background-color: #ffffff !important;
    border: 1px solid var(--line) !important;
    border-radius: 6px !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    caret-color: #000000 !important;
    padding: 18px !important;
    font-size: 17px !important;
}
.stTextInput input:focus, .stTextArea textarea:focus, .stNumberInput input:focus {
    border-color: var(--gold) !important;
    box-shadow: 0 0 0 2px rgba(212,175,55,0.2) !important;
}

/* 셀렉트박스 컨테이너 */
.stSelectbox > div > div {
    background: var(--card) !important;
    border: 1px solid var(--line) !important;
    border-radius: 0;
}
/* 셀렉트박스 선택된 값 - 흰색 */
.stSelectbox [data-baseweb="select"] > div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* 스코어 카드 - 럭셔리 */
.score-card {
    background: linear-gradient(145deg, var(--card) 0%, rgba(30,30,30,0.95) 100%) !important;
    border: 2px solid var(--gold);
    border-radius: 20px;
    padding: 50px 40px;
    text-align: center;
    animation: fadeInUp 0.6s ease-out;
    transition: all 0.5s cubic-bezier(0.4, 0, 0.2, 1);
    position: relative;
    overflow: hidden;
    box-shadow: 0 10px 40px rgba(212,175,55,0.15);
}
.score-card::before {
    content: '';
    position: absolute;
    top: 0;
    left: 0;
    right: 0;
    height: 3px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
    opacity: 1;
}
.score-card:hover {
    border-color: var(--gold);
    box-shadow: 0 20px 60px rgba(212,175,55,0.3), inset 0 1px 0 rgba(212,175,55,0.1);
    transform: translateY(-5px);
}
.score-card:hover::before {
    opacity: 1;
}
.score-number {
    font-size: 140px;
    font-weight: 300;
    background: linear-gradient(135deg, var(--gold-light) 0%, var(--gold) 50%, var(--gold-dark) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    line-height: 1;
    letter-spacing: -4px;
    animation: fadeInUp 0.8s ease-out;
    filter: drop-shadow(0 2px 4px rgba(212,175,55,0.3));
}

/* 정보 카드 + 애니메이션 */
.info-card {
    background: transparent !important;
    border: none;
    border-left: 2px solid var(--gold);
    padding: 20px 24px;
    margin: 20px 0;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.3s ease;
}
.info-card:hover {
    background: rgba(201,169,98,0.05) !important;
    border-left-width: 4px;
    padding-left: 22px;
}

/* 스탯 박스 + 애니메이션 */
.stat-box {
    background: var(--card) !important;
    border: 1px solid var(--line);
    padding: 32px;
    text-align: center;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.4s ease;
}
.stat-box:hover {
    transform: translateY(-4px);
    box-shadow: 0 8px 25px rgba(0,0,0,0.3);
    border-color: var(--gold);
}
.stat-value {
    font-size: 42px;
    font-weight: 200;
    color: var(--gold) !important;
    letter-spacing: -2px;
    transition: transform 0.3s ease;
}
.stat-box:hover .stat-value {
    transform: scale(1.05);
}
.stat-label {
    font-size: 11px;
    color: var(--text2) !important;
    margin-top: 12px;
    text-transform: uppercase;
    letter-spacing: 3px;
}

/* 데이터 카드 + 애니메이션 */
.data-card {
    background: var(--card) !important;
    border-left: 2px solid var(--gold);
    padding: 20px 24px;
    margin: 16px 0;
    animation: fadeInUp 0.4s ease-out;
    transition: all 0.3s ease;
}
.data-card:hover {
    border-left-width: 4px;
    background: var(--card2) !important;
}

/* 서머리 허브 + 애니메이션 */
.summary-hub {
    background: var(--card) !important;
    border: 1px solid var(--line);
    padding: 40px;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.4s ease;
}
.summary-hub:hover {
    border-color: var(--gold);
}

/* 배지 - 미니멀 + 펄스 */
.verdict-go {
    background: transparent !important;
    color: var(--success) !important;
    border: 1px solid var(--success);
    padding: 12px 32px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    text-transform: uppercase;
    animation: fadeInUp 0.6s ease-out;
}
.verdict-wait {
    background: transparent !important;
    color: var(--warning) !important;
    border: 1px solid var(--warning);
    padding: 12px 32px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    animation: fadeInUp 0.6s ease-out;
}
.verdict-no {
    background: transparent !important;
    color: var(--danger) !important;
    border: 1px solid var(--danger);
    padding: 12px 32px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    animation: fadeInUp 0.6s ease-out;
}

/* 네비게이션 */
.premium-nav-container {
    background: transparent;
    border-top: 1px solid var(--line);
    border-bottom: 1px solid var(--line);
    padding: 0;
    margin-bottom: 48px;
}
.nav-item {
    padding: 18px 12px;
    text-align: center;
    font-size: 14px;
    color: var(--text2);
    letter-spacing: 1px;
    transition: all 0.3s ease;
}
.nav-item.active {
    background: linear-gradient(135deg, rgba(212,175,55,0.2) 0%, rgba(212,175,55,0.1) 100%);
    color: var(--gold) !important;
    font-weight: 600;
    border-bottom: 3px solid var(--gold);
    box-shadow: 0 4px 15px rgba(212,175,55,0.2);
}

/* 섹션 타이틀 - 화려하게 */
.section-title-box {
    background: linear-gradient(135deg, rgba(212,175,55,0.15) 0%, rgba(183,110,121,0.1) 50%, rgba(212,175,55,0.15) 100%);
    border: 2px solid rgba(212,175,55,0.4);
    border-radius: 16px;
    padding: 30px 40px;
    margin-bottom: 35px;
    text-align: center;
    position: relative;
    overflow: hidden;
    animation: fadeInUp 0.5s ease-out;
}
.section-title-box::before {
    content: '';
    position: absolute;
    top: 0;
    left: -100%;
    width: 200%;
    height: 100%;
    background: linear-gradient(90deg, transparent, rgba(212,175,55,0.1), transparent);
    animation: shimmer 3s ease-in-out infinite;
}
.section-title-box h2 {
    font-size: 32px !important;
    background: linear-gradient(135deg, var(--gold-light) 0%, var(--gold) 50%, var(--gold-dark) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    margin: 0 0 10px 0 !important;
    font-weight: 600 !important;
    letter-spacing: 3px;
}
.section-title-box p {
    color: var(--text) !important;
    font-size: 17px !important;
    margin: 0 !important;
    opacity: 0.85;
}
.section-step {
    display: inline-block;
    background: var(--gold);
    color: var(--dark) !important;
    font-size: 13px;
    font-weight: 700;
    padding: 6px 16px;
    border-radius: 20px;
    margin-bottom: 15px;
    letter-spacing: 2px;
}

/* 제목 카드 + 애니메이션 */
.title-card {
    background: transparent;
    border: 1px solid var(--line);
    padding: 28px;
    margin: 16px 0;
    transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
    animation: fadeInUp 0.5s ease-out;
}
.title-card:hover {
    border-color: var(--gold);
    background: rgba(201,169,98,0.05);
    transform: translateX(8px);
    box-shadow: -4px 0 20px rgba(201,169,98,0.15);
}
.title-main {
    font-size: 18px;
    font-weight: 400;
    color: var(--text) !important;
    letter-spacing: 1px;
    transition: color 0.3s ease;
}
.title-card:hover .title-main {
    color: var(--gold) !important;
}
.title-sub {
    font-size: 13px;
    color: var(--text2) !important;
    margin-top: 8px;
}

/* 로그인 - 럭셔리 */
.login-card {
    max-width: 420px;
    margin: 100px auto;
    padding: 70px 50px;
    background: linear-gradient(145deg, rgba(25,25,25,0.98) 0%, rgba(15,15,15,0.98) 100%);
    border: 1px solid var(--line);
    text-align: center;
    animation: fadeInUp 0.8s ease-out;
    position: relative;
    box-shadow: 0 25px 80px rgba(0,0,0,0.5), 0 0 40px rgba(212,175,55,0.05);
}
.login-card::before {
    content: '';
    position: absolute;
    top: -1px;
    left: 20%;
    right: 20%;
    height: 2px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
}
.login-card::after {
    content: '';
    position: absolute;
    bottom: -1px;
    left: 20%;
    right: 20%;
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--gold-dark), transparent);
}
.login-title {
    font-size: 32px;
    font-weight: 300;
    background: linear-gradient(135deg, var(--cream) 0%, var(--gold-light) 50%, var(--gold) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: 10px;
    animation: fadeInUp 1s ease-out;
}
.login-subtitle {
    font-size: 11px;
    color: var(--gold) !important;
    margin-top: 20px;
    letter-spacing: 4px;
    text-transform: uppercase;
    animation: fadeInUp 1.2s ease-out;
    opacity: 0.8;
}

/* 헤더 - 럭셔리 */
.main-header {
    text-align: center;
    padding: 80px 20px 60px;
    margin-bottom: 50px;
    border-bottom: 1px solid var(--line);
    animation: fadeInUp 0.6s ease-out;
    position: relative;
    background: linear-gradient(180deg, rgba(212,175,55,0.02) 0%, transparent 100%);
}
.main-header::after {
    content: '';
    position: absolute;
    bottom: -1px;
    left: 10%;
    right: 10%;
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
}
.main-header-brand {
    font-size: 11px;
    color: var(--gold) !important;
    letter-spacing: 10px;
    text-transform: uppercase;
    animation: fadeInUp 0.8s ease-out;
    text-shadow: 0 0 20px rgba(212,175,55,0.3);
}
.main-header-title {
    font-size: 42px;
    font-weight: 200;
    background: linear-gradient(135deg, var(--cream) 0%, var(--gold-light) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: 8px;
    margin-top: 24px;
    animation: fadeInUp 1s ease-out;
}
.header-tagline {
    font-size: 13px;
    color: var(--text2) !important;
    margin-top: 24px;
    letter-spacing: 3px;
    animation: fadeInUp 1.2s ease-out;
}

/* Expander + 애니메이션 */
.stExpander {
    background: var(--card) !important;
    border: 1px solid var(--line) !important;
    border-radius: 0 !important;
    animation: fadeInUp 0.4s ease-out;
    transition: border-color 0.3s ease;
}
.stExpander:hover {
    border-color: var(--gold) !important;
}
.stProgress > div > div > div {
    background: linear-gradient(90deg, var(--gold-dark), var(--gold), var(--gold-light), var(--gold), var(--gold-dark)) !important;
    background-size: 300% 100%;
    animation: goldShine 3s ease infinite;
    border-radius: 4px;
    box-shadow: 0 0 15px rgba(212,175,55,0.4);
}
.stProgress > div > div {
    background: rgba(20,20,20,0.8);
    border-radius: 4px;
    border: 1px solid var(--line);
}

/* 라디오 & 탭 */
.stRadio > div { background: transparent; border: 1px solid var(--line); padding: 16px; }
.stTabs [data-baseweb="tab-list"] { background: transparent; border-bottom: 1px solid var(--line); }
.stTabs [aria-selected="true"] {
    background: transparent !important;
    color: var(--gold) !important;
    border-bottom: 2px solid var(--gold) !important;
}

/* 알림 */
.stSuccess > div { background: rgba(74,124,89,0.1) !important; border: 1px solid rgba(74,124,89,0.3) !important; border-radius: 0; }
.stWarning > div { background: rgba(196,154,61,0.1) !important; border: 1px solid rgba(196,154,61,0.3) !important; border-radius: 0; }
.stError > div { background: rgba(139,64,73,0.1) !important; border: 1px solid rgba(139,64,73,0.3) !important; border-radius: 0; }
.stInfo > div { background: rgba(201,169,98,0.1) !important; border: 1px solid var(--line) !important; border-radius: 0; }

/* 스크롤바 */
::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: var(--dark); }
::-webkit-scrollbar-thumb { background: var(--gold-dark); }

/* 다운로드 버튼 - 럭셔리 골드 */
.stDownloadButton button {
    background: linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 50%, var(--gold) 100%) !important;
    background-size: 200% 100%;
    color: var(--dark) !important;
    border: none !important;
    border-radius: 2px;
    font-weight: 600;
    letter-spacing: 3px;
    box-shadow: 0 4px 20px rgba(212,175,55,0.3);
    transition: all 0.4s ease;
    text-shadow: 0 1px 1px rgba(255,255,255,0.2);
}
.stDownloadButton button:hover {
    background-position: 100% 0 !important;
    box-shadow: 0 8px 35px rgba(212,175,55,0.5);
    transform: translateY(-2px);
}

/* 구분선 */
hr { border: none; height: 1px; background: var(--line); margin: 40px 0; }

/* 표지 미리보기 - 실제 책처럼 */
.book-wrapper {
    perspective: 1000px;
    display: flex;
    justify-content: center;
    padding: 30px;
    background: linear-gradient(135deg, #1a1a1a 0%, #2d2d2d 100%);
    border-radius: 8px;
}
.ebook-cover {
    font-family: 'Pretendard', sans-serif !important;
    box-shadow:
        0 0 5px rgba(0,0,0,0.3),
        5px 5px 15px rgba(0,0,0,0.4),
        10px 10px 30px rgba(0,0,0,0.3),
        15px 15px 50px rgba(0,0,0,0.2),
        inset -3px 0 10px rgba(0,0,0,0.2);
    transform: rotateY(-3deg);
    border-radius: 0 3px 3px 0;
    position: relative;
}
.ebook-cover::before {
    content: '';
    position: absolute;
    left: 0;
    top: 0;
    bottom: 0;
    width: 25px;
    background: linear-gradient(90deg,
        rgba(0,0,0,0.4) 0%,
        rgba(0,0,0,0.1) 30%,
        rgba(255,255,255,0.05) 50%,
        rgba(0,0,0,0.1) 70%,
        rgba(0,0,0,0.3) 100%);
    border-radius: 3px 0 0 3px;
}
.ebook-cover::after {
    content: '';
    position: absolute;
    right: 0;
    top: 2px;
    bottom: 2px;
    width: 8px;
    background: linear-gradient(90deg,
        rgba(255,255,255,0.03) 0%,
        rgba(255,255,255,0.08) 50%,
        rgba(0,0,0,0.1) 100%);
}
.ebook-cover * {
    color: inherit !important;
    -webkit-text-fill-color: inherit !important;
}

/* ============================================
   입력 필드 텍스트 색상 - 최우선 적용
   ============================================ */

/* 모든 입력 필드 - 흰 배경 + 검은 글씨 */
.stTextInput input,
.stTextArea textarea,
.stNumberInput input,
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stNumberInput"] input {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    caret-color: #000000 !important;
}

/* Placeholder 색상 */
input::placeholder,
textarea::placeholder {
    color: #888888 !important;
    -webkit-text-fill-color: #888888 !important;
}

/* 셀렉트박스 - 선택된 값 (어두운 배경에 흰 글씨) */
.stSelectbox [data-baseweb="select"] > div,
.stSelectbox [data-baseweb="select"] span,
.stSelectbox > div > div > div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* ============================================
   드롭다운/팝오버 - 검은 글씨 (흰 배경)
   ============================================ */
[data-baseweb="popover"],
[data-baseweb="popover"] *,
[data-baseweb="menu"],
[data-baseweb="menu"] *,
[data-baseweb="list"],
[data-baseweb="list"] *,
[role="listbox"],
[role="listbox"] *,
[role="option"],
[role="option"] *,
.stSelectbox ul,
.stSelectbox ul *,
.stSelectbox li,
.stSelectbox li * {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
}

/* 드롭다운 옵션 호버 */
[role="option"]:hover,
[data-baseweb="menu"] li:hover,
.stSelectbox li:hover {
    background: #f0f0f0 !important;
    background-color: #f0f0f0 !important;
}

/* select 요소 */
select,
select option {
    color: #000000 !important;
    background: #ffffff !important;
}

/* Expander 스타일 정리 */
.stExpander details summary {
    background: var(--card) !important;
    overflow: hidden !important;
}
/* 모든 텍스트 숨기기 (keyboard_arrow 등 영어 텍스트 포함) */
.stExpander details summary * {
    font-size: 0 !important;
    color: transparent !important;
    -webkit-text-fill-color: transparent !important;
}
/* 한국어 제목만 보이게 */
.stExpander details summary p {
    font-size: 15px !important;
    color: var(--text) !important;
    -webkit-text-fill-color: var(--text) !important;
}
/* 화살표 아이콘만 보이게 */
.stExpander details summary svg {
    width: 20px !important;
    height: 20px !important;
    color: var(--gold) !important;
    fill: var(--gold) !important;
}

/* 버튼 앞 불필요한 라벨 숨기기 */
.stButton > div:not([data-testid="baseButton-secondary"]):not([data-testid="baseButton-primary"]) > p,
.stButton > div > div > p:first-child:not(:last-child),
.stButton label,
.stExpander .stButton > div:first-child > p {
    display: none !important;
}
/* 링크버튼 라벨 숨기기 */
.stLinkButton > div:first-child > p {
    display: none !important;
}
</style>
""", unsafe_allow_html=True)


# 인증
CORRECT_PASSWORD = "cashmaker2024"

# JSONBin 초기화 (첫 실행 시 Bin 생성)
if not get_bin_id():
    with st.spinner("초기 설정 중..."):
        create_jsonbin()

# 로컬 파일에서 저장된 로그인 정보 로드
saved_config = load_config()
if saved_config.get('authenticated') and saved_config.get('device_id'):
    st.session_state['browser_device_id'] = saved_config.get('device_id')
    st.session_state['authenticated'] = True

if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

if not st.session_state['authenticated']:
    st.markdown("""
    <div class="login-card">
        <div class="login-subtitle">CASHMAKER</div>
        <div class="login-title">Writey</div>
        <div class="login-subtitle">Premium E-Book Studio</div>
    </div>
    """, unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        device_id = get_device_id()

        # 브라우저 ID가 없으면 수동 입력 모드
        if not device_id:
            st.info("🔑 기기 코드를 입력하세요")

            # 새 기기 등록
            if 'new_device_code' not in st.session_state:
                import random
                import string
                st.session_state['new_device_code'] = 'DEV_' + ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))

            tab1, tab2 = st.tabs(["기존 코드 입력", "새 기기 등록"])

            with tab1:
                manual_id = st.text_input("기기 코드", key="manual_device_id", placeholder="DEV_XXXXXXXX")
                if st.button("확인", key="btn_manual_id"):
                    if manual_id:
                        st.session_state['browser_device_id'] = manual_id
                        save_device_id_to_browser(manual_id)
                        st.rerun()

            with tab2:
                new_code = st.session_state['new_device_code']
                st.markdown(f'<div style="background:#f0f0f0;padding:15px;border-radius:8px;text-align:center;"><code style="color:#000;font-size:18px;font-weight:bold;">{new_code}</code></div>', unsafe_allow_html=True)
                st.caption("위 코드를 복사해서 저장해두세요!")
                if st.button("이 코드로 등록 요청", key="btn_new_device"):
                    st.session_state['browser_device_id'] = new_code
                    save_device_id_to_browser(new_code)
                    add_device_to_pending(new_code)
                    st.rerun()

            # 관리자 자동 승인
            st.markdown("---")
            st.caption("👑 관리자이신가요?")
            admin_pw_first = st.text_input("관리자 비밀번호", type="password", key="admin_pw_first", placeholder="관리자 비밀번호 입력...")
            if st.button("🔓 관리자로 바로 접속", key="btn_admin_first"):
                if admin_pw_first == ADMIN_PASSWORD:
                    new_code = st.session_state['new_device_code']
                    st.session_state['browser_device_id'] = new_code
                    save_device_id_to_browser(new_code)
                    add_device_to_approved(new_code)
                    st.success("✅ 관리자 기기로 승인되었습니다!")
                    st.rerun()
                else:
                    st.error("관리자 비밀번호가 틀렸습니다.")

        # 브라우저 ID가 있으면 승인 여부 확인
        elif is_device_approved_jsonbin():
            # 로컬 파일에서 저장된 비밀번호 확인
            saved_pw = saved_config.get('password', '')
            if saved_pw and saved_pw == CORRECT_PASSWORD:
                st.session_state['authenticated'] = True
                # 로컬 파일에서 API 키 로드
                if saved_config.get('api_key'):
                    st.session_state['api_key'] = saved_config.get('api_key')
                st.rerun()

            pw = st.text_input("비밀번호", type="password", key="pw_login", placeholder="Enter password...")
            remember_pw = st.checkbox("비밀번호 저장", value=True, key="remember_pw")
            if st.button("입장", key="btn_login"):
                if pw == CORRECT_PASSWORD:
                    st.session_state['authenticated'] = True
                    if remember_pw:
                        # 로컬 파일에 로그인 정보 저장
                        save_config({
                            'authenticated': True,
                            'password': pw,
                            'device_id': get_device_id()
                        })
                    st.rerun()
                else:
                    st.error("비밀번호가 올바르지 않습니다")
        else:
            # 승인되지 않은 기기
            st.warning(f"🖥️ 승인 대기 중\n\n기기 ID: `{device_id}`")
            st.info("관리자에게 위 기기 ID를 전달하고 승인을 요청하세요.")

            # 대기 목록에 추가
            add_device_to_pending(device_id)

            if st.button("🔄 승인 상태 확인", key="btn_check_approval"):
                if is_device_approved_jsonbin():
                    st.success("✅ 승인되었습니다! 페이지를 새로고침합니다.")
                    st.rerun()
                else:
                    st.warning("아직 승인되지 않았습니다. 관리자에게 문의하세요.")

            # 관리자 자동 승인
            st.markdown("---")
            st.caption("👑 관리자이신가요?")
            admin_pw = st.text_input("관리자 비밀번호", type="password", key="admin_pw_approve", placeholder="관리자 비밀번호 입력...")
            if st.button("🔓 관리자로 승인", key="btn_admin_approve"):
                if admin_pw == ADMIN_PASSWORD:
                    add_device_to_approved(device_id)
                    save_device_id_to_browser(device_id)
                    st.success("✅ 관리자 기기로 승인되었습니다!")
                    st.rerun()
                else:
                    st.error("관리자 비밀번호가 틀렸습니다.")

            # 다른 기기 코드로 변경
            if st.button("🔄 다른 코드로 변경", key="btn_change_device"):
                del st.session_state['browser_device_id']
                st.rerun()
    st.stop()

# 세션 초기화
defaults = {
    'topic': '', 'target_persona': '', 'pain_points': '',
    'outline': [], 'chapters': {}, 'book_title': '', 'subtitle': '',
    'score_details': None, 'generated_titles': None, 'suggested_targets': None,
    'analyzed_pains': None, 'review_analysis': None, 'market_gaps': None,
    'knowledge_hub': [], 'study_summary': None, 'current_page': 0,
    'recommended_refs': None, 'generated_ideas': None,
    # 인터뷰 관련 변수
    'interview_completed': False,
    'interview_data': {},
    'author_name': '',
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# 사이드바
with st.sidebar:
    # API 키 섹션 (접기/펼치기 가능)
    if 'show_api_section' not in st.session_state:
        st.session_state['show_api_section'] = True

    if 'api_key' not in st.session_state:
        # JSONBin에서 API 키 불러오기
        device_id = get_device_id()
        if device_id:
            settings = get_device_settings(device_id)
            st.session_state['api_key'] = settings.get('api_key', '')
        else:
            st.session_state['api_key'] = ''

    # API 키가 입력되어 있으면 기본적으로 접힌 상태로
    api_key_exists = bool(st.session_state['api_key'])

    col_title, col_toggle = st.columns([4, 1])
    with col_title:
        st.markdown("### 🔑 Claude API 키")
    with col_toggle:
        toggle_label = "▼" if st.session_state['show_api_section'] else "▶"
        if st.button(toggle_label, key="toggle_api_section", help="접기/펼치기"):
            st.session_state['show_api_section'] = not st.session_state['show_api_section']
            st.rerun()

    if st.session_state['show_api_section']:
        api_key = st.text_input("키 입력", value=st.session_state['api_key'], type="password", key="api_sidebar", label_visibility="collapsed", placeholder="sk-ant-api03-... 형식")
        if api_key != st.session_state['api_key']:
            st.session_state['api_key'] = api_key
            # JSONBin에 저장
            device_id = get_device_id()
            if device_id:
                save_device_settings(device_id, api_key=api_key)

        if api_key:
            st.success("✅ Claude 키 입력 완료!")
        else:
            st.error("⚠️ Claude API 키를 입력하세요")
    else:
        # 접힌 상태에서 간단한 상태 표시
        if st.session_state['api_key']:
            st.caption("✅ API 키 설정됨")
        else:
            st.caption("⚠️ API 키 필요")

    # 모델 선택
    st.markdown("### 🤖 모델 선택")
    if 'claude_model' not in st.session_state:
        st.session_state['claude_model'] = "claude-sonnet-4-20250514"

    model_options = {
        "Claude Sonnet 4 (추천)": "claude-sonnet-4-20250514",
        "Claude Sonnet 3.5 v2": "claude-3-5-sonnet-20241022",
        "Claude Haiku 3.5 (저렴)": "claude-3-5-haiku-20241022"
    }
    selected_model = st.selectbox(
        "모델 선택",
        options=list(model_options.keys()),
        index=0,
        label_visibility="collapsed"
    )
    st.session_state['claude_model'] = model_options[selected_model]

    if "Haiku" in selected_model:
        st.info("💰 저렴하지만 품질이 낮을 수 있음")
    elif "Sonnet 4" in selected_model:
        st.info("💎 최신 모델, 최고 품질")
    else:
        st.info("⚡ 안정적인 성능")

    # API 키 발급 방법 안내
    with st.expander("📖 Claude API 키 발급 방법 (상세)", expanded=False):
        st.markdown("""
        ### 🟣 1단계: Anthropic 회원가입

        1. 아래 버튼을 클릭하세요
        2. **"Sign up"** 클릭
        3. Google 계정 또는 이메일로 가입
        """)
        st.link_button("🔗 Anthropic 가입 페이지", "https://console.anthropic.com/", use_container_width=True)

        st.markdown("""
        ---
        ### 💳 2단계: 결제 수단 등록

        1. 로그인 후 왼쪽 메뉴에서 **"Settings"** 클릭
        2. **"Billing"** 클릭
        3. **"Add payment method"** 클릭
        4. 카드 정보 입력 후 저장
        5. **"Add credits"**로 크레딧 충전 ($5~10 추천)
        """)
        st.link_button("🔗 Billing 페이지 바로가기", "https://console.anthropic.com/settings/billing", use_container_width=True)

        st.markdown("""
        ---
        ### 🔑 3단계: API 키 발급

        1. 왼쪽 메뉴에서 **"API Keys"** 클릭
        2. **"Create Key"** 버튼 클릭
        3. 이름 입력 (예: ebook)
        4. **"Create Key"** 클릭
        5. 생성된 키 **복사** (sk-ant-api03-... 형식)
        6. 위 입력창에 **붙여넣기**
        """)
        st.link_button("🔗 API Keys 페이지 바로가기", "https://console.anthropic.com/settings/keys", use_container_width=True)

        st.markdown("---")
        st.warning("⚠️ API 키는 한 번만 보여줍니다. 복사해두세요!")
        st.success("💰 예상 비용: 전자책 1권당 약 200~500원")

    st.markdown("---")
    st.markdown("### 📊 진행 상황")
    progress = sum([bool(st.session_state['topic']), bool(st.session_state['target_persona']), bool(st.session_state['outline']), len(st.session_state['chapters']) > 0]) / 4
    st.progress(progress)

    st.markdown("---")
    st.markdown("### 🚀 빠른 이동")
    sidebar_pages = ["① 주제", "② 목차", "③ 본문", "④ 완성"]
    sidebar_mapping = [0, 4, 5, 7]
    for i, p in enumerate(sidebar_pages):
        if st.button(p, key=f"sidebar_nav_{i}", use_container_width=True):
            st.session_state['current_page'] = sidebar_mapping[i]
            st.rerun()

    # 관리자 메뉴 (기기 승인 관리)
    st.markdown("---")
    with st.expander("🔐 관리자 설정", expanded=False):
        if 'admin_logged_in' not in st.session_state:
            st.session_state['admin_logged_in'] = False

        if not st.session_state['admin_logged_in']:
            admin_pw = st.text_input("관리자 비밀번호", type="password", key="admin_pw", placeholder="관리자 비밀번호...")
            if st.button("로그인", key="btn_admin_login"):
                if admin_pw == ADMIN_PASSWORD:
                    st.session_state['admin_logged_in'] = True
                    st.rerun()
                else:
                    st.error("비밀번호가 틀렸습니다")
        else:
            st.success("✅ 관리자 모드")

            # 대기 중인 기기 목록
            st.markdown("**📋 승인 대기 중:**")
            pending = get_pending_devices()
            if pending:
                for pid in pending:
                    col_a, col_b = st.columns([3, 1])
                    with col_a:
                        st.code(pid, language=None)
                    with col_b:
                        if st.button("✅", key=f"approve_{pid}"):
                            if add_device_to_approved(pid):
                                st.success(f"승인됨!")
                                st.rerun()
            else:
                st.caption("대기 중인 기기 없음")

            st.markdown("---")

            # 승인된 기기 목록
            st.markdown("**✅ 승인된 기기:**")
            approved = get_approved_devices()
            if approved:
                for aid in approved:
                    col_a, col_b = st.columns([3, 1])
                    with col_a:
                        st.code(aid, language=None)
                    with col_b:
                        if st.button("❌", key=f"remove_{aid}"):
                            if remove_device_from_approved(aid):
                                st.warning(f"삭제됨!")
                                st.rerun()
            else:
                st.caption("승인된 기기 없음")

            st.markdown("---")

            # 수동 기기 추가
            new_device_id = st.text_input("기기 ID 직접 추가", key="manual_device_id", placeholder="기기 ID 입력...")
            if st.button("➕ 추가", key="btn_add_device"):
                if new_device_id.strip():
                    if add_device_to_approved(new_device_id.strip()):
                        st.success("추가됨!")
                        st.rerun()

            if st.button("🚪 관리자 로그아웃", key="btn_admin_logout"):
                st.session_state['admin_logged_in'] = False
                st.rerun()

    # 사이드바 하단 제작자 정보
    st.markdown("---")
    st.markdown("""
    <div style="text-align:center; padding:10px 0; color:#d4af37 !important; font-size:12px;">
        <strong>CASHMAKER</strong><br>
        <span style="color:#ffffff !important;">제작: 남현우 작가</span>
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 헬퍼 함수
# ==========================================
def get_api_key():
    return st.session_state.get('api_key', '')

def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*([^*]+)\*\*', r'「\1」', text)
    text = text.replace('**', '').replace('*', '').replace('###', '').replace('##', '').replace('#', '')
    return text.strip()

def clean_content(text):
    if not text:
        return ""
    # 마크다운 제거
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    # **굵은글씨** 패턴 완전 제거
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = text.replace('**', '').replace('*', '').replace('###', '').replace('##', '').replace('#', '')
    # 연속 줄바꿈 정리
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def parse_json(response):
    """JSON 파싱 - 개선된 에러 처리"""
    if not response:
        return None
    try:
        # 먼저 전체 응답에서 JSON 블록 찾기
        json_match = re.search(r'```json\s*([\s\S]*?)\s*```', response)
        if json_match:
            return json.loads(json_match.group(1))

        # JSON 블록이 없으면 중괄호로 시작하는 객체 찾기
        match = re.search(r'\{[\s\S]*\}', response)
        if match:
            json_str = match.group()
            # 불완전한 JSON 수정 시도
            json_str = re.sub(r',\s*}', '}', json_str)  # 마지막 쉼표 제거
            json_str = re.sub(r',\s*]', ']', json_str)  # 배열 마지막 쉼표 제거
            return json.loads(json_str)
    except json.JSONDecodeError as e:
        st.warning(f"JSON 파싱 경고: {str(e)[:50]}")
    except Exception as e:
        st.warning(f"파싱 오류: {str(e)[:50]}")
    return None

def ask_ai(prompt, temp=0.7):
    """Claude API 호출"""
    api_key = get_api_key()
    if not api_key:
        st.error("Claude API 키를 입력해주세요")
        return None

    if not CLAUDE_AVAILABLE:
        st.error("anthropic 패키지가 설치되지 않았습니다. pip install anthropic")
        return None

    # 선택된 모델 가져오기 (기본값: Sonnet 4)
    model = st.session_state.get('claude_model', 'claude-sonnet-4-20250514')

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model=model,
            max_tokens=8000,
            temperature=temp,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return message.content[0].text
    except anthropic.AuthenticationError:
        st.error("API 키가 유효하지 않습니다. Claude API 키를 확인해주세요.")
        return None
    except anthropic.RateLimitError:
        st.error("API 할당량이 초과되었습니다. 잠시 후 다시 시도해주세요.")
        return None
    except anthropic.BadRequestError as e:
        st.error(f"요청 오류: {str(e)[:100]}")
        return None
    except Exception as e:
        st.error(f"AI 오류: {str(e)[:100]}")
        return None

def generate_cover_image_gemini(title, subtitle, theme_keywords):
    """Google Gemini로 표지 배경 이미지 생성"""

    api_key = get_api_key()
    if not api_key:
        return None, "Gemini API 키가 필요합니다."

    if not IMAGEN_AVAILABLE:
        return None, "google-genai 패키지가 필요합니다: pip install google-genai"

    try:
        client = google_genai.Client(api_key=api_key)

        # 베스트셀러급 고급 표지 프롬프트 - 텍스트 절대 금지 강조
        prompt = f"""Create an ABSTRACT background image for a book cover.

Theme keywords: {theme_keywords}

STYLE: Dark, moody, cinematic atmosphere. Abstract shapes, gradients, smoke, light rays, or geometric patterns. Luxury aesthetic with gold/amber accent lighting on deep black background.

CRITICAL RULES:
- ONLY abstract visuals: smoke, light, shadows, gradients, textures
- NO objects, NO people, NO faces, NO hands
- NO text, NO letters, NO words, NO numbers, NO symbols, NO characters of ANY language
- NO Korean, NO English, NO Chinese, NO Japanese characters
- Pure abstract art only

OUTPUT: Dark dramatic background with subtle golden light accents, suitable for text overlay."""

        # Gemini 이미지 생성
        response = client.models.generate_content(
            model='gemini-2.0-flash-exp-image-generation',
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                response_modalities=['IMAGE', 'TEXT']
            )
        )

        if response.candidates and response.candidates[0].content.parts:
            for part in response.candidates[0].content.parts:
                if hasattr(part, 'inline_data') and part.inline_data:
                    image_base64 = base64.b64encode(part.inline_data.data).decode('utf-8')
                    return image_base64, None

        return None, "이미지가 생성되지 않았습니다."

    except Exception as e:
        error_msg = str(e)
        if "quota" in error_msg.lower() or "limit" in error_msg.lower():
            return None, "API 할당량 초과. 잠시 후 다시 시도해주세요."
        elif "safety" in error_msg.lower():
            return None, "안전 필터에 의해 차단되었습니다. 다른 키워드로 시도해주세요."
        return None, f"이미지 생성 오류: {error_msg[:80]}"

def generate_cover_prompt_ai(title, subtitle, topic):
    """AI가 표지 디자인 컨셉과 이미지 프롬프트 생성"""
    prompt = f"""당신은 베스트셀러 책 표지 디자이너입니다.

책 제목: {title}
부제: {subtitle}
주제: {topic}

이 책의 표지 이미지를 위한 영문 프롬프트를 만들어주세요.

[요구사항]
1. 실제 베스트셀러 표지 스타일 분석 기반
2. 제목의 핵심 메시지를 시각적으로 표현
3. 고급스럽고 전문적인 느낌
4. 텍스트 오버레이를 위한 여백 고려
5. 추상적이거나 상징적인 이미지

[출력 형식]
IMAGE_PROMPT: (영문 이미지 생성 프롬프트, 50단어 이내)
COLOR_SCHEME: (추천 컬러 팔레트, 예: dark, gold, minimal)
STYLE: (디자인 스타일, 예: editorial, bold, elegant)

영문 프롬프트만 출력하세요. 한국어 설명 불필요."""

    result = ask_ai(prompt, temp=0.7)
    if result:
        # 파싱
        image_prompt = ""
        color_scheme = "dark"
        style = "editorial"

        for line in result.split('\n'):
            if 'IMAGE_PROMPT:' in line:
                image_prompt = line.split('IMAGE_PROMPT:')[-1].strip()
            elif 'COLOR_SCHEME:' in line:
                color_scheme = line.split('COLOR_SCHEME:')[-1].strip().lower()
            elif 'STYLE:' in line:
                style = line.split('STYLE:')[-1].strip().lower()

        return image_prompt, color_scheme, style
    return None, "dark", "editorial"

def extract_video_id(url):
    """YouTube URL에서 video ID 추출"""
    patterns = [
        r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([^&\n?#]+)',
        r'youtube\.com\/watch\?.*v=([^&\n?#]+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def get_youtube_transcript(video_id):
    """YouTube 자막 가져오기"""
    if not YOUTUBE_TRANSCRIPT_AVAILABLE:
        return None, "youtube-transcript-api가 설치되지 않았습니다. pip install youtube-transcript-api"

    try:
        # 한국어 자막 우선, 없으면 영어, 없으면 자동생성
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

        transcript = None
        # 수동 자막 먼저 시도
        for lang in ['ko', 'en']:
            try:
                transcript = transcript_list.find_transcript([lang])
                break
            except:
                continue

        # 수동 자막 없으면 자동 생성 자막
        if not transcript:
            try:
                transcript = transcript_list.find_generated_transcript(['ko', 'en'])
            except:
                # 아무 자막이나 가져오기
                for t in transcript_list:
                    transcript = t
                    break

        if transcript:
            fetched = transcript.fetch()
            full_text = ' '.join([item['text'] for item in fetched])
            return full_text, None
        else:
            return None, "자막을 찾을 수 없습니다"

    except Exception as e:
        return None, f"자막 추출 오류: {str(e)[:100]}"

def analyze_youtube_video_direct(url):
    """YouTube 영상 자막 기반 분석 (빠르고 정확)"""
    api_key = get_api_key()
    if not api_key:
        st.error("API 키를 입력해주세요")
        return None

    # 1. Video ID 추출
    video_id = extract_video_id(url)
    if not video_id:
        st.error("올바른 YouTube URL이 아닙니다")
        return None

    # 2. 자막 가져오기
    transcript, error = get_youtube_transcript(video_id)
    if error:
        st.warning(f"자막 추출 실패: {error}")
        st.info("자막이 없는 영상입니다. 텍스트 입력으로 직접 내용을 입력해주세요.")
        return None

    if not transcript or len(transcript) < 50:
        st.warning("자막 내용이 너무 짧습니다")
        return None

    # 3. 자막 기반 분석
    prompt = f"""다음은 YouTube 영상의 자막입니다. 이 내용을 분석해주세요.

[자막 내용]
{transcript[:15000]}

[분석 요청]
위 자막 내용을 바탕으로 분석해주세요. 자막에 없는 내용은 추측하지 마세요.

JSON 형식으로 응답:
{{
    "title": "영상의 핵심 주제 (자막 기반 추론)",
    "creator": "알 수 없음",
    "main_topic": "메인 주제 한 줄 요약",
    "key_points": ["핵심 포인트 1", "핵심 포인트 2", "핵심 포인트 3", "핵심 포인트 4", "핵심 포인트 5"],
    "detailed_notes": ["상세 내용 1", "상세 내용 2", "상세 내용 3"],
    "actionable_tips": ["실천 팁 1", "실천 팁 2", "실천 팁 3"],
    "quotes": ["인상적인 문장 1", "인상적인 문장 2"],
    "vocabulary": [{{"term": "용어", "definition": "설명"}}],
    "study_questions": ["학습 질문 1", "학습 질문 2"],
    "summary": "전체 내용 5-7문장 요약"
}}"""

    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"분석 오류: {str(e)[:150]}")
        return None

def get_full_content():
    full = ""
    for ch in st.session_state.get('outline', []):
        if ch in st.session_state.get('chapters', {}):
            ch_data = st.session_state['chapters'][ch]
            ch_content = ""
            for s in ch_data.get('subtopics', []):
                c = ch_data.get('subtopic_data', {}).get(s, {}).get('content', '')
                if c:
                    ch_content += f"\n\n【{s}】\n\n{clean_content(c)}"
            if ch_content:
                full += f"\n\n{'='*50}\n{ch}\n{'='*50}{ch_content}"
    return full.strip()

def create_ebook_docx(title, subtitle, author, chapters_data, outline, interview_data=None):
    """베스트셀러 스타일의 전문적인 워드 문서 생성"""
    if not DOCX_AVAILABLE:
        return None, "python-docx 패키지가 필요합니다: pip install python-docx"

    try:
        doc = Document()

        # 페이지 설정 (A5 크기 - 전자책에 적합)
        section = doc.sections[0]
        section.page_width = Cm(14.8)
        section.page_height = Cm(21)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        def set_font(run, size, bold=False, color=None, italic=False):
            run.font.size = Pt(size)
            run.font.name = 'Malgun Gothic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = RGBColor(*color)

        def add_bookmark(paragraph, bookmark_name):
            """문단에 북마크 추가"""
            # 북마크 이름에서 특수문자 제거 (Word 북마크 규칙)
            clean_name = re.sub(r'[^\w가-힣]', '_', bookmark_name)[:40]

            bookmark_start = OxmlElement('w:bookmarkStart')
            bookmark_start.set(qn('w:id'), str(hash(clean_name) % 10000))
            bookmark_start.set(qn('w:name'), clean_name)

            bookmark_end = OxmlElement('w:bookmarkEnd')
            bookmark_end.set(qn('w:id'), str(hash(clean_name) % 10000))

            paragraph._p.insert(0, bookmark_start)
            paragraph._p.append(bookmark_end)
            return clean_name

        def add_hyperlink(paragraph, text, bookmark_name, font_size=10, bold=False, color=(70, 70, 70)):
            """북마크로 연결되는 하이퍼링크 추가"""
            # 북마크 이름 정리
            clean_name = re.sub(r'[^\w가-힣]', '_', bookmark_name)[:40]

            # 하이퍼링크 요소 생성
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), clean_name)

            # 텍스트 실행 요소
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            # 폰트 설정
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Malgun Gothic')
            rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            rFonts.set(qn('w:hAnsi'), 'Malgun Gothic')
            rPr.append(rFonts)

            # 폰트 크기
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(font_size * 2))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(font_size * 2))
            rPr.append(szCs)

            # 볼드
            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)

            # 색상
            if color:
                c = OxmlElement('w:color')
                c.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color))
                rPr.append(c)

            new_run.append(rPr)

            # 텍스트
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            new_run.append(text_elem)

            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

            return hyperlink

        # ══════════════════════════════════════════════════════════════
        # 표지 페이지 (미니멀 고급 스타일)
        # ══════════════════════════════════════════════════════════════
        for _ in range(8):
            doc.add_paragraph()

        # 메인 타이틀
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        set_font(title_run, 28, bold=True)
        title_para.paragraph_format.space_after = Pt(16)

        # 부제
        if subtitle:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = subtitle_para.add_run(subtitle)
            set_font(sub_run, 12, color=(80, 80, 80))
            subtitle_para.paragraph_format.space_before = Pt(8)

        # 저자명
        for _ in range(10):
            doc.add_paragraph()
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(author if author else "저자")
        set_font(author_run, 13)

        doc.add_page_break()

        # ══════════════════════════════════════════════════════════════
        # 판권 페이지
        # ══════════════════════════════════════════════════════════════
        for _ in range(18):
            doc.add_paragraph()

        copyright_lines = [
            f"{title}",
            "",
            f"지은이: {author if author else '저자'}",
            "",
            "이 책의 저작권은 저자에게 있습니다.",
            "무단 전재와 복제를 금합니다."
        ]

        for line in copyright_lines:
            cp_para = doc.add_paragraph()
            cp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if line:
                cp_run = cp_para.add_run(line)
                set_font(cp_run, 9, color=(120, 120, 120))
            cp_para.paragraph_format.space_after = Pt(2)

        doc.add_page_break()

        # ══════════════════════════════════════════════════════════════
        # 프롤로그 (미니멀 스타일)
        # ══════════════════════════════════════════════════════════════
        for _ in range(4):
            doc.add_paragraph()

        # 프롤로그 제목
        prologue_title = doc.add_paragraph()
        prologue_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt_run = prologue_title.add_run("Prologue")
        set_font(pt_run, 14, bold=True)
        prologue_title.paragraph_format.space_after = Pt(30)

        # 프롤로그 내용 - AI가 인터뷰 내용을 참고해서 자연스럽게 작성
        prologue_text = None
        if interview_data:
            prologue_prompt = f"""당신은 자청 스타일로 글을 쓰는 베스트셀러 작가입니다. 프롤로그를 작성하세요.

[저자 정보 - 참고용, 그대로 복사하지 말 것]
- 분야: {interview_data.get('field', '')}
- 경력: {interview_data.get('experience_years', '')}
- 책 주제: {interview_data.get('topic', '')}
- 타겟 독자: {interview_data.get('target_reader', '')}
- 독자의 고민: {interview_data.get('target_problem', '')}
- 집필 동기: {interview_data.get('why_write', '')}

[프롤로그 작성 원칙]
1. 나의 실패담이나 솔직한 고백으로 시작
2. "저도 처음엔 몰랐습니다" 공감
3. 이 책에서 뭘 얻어갈 수 있는지 힌트
4. 짧은 문장 (한 문장에 생각 하나)
5. 짧은 문단 (2-4문장 MAX)
6. 구어체 + 합쇼체 ("~거든요", "~잖아요" OK)

[분량] 400-600자

[금지 - 절대 쓰지 말 것]
- 위 저자 정보를 그대로 복사 붙여넣기
- 교과서 표현: "여정", "발걸음", "함께 하시길 바랍니다", "진정한"
- AI 표현: "~의 중요성", "다양한", "효과적인", "~를 통해"
- 과장: "놀라운", "혁신적인", "충격적인"
- 뻔한 말: "포기하지 마세요", "꾸준히 하세요"
- 마크다운 문법

프롤로그만 출력하세요."""

            generated_prologue = ask_ai(prologue_prompt, 0.7)
            if generated_prologue:
                prologue_text = generated_prologue

        if not prologue_text:
            prologue_text = """이 책을 쓰게 된 이유는 단순합니다.

제가 직접 경험하고 배운 것들을 나누고 싶었습니다.

처음에는 저도 막막했습니다. 하지만 포기하지 않았고, 결국 방법을 찾았습니다.

이 책은 단순한 이론서가 아닙니다. 직접 해보고, 실패하고, 다시 일어나며 터득한 실전 노하우입니다.

당신도 할 수 있습니다.

자, 이제 시작합니다."""

        # **프롤로그** 같은 마크다운 제목 제거
        prologue_text = prologue_text.replace('**프롤로그**', '').replace('**Prologue**', '').strip()

        for para_text in prologue_text.split('\n\n'):
            if para_text.strip():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 왼쪽 정렬 (본문과 동일)
                para_run = para.add_run(para_text.strip())
                set_font(para_run, 10)
                para_format = para.paragraph_format
                para_format.line_spacing = 1.6
                # 문단 사이 간격 (한 줄 띄우기 효과)
                para_format.space_after = Pt(14)

        doc.add_page_break()

        # ══════════════════════════════════════════════════════════════
        # 목차 (프리미엄 미니멀 디자인)
        # ══════════════════════════════════════════════════════════════

        # 상단 여백
        for _ in range(4):
            doc.add_paragraph()

        # 목차 제목 (미니멀 타이포그래피)
        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_title.add_run("CONTENTS")
        set_font(toc_run, 16, bold=False, color=(40, 40, 40))
        toc_title.paragraph_format.space_after = Pt(60)

        for idx, chapter in enumerate(outline):
            # 챕터 제목 정리 (PART X. 등 접두사 제거)
            clean_chapter = chapter
            for prefix in [f"PART {idx + 1}.", f"PART{idx + 1}.", f"PART {idx + 1} ", f"PART{idx + 1} ", f"{idx + 1}.", f"{idx + 1})"]:
                clean_chapter = clean_chapter.replace(prefix, "").strip()

            # ─────────────────────────────────────
            # 챕터 번호 (큰 숫자)
            # ─────────────────────────────────────
            ch_num_para = doc.add_paragraph()
            ch_num_para.paragraph_format.space_before = Pt(28)
            ch_num_para.paragraph_format.space_after = Pt(4)
            ch_num_run = ch_num_para.add_run(f"{idx + 1:02d}")
            set_font(ch_num_run, 24, bold=False, color=(200, 200, 200))

            # ─────────────────────────────────────
            # 챕터 제목 (하이퍼링크)
            # ─────────────────────────────────────
            ch_title_para = doc.add_paragraph()
            ch_title_para.paragraph_format.space_after = Pt(14)
            chapter_bookmark_name = f"chapter_{idx + 1}"
            add_hyperlink(ch_title_para, clean_chapter, chapter_bookmark_name, font_size=12, bold=True, color=(30, 30, 30))

            # ─────────────────────────────────────
            # 소제목들 (심플한 리스트)
            # ─────────────────────────────────────
            if chapter in chapters_data:
                ch_data = chapters_data[chapter]
                subtopics = ch_data.get('subtopics', [])

                for sub_idx, sub in enumerate(subtopics):
                    toc_sub = doc.add_paragraph()
                    toc_sub.paragraph_format.left_indent = Cm(0.3)
                    toc_sub.paragraph_format.space_after = Pt(6)

                    # 작은 점 불릿
                    bullet_run = toc_sub.add_run("·  ")
                    set_font(bullet_run, 10, color=(180, 180, 180))

                    # 소제목 텍스트 (하이퍼링크로 연결)
                    subtopic_bookmark_name = f"subtopic_{idx + 1}_{sub_idx + 1}"
                    add_hyperlink(toc_sub, sub, subtopic_bookmark_name, font_size=10, bold=False, color=(80, 80, 80))

        # 하단 여백
        for _ in range(3):
            doc.add_paragraph()

        doc.add_page_break()

        # ══════════════════════════════════════════════════════════════
        # 본문 (프리미엄 에디토리얼 스타일)
        # ══════════════════════════════════════════════════════════════

        def add_horizontal_line(doc, width_cm=3, color=(220, 220, 220)):
            """가로 구분선 추가"""
            line_para = doc.add_paragraph()
            line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line_run = line_para.add_run("─" * 12)
            set_font(line_run, 10, color=color)
            line_para.paragraph_format.space_before = Pt(20)
            line_para.paragraph_format.space_after = Pt(20)
            return line_para

        def add_chapter_opener(doc, idx, chapter_title):
            """챕터 시작 페이지 - 프리미엄 에디토리얼 스타일"""
            # 상단 넓은 여백 (페이지 중앙 정도에 위치)
            for _ in range(7):
                doc.add_paragraph()

            # PART 라벨 (작은 대문자)
            part_label = doc.add_paragraph()
            part_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            part_run = part_label.add_run(f"P A R T")
            set_font(part_run, 9, color=(160, 160, 160))
            part_label.paragraph_format.space_after = Pt(8)

            # 챕터 번호 (매우 큰 숫자)
            ch_num_para = doc.add_paragraph()
            ch_num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ch_num_run = ch_num_para.add_run(f"{idx + 1}")
            set_font(ch_num_run, 48, bold=False, color=(40, 40, 40))
            ch_num_para.paragraph_format.space_after = Pt(16)

            # 구분선
            line_para = doc.add_paragraph()
            line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line_run = line_para.add_run("───────────")
            set_font(line_run, 10, color=(200, 200, 200))
            line_para.paragraph_format.space_after = Pt(20)

            # 챕터 제목
            ch_name = doc.add_paragraph()
            ch_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chn_run = ch_name.add_run(chapter_title)
            set_font(chn_run, 14, bold=True, color=(30, 30, 30))
            ch_name.paragraph_format.space_after = Pt(60)

            return ch_name

        def add_subtopic_header(doc, subtopic_text, sub_idx):
            """소제목 - 베스트셀러 스타일"""
            # 소제목 전 넓은 여백 (시각적 구분)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(40)

            # 상단 미니멀 구분선
            line_para = doc.add_paragraph()
            line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            line_run = line_para.add_run("──")
            set_font(line_run, 10, color=(200, 200, 200))
            line_para.paragraph_format.space_after = Pt(12)

            # 소제목 텍스트 (대담하고 명확하게)
            sub_title = doc.add_paragraph()
            sub_run = sub_title.add_run(subtopic_text)
            set_font(sub_run, 13, bold=True, color=(25, 25, 25))
            sub_title.paragraph_format.space_after = Pt(24)

            return sub_title

        def format_body_paragraph(doc, text, is_first=False):
            """본문 문단 - 베스트셀러 가독성 스타일"""
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if is_first and len(text) > 0:
                # 첫 문단 첫 글자 드롭캡 (세련된 버전)
                first_char = text[0]
                rest_text = text[1:]

                first_run = para.add_run(first_char)
                set_font(first_run, 18, bold=True, color=(40, 40, 40))

                rest_run = para.add_run(rest_text)
                set_font(rest_run, 10.5, color=(45, 45, 45))
            else:
                p_run = para.add_run(text)
                set_font(p_run, 10.5, color=(45, 45, 45))

            para_format = para.paragraph_format
            para_format.line_spacing = 1.85  # 베스트셀러 표준 줄간격
            para_format.space_after = Pt(14)
            para_format.first_line_indent = Cm(0.6)  # 들여쓰기

            return para

        def is_table_text(text):
            """텍스트가 표 형식인지 감지"""
            lines = text.strip().split('\n')
            if len(lines) < 2:
                return False

            # 마크다운 테이블 감지: | 로 시작하고 | 로 끝남
            pipe_lines = sum(1 for line in lines if line.strip().startswith('|') and line.strip().endswith('|'))
            if pipe_lines >= 2:
                return True

            # 파이프로 구분된 테이블 (| 가 있지만 시작/끝이 아닐 수 있음)
            pipe_content_lines = sum(1 for line in lines if '|' in line and len(line.split('|')) >= 2)
            if pipe_content_lines >= 2:
                return True

            # 탭 구분 테이블
            tab_lines = sum(1 for line in lines if '\t' in line)
            if tab_lines >= 2:
                return True

            # 콜론 기반 비교 테이블 감지 (Before: xxx / After: xxx)
            colon_lines = sum(1 for line in lines
                            if ':' in line
                            and not line.strip().startswith('http')
                            and len(line.split(':')[0]) < 30)
            if colon_lines >= 2 and colon_lines >= len(lines) * 0.6:
                return True

            return False

        def parse_table_data(text):
            """표 텍스트를 파싱하여 2D 배열로 변환"""
            lines = text.strip().split('\n')
            table_data = []

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # 순수 구분선만 스킵 (실제 내용이 없는 라인)
                # 마크다운 구분선: |---|---| 또는 |:---:|:---:|
                if re.match(r'^\|[\s\-:]+\|$', line):
                    continue
                # 박스 그리기 문자만 있는 라인
                if re.match(r'^[─━┌┬┐├┼┤└┴┘│┃]+$', line):
                    continue
                # 하이픈만 있는 라인 (--- 또는 - - -)
                if re.match(r'^[\s\-]+$', line) and len(line.replace(' ', '').replace('-', '')) == 0:
                    continue

                # 마크다운 테이블 파싱 (| cell | cell |)
                if line.startswith('|') and line.endswith('|'):
                    cells = [cell.strip() for cell in line.split('|')]
                    cells = [c for c in cells if c]  # 빈 셀 제거
                    if cells:
                        table_data.append(cells)
                # 일반 파이프 구분 (cell | cell)
                elif '|' in line and not line.startswith('|'):
                    cells = [cell.strip() for cell in line.split('|')]
                    cells = [c for c in cells if c]
                    if cells:
                        table_data.append(cells)
                # 탭 구분 테이블
                elif '\t' in line:
                    cells = [cell.strip() for cell in line.split('\t')]
                    cells = [c for c in cells if c]
                    if len(cells) >= 2:
                        table_data.append(cells)
                # 콜론 기반 파싱 (Before: xxx) - 단, URL이 아닌 경우
                elif ':' in line and not line.startswith('http'):
                    # 첫 번째 콜론으로만 분리
                    parts = line.split(':', 1)
                    if len(parts) == 2 and len(parts[0]) < 30:  # 키가 너무 길면 제외
                        table_data.append([parts[0].strip(), parts[1].strip()])

            return table_data

        def add_premium_table(doc, table_data):
            """인포그래픽 스타일 테이블 - 시각적 이해도 향상"""
            if not table_data or len(table_data) < 1:
                return None

            rows = len(table_data)
            cols = max(len(row) for row in table_data)

            # 2열 비교 테이블인 경우 (Before/After, 항목/설명 등)
            is_comparison = cols == 2 and rows >= 2

            # 테이블 생성
            table = doc.add_table(rows=rows, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 테이블 전체 너비 설정
            table.autofit = False
            for col_idx in range(cols):
                for row in table.rows:
                    if col_idx < len(row.cells):
                        if is_comparison:
                            # 2열 비교: 첫 열 좁게, 둘째 열 넓게
                            width = Cm(3) if col_idx == 0 else Cm(7)
                        else:
                            width = Cm(10 / cols)
                        row.cells[col_idx].width = width

            # 각 셀 스타일링
            for i, row_data in enumerate(table_data):
                row = table.rows[i]
                row.height = Cm(1.0)  # 행 높이 증가

                for j, cell_text in enumerate(row_data):
                    if j < cols:
                        cell = row.cells[j]
                        cell.text = ''

                        para = cell.paragraphs[0]

                        # 첫 번째 행(헤더) - 진한 배경
                        if i == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run(str(cell_text))
                            set_font(run, 9, bold=True, color=(255, 255, 255))
                            # 헤더 배경색 (진한 회색)
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), '4A4A4A')
                            cell._tc.get_or_add_tcPr().append(shading)

                        # 첫 번째 열 (라벨/항목) - 2열 비교 테이블
                        elif is_comparison and j == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run(str(cell_text))
                            set_font(run, 9, bold=True, color=(50, 50, 50))
                            # 연한 배경
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), 'F8F8F8')
                            cell._tc.get_or_add_tcPr().append(shading)

                        # 일반 내용 셀
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            para.paragraph_format.left_indent = Pt(8)
                            run = para.add_run(str(cell_text))
                            set_font(run, 9, color=(60, 60, 60))
                            # 짝수 행 배경 (줄무늬 효과)
                            if i % 2 == 0:
                                shading = OxmlElement('w:shd')
                                shading.set(qn('w:fill'), 'FAFAFA')
                                cell._tc.get_or_add_tcPr().append(shading)

                        # 셀 여백 설정
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcMar = OxmlElement('w:tcMar')
                        for margin_name, margin_val in [('top', '80'), ('left', '120'), ('bottom', '80'), ('right', '120')]:
                            margin = OxmlElement(f'w:{margin_name}')
                            margin.set(qn('w:w'), margin_val)
                            margin.set(qn('w:type'), 'dxa')
                            tcMar.append(margin)
                        tcPr.append(tcMar)

                        # 셀 수직 정렬 (가운데)
                        vAlign = OxmlElement('w:vAlign')
                        vAlign.set(qn('w:val'), 'center')
                        tcPr.append(vAlign)

            # 테이블 테두리 스타일 (깔끔한 라인)
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblBorders = OxmlElement('w:tblBorders')

            # 외곽선만 진하게, 내부선은 연하게
            for border_name in ['top', 'bottom']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '8')  # 진한 선
                border.set(qn('w:color'), '4A4A4A')
                tblBorders.append(border)

            for border_name in ['left', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')  # 좌우 테두리 없음
                tblBorders.append(border)

            for border_name in ['insideH']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), 'E0E0E0')
                tblBorders.append(border)

            for border_name in ['insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), 'E0E0E0')
                tblBorders.append(border)

            tblPr.append(tblBorders)

            # 테이블 후 여백
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(15)

            return table

        def process_content_with_tables(doc, text):
            """본문 텍스트에서 표를 감지하고 처리"""
            blocks = []
            current_block = []
            table_lines = []

            lines = text.split('\n')
            i = 0

            def is_table_start(line, next_line=None):
                """표 시작 라인인지 확인"""
                stripped = line.strip()
                # 마크다운 테이블 (| cell | cell |)
                if stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 2:
                    return True
                # 파이프로 구분된 내용 (cell | cell)
                if '|' in stripped and len(stripped.split('|')) >= 2:
                    parts = stripped.split('|')
                    if any(p.strip() and not re.match(r'^[\s\-:]+$', p) for p in parts):
                        return True
                # 콜론 기반 테이블 (키: 값) - 연속된 경우
                if next_line and ':' in stripped and ':' in next_line:
                    if len(stripped.split(':')[0].strip()) < 30 and len(next_line.split(':')[0].strip()) < 30:
                        return True
                return False

            def is_table_continue(line):
                """표 계속 라인인지 확인"""
                stripped = line.strip()
                # 빈 줄은 표 종료
                if not stripped:
                    return False
                # 마크다운 테이블
                if stripped.startswith('|') and stripped.endswith('|'):
                    return True
                # 마크다운 구분선
                if re.match(r'^\|[\s\-:]+\|$', stripped):
                    return True
                # 파이프로 구분된 내용
                if '|' in stripped:
                    return True
                # 콜론 기반 (키: 값)
                if ':' in stripped and len(stripped.split(':')[0].strip()) < 30:
                    return True
                return False

            while i < len(lines):
                line = lines[i]
                next_line = lines[i + 1] if i + 1 < len(lines) else None

                if is_table_start(line, next_line):
                    # 이전 일반 텍스트 저장
                    if current_block:
                        blocks.append(('text', '\n'.join(current_block)))
                        current_block = []

                    # 표 라인 수집
                    table_lines = [line]
                    i += 1
                    while i < len(lines) and is_table_continue(lines[i]):
                        table_lines.append(lines[i])
                        i += 1

                    if len(table_lines) >= 2:
                        blocks.append(('table', '\n'.join(table_lines)))
                    else:
                        current_block.extend(table_lines)
                    table_lines = []
                else:
                    current_block.append(line)
                    i += 1

            # 마지막 블록 저장
            if current_block:
                blocks.append(('text', '\n'.join(current_block)))

            return blocks

        for idx, chapter in enumerate(outline):
            if chapter in chapters_data:
                # ─────────────────────────────────────
                # 챕터 시작 페이지 (프리미엄 오프너)
                # ─────────────────────────────────────
                clean_chapter = chapter
                for prefix in [f"PART {idx + 1}.", f"PART{idx + 1}.", f"PART {idx + 1} ", f"PART{idx + 1} ", f"{idx + 1}.", f"{idx + 1})"]:
                    clean_chapter = clean_chapter.replace(prefix, "").strip()

                ch_name = add_chapter_opener(doc, idx, clean_chapter)
                add_bookmark(ch_name, f"chapter_{idx + 1}")

                doc.add_page_break()

                # ─────────────────────────────────────
                # 본문 시작
                # ─────────────────────────────────────
                ch_data = chapters_data[chapter]
                subtopics = ch_data.get('subtopics', [])

                for sub_idx, sub in enumerate(subtopics):
                    content = ch_data.get('subtopic_data', {}).get(sub, {}).get('content', '')
                    if content:
                        # 소제목마다 새 페이지에서 시작 (첫 번째 제외)
                        if sub_idx > 0:
                            doc.add_page_break()

                        # 소제목 (프리미엄 스타일)
                        sub_title = add_subtopic_header(doc, sub, sub_idx)
                        add_bookmark(sub_title, f"subtopic_{idx + 1}_{sub_idx + 1}")

                        # 본문 내용 (표 감지 및 처리 포함)
                        cleaned = clean_content(content)

                        # 표가 포함된 콘텐츠 처리
                        content_blocks = process_content_with_tables(doc, cleaned)

                        is_first_para = True
                        for block_type, block_content in content_blocks:
                            if block_type == 'table':
                                # 표 데이터 파싱 및 프리미엄 테이블 생성
                                table_data = parse_table_data(block_content)
                                if table_data and len(table_data) >= 2:
                                    # 표 전 여백
                                    spacer = doc.add_paragraph()
                                    spacer.paragraph_format.space_after = Pt(10)
                                    add_premium_table(doc, table_data)
                                    is_first_para = False
                            else:
                                # 일반 텍스트 처리
                                paragraphs = block_content.split('\n\n')
                                if not paragraphs or not any(p.strip() for p in paragraphs):
                                    paragraphs = block_content.split('\n')

                                for para_text in paragraphs:
                                    if para_text.strip():
                                        format_body_paragraph(doc, para_text.strip(), is_first=is_first_para)
                                        is_first_para = False

                        # 소제목 사이 구분 (마지막 소제목 제외)
                        if sub_idx < len(subtopics) - 1:
                            separator = doc.add_paragraph()
                            separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            sep_run = separator.add_run("· · ·")
                            set_font(sep_run, 10, color=(200, 200, 200))
                            separator.paragraph_format.space_before = Pt(30)
                            separator.paragraph_format.space_after = Pt(30)

                doc.add_page_break()

        # ══════════════════════════════════════════════════════════════
        # 에필로그 (프리미엄 에디토리얼 스타일)
        # ══════════════════════════════════════════════════════════════

        # 상단 넓은 여백
        for _ in range(6):
            doc.add_paragraph()

        # 에필로그 라벨
        ep_label = doc.add_paragraph()
        ep_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep_label_run = ep_label.add_run("E P I L O G U E")
        set_font(ep_label_run, 9, color=(160, 160, 160))
        ep_label.paragraph_format.space_after = Pt(16)

        # 구분선
        ep_line = doc.add_paragraph()
        ep_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep_line_run = ep_line.add_run("───────────")
        set_font(ep_line_run, 10, color=(200, 200, 200))
        ep_line.paragraph_format.space_after = Pt(20)

        # 에필로그 제목
        epilogue_title = doc.add_paragraph()
        epilogue_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep_run = epilogue_title.add_run("마치며")
        set_font(ep_run, 14, bold=True, color=(40, 40, 40))
        epilogue_title.paragraph_format.space_after = Pt(40)

        # 에필로그 내용 - AI가 인터뷰 내용을 참고해서 자연스럽게 작성
        epilogue_text = None
        if interview_data:
            epilogue_prompt = f"""당신은 자청 스타일로 글을 쓰는 베스트셀러 작가입니다. 에필로그를 작성하세요.

[저자 정보 - 참고용, 그대로 복사하지 말 것]
- 분야: {interview_data.get('field', '')}
- 경력 기간: {interview_data.get('experience_years', '')}
- 책 주제: {interview_data.get('topic', '')}
- 타겟 독자: {interview_data.get('target_reader', '')}
- 독자에게 전하고 싶은 말: {interview_data.get('final_message', '')}
- 작가 경력/경험: {interview_data.get('author_career', '')}
- 어려움/실패 경험: {interview_data.get('struggle_story', '')}
- 극복 스토리: {interview_data.get('breakthrough', '')}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

## 에필로그 작성 원칙 (스토리텔링)

### 1. 나의 스토리로 시작 (3-4문장)
- 작가 경력/경험을 자연스럽게 녹여서
- "저는 ~했습니다" 형식으로 간결하게
- 구체적 숫자나 사실 포함

### 2. 왜 이 책을 썼는지 (2-3문장)
- 내가 겪은 어려움 + 극복 과정 힌트
- 독자를 위해 책을 쓴 진심

### 3. 독자에게 한마디 (2-3문장)
- 지금 당장 할 수 있는 구체적 행동 하나
- 진심 어린 마무리 (근데 뻔하지 않게)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[문체]
- 짧은 문장, 짧은 문단 (2-4문장)
- 구어체 + 합쇼체 ("~거든요", "~잖아요" OK)

[분량] 400-500자

[금지 - 절대 쓰지 말 것]
- 저자 정보를 그대로 복사 붙여넣기
- 교과서 표현: "여정", "발걸음", "함께 하시길 바랍니다", "진정한"
- AI 표현: "~의 중요성", "다양한", "효과적인", "~를 통해"
- 과장: "놀라운", "혁신적인", "충격적인"
- 뻔한 말: "포기하지 마세요", "꾸준히 하세요", "화이팅"
- 마크다운 문법

에필로그만 출력하세요."""

            generated_epilogue = ask_ai(epilogue_prompt, 0.7)
            if generated_epilogue:
                epilogue_text = generated_epilogue

        if not epilogue_text:
            epilogue_text = """여기까지 읽어주셔서 감사합니다.

이 책에 담긴 내용이 당신의 삶에 작은 변화라도 만들어낸다면 그것으로 충분합니다.

완벽할 필요 없습니다. 지금 당장 할 수 있는 것 하나만 시작해보세요.

작은 시작이 큰 결과를 만듭니다.

항상 응원합니다."""

        for para_text in epilogue_text.split('\n\n'):
            if para_text.strip():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para_run = para.add_run(para_text.strip())
                set_font(para_run, 10, color=(50, 50, 50))
                para_format = para.paragraph_format
                para_format.line_spacing = 1.7
                para_format.space_after = Pt(14)
                para_format.first_line_indent = Cm(0.5)

        # 저자 서명 (프리미엄 스타일)
        for _ in range(3):
            doc.add_paragraph()

        # 서명 라인
        sign_line = doc.add_paragraph()
        sign_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_line_run = sign_line.add_run("─────")
        set_font(sign_line_run, 10, color=(200, 200, 200))
        sign_line.paragraph_format.space_after = Pt(10)

        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_run = sign_para.add_run(f"{author if author else '저자'}")
        set_font(sign_run, 11, italic=True, color=(80, 80, 80))

        doc.add_page_break()

        # ══════════════════════════════════════════════════════════════
        # 저자 소개 페이지 (프리미엄 에디토리얼 스타일)
        # ══════════════════════════════════════════════════════════════

        # 상단 넓은 여백
        for _ in range(6):
            doc.add_paragraph()

        # 저자 소개 라벨
        about_label = doc.add_paragraph()
        about_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        about_label_run = about_label.add_run("A B O U T")
        set_font(about_label_run, 9, color=(160, 160, 160))
        about_label.paragraph_format.space_after = Pt(16)

        # 구분선
        about_line = doc.add_paragraph()
        about_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        about_line_run = about_line.add_run("───────────")
        set_font(about_line_run, 10, color=(200, 200, 200))
        about_line.paragraph_format.space_after = Pt(20)

        # 저자명 (크게)
        author_name_para = doc.add_paragraph()
        author_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_name_run = author_name_para.add_run(author if author else "저자")
        set_font(author_name_run, 16, bold=True, color=(40, 40, 40))
        author_name_para.paragraph_format.space_after = Pt(30)

        # 저자 소개 내용
        if interview_data:
            career_text = interview_data.get('author_career', '')
            field = interview_data.get('field', '')
            exp = interview_data.get('experience_years', '')
            method = interview_data.get('core_method', '')

            if career_text:
                author_bio = f"""{field} 분야에서 {exp}간 활동해온 실전가.

{career_text}

{method[:100] if method else ''}"""
            else:
                author_bio = f"""{field} 분야에서 {exp}간 활동해온 실전가.

{method}"""
        else:
            author_bio = """실전에서 직접 부딪히며 쌓은 노하우를 독자들과 나누고자 이 책을 썼다."""

        for para_text in author_bio.split('\n\n'):
            if para_text.strip():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para_run = para.add_run(para_text.strip())
                set_font(para_run, 10, color=(80, 80, 80))
                para_format = para.paragraph_format
                para_format.line_spacing = 1.6
                para_format.space_after = Pt(14)

        # 하단 장식
        for _ in range(4):
            doc.add_paragraph()

        end_mark = doc.add_paragraph()
        end_mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end_run = end_mark.add_run("◆")
        set_font(end_run, 12, color=(200, 200, 200))

        # 메모리에 저장
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue(), None

    except Exception as e:
        return None, f"문서 생성 오류: {str(e)}"

def generate_outline_only(interview_data, progress_placeholder):
    """인터뷰 데이터를 기반으로 목차까지만 생성 (본문 제외)"""
    try:
        topic = interview_data.get('topic', '')
        if not topic:
            return False

        # 1. 타겟 자동 설정
        progress_placeholder.info("🎯 1/4 타겟 독자 분석 중...")
        target = f"{interview_data.get('target_reader', '')} - {interview_data.get('target_problem', '')}"
        st.session_state['target_persona'] = target

        # 2. 책 고유 컨셉 생성 (가장 중요!)
        progress_placeholder.info("💡 2/4 책 고유 컨셉 설계 중...")
        concept_prompt = f"""당신은 크몽/클래스101 베스트셀러 전자책 기획자입니다.
이 책만의 '고유한 시스템/공식'을 만들어야 합니다.

[저자 정보]
주제: {topic}
핵심 방법: {interview_data.get('core_method', '')}
저자만의 차별점: {interview_data.get('unique_point', '')}
타겟의 고민: {interview_data.get('target_problem', '')}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 실제 잘 팔리는 전자책의 고유 시스템/공식 예시
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[영어+알파벳 조합]
• SLP 공식 (Search-Learn-Produce)
• CPM 전략 (Contents-Profit-Multiply)
• 3R 시스템 (Research-Run-Repeat)

[한글 조어]
• 역행 루프
• 수익 사이클
• 복리 엔진

[분야별 비유 - 다양하게!]
• 주식: 스노우볼, 배당 파이프라인, 현금흐름 엔진
• 블로그: 검색 알고리즘, 트래픽 자석, 상위노출 공식
• 마케팅: 전환 퍼널, 구매 트리거, 설득 코드
• 습관: 루틴 시스템, 자동화 루프, 습관 스택
• 투자: 리스크 헤지, 분산 매트릭스, 안전마진

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 만들어야 할 것
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. 고유 시스템/공식 이름 (가장 중요!)
   - 영어 약자 + 한글 설명 (예: "CPM 전략")
   - 또는 직관적인 한글 조어 (예: "역행 루프")
   - 목차 전체에서 이 용어가 반복되어야 함

2. 핵심 관점
   - 이 주제를 어떤 새로운 시각으로 보는가?
   - 남들과 다른 접근법

3. 핵심 메시지
   - "[시스템명]만 알면 ~할 수 있다" 형식

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 부자연스러운 과장:
- "제국을 건설", "왕좌에 오르다", "언더그라운드"
- "전설의", "역사를 바꾼", "세계 최초"

❌ 유치한 단어:
- 황금, 보물, 비밀, 마법, 연금술

❌ 모든 분야에 같은 비유 사용:
- 부동산 비유만 반복하지 말 것
- 주제에 맞는 다양한 비유 사용

❌ 이미 유명한 이름:
- 역행자, 추월차선, 아토믹 해빗 등 그대로 사용

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 형식
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[고유 시스템/공식 이름]
(영어 약자+한글 또는 참신한 한글 조어)

[핵심 관점]
(2~3문장, 자연스럽게)

[핵심 메시지]
(한 문장, "[시스템명]으로 ~하는 방법")

[목차에서 반복할 키워드]
(시스템 이름 또는 핵심 단어 1~2개)"""

        book_concept = ask_ai(concept_prompt, 0.8)
        st.session_state['book_concept'] = book_concept

        # 3. 제목 생성
        progress_placeholder.info("📝 3/4 제목 생성 중...")
        title_prompt = f"""당신은 크몽/클래스101 베스트셀러 전자책 제목을 만드는 전문가입니다.
결제 버튼을 누르게 만드는 제목을 써주세요.

[이 책의 컨셉]
{book_concept}

[주제]
{topic}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 실제 잘 팔리는 전자책 제목 분석
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[자청/프드프 스타일 - 컨셉 중심]
• 역행자 - 부자의 시간을 사는 법
• 돈의 속성 - 최소한 이것만은 알아야 할
• 부의 추월차선 - 부자들이 말해주지 않는 것

[신사임당/클래스101 스타일 - 결과 중심]
• 퇴사 후 월 1000만원 버는 글쓰기
• 블로그로 월 300 만드는 현실적인 방법
• 투잡러의 시간관리 비법

[크몽 베스트셀러 - 구체적 약속]
• 30일 만에 첫 수익 내는 스마트스토어
• 3개월 안에 월 100 만드는 전자책 공식
• 회사 다니면서 월 200 추가 수입 만들기

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 좋은 제목의 공식
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[공식 1] 컨셉형 (2~4단어)
저자만의 프레임워크/용어가 들어간 제목
예: "역행자", "추월차선", "언카피어블"

[공식 2] 결과형 (구체적 숫자 포함)
기간 + 결과가 명확한 제목
예: "3개월 만에 월 300", "100일 글쓰기"

[공식 3] 타겟형 (누구를 위한)
특정 대상의 고민을 건드리는 제목
예: "퇴사 준비생의 월급 독립기", "직장인의 두 번째 월급"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 유치한 비유:
나침반, 지도, 열쇠, 보물, 황금, 마법, 연금술

❌ AI스러운 제목:
"~의 이해", "~가이드", "~완벽 정복"
"효과적인 ~", "성공적인 ~"

❌ 너무 추상적:
의미를 알 수 없는 신조어
무슨 내용인지 전혀 감이 안 오는 제목

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

제목: 컨셉이 드러나면서도 무슨 책인지 알 수 있게
부제: 구체적인 결과/약속을 담아 15~25자

JSON만 출력:
{{
    "title": "제목 (컨셉+내용이 드러나게, 3~8단어)",
    "subtitle": "부제 (구체적 결과/약속, 15~25자)"
}}"""

        title_result = ask_ai(title_prompt, 0.4)
        title_data = parse_json(title_result)
        if title_data:
            st.session_state['book_title'] = title_data.get('title', topic)
            st.session_state['subtitle'] = title_data.get('subtitle', '')

        # 4. 목차 생성 (책 컨셉 기반)
        progress_placeholder.info("📋 4/4 목차 설계 중...")
        outline_prompt = f"""당신은 크몽/클래스101 베스트셀러 전자책 기획자입니다.
목차만 보고 결제 버튼을 누르게 만드세요.

[이 책의 고유 시스템/공식]
{book_concept}

[주제]: {topic}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 목차 구성 원칙
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. 위 컨셉의 "시스템/공식 이름"을 목차 전체에서 활용
2. 챕터마다 그 시스템의 다른 측면을 다룸
3. 독자가 "이 시스템을 배우면 나도 할 수 있겠다" 느끼게

예시) "CPM 전략" 컨셉이라면:
- PART 1. 콘텐츠 없이 시작하면 망하는 이유
- PART 2. 수익이 자동으로 굴러가는 구조
- PART 3. 하나로 열 개를 만드는 복제 기술

예시) "상위노출 알고리즘" 컨셉이라면:
- PART 1. 네이버가 순위를 정하는 진짜 기준
- PART 2. 경쟁 없이 1등 하는 틈새 공략법
- PART 3. 한 번 올리면 계속 오르는 구조

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📌 챕터 제목 공식
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

"PART X. [자극적인 제목]" (콜론 사용 금지!)

✅ 좋은 예:
- "PART 1. 왜 열심히 해도 결과가 안 나오는가"
- "PART 2. 상위 1%만 아는 검색 원리"
- "PART 3. 한 번 세팅하면 알아서 굴러가는 구조"
- "PART 4. 30일 안에 첫 결과를 만드는 순서"
- "PART 5. 월 100에서 월 1000으로 가는 로드맵"

❌ 나쁜 예:
- "PART 1. 시작" (추상적)
- "PART 2. 제국을 건설하다" (부자연스러운 과장)
- "PART 3. 기초 이해" (설명서 같음)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 소제목 공식
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

각 소제목은 다른 형식으로 (같은 패턴 반복 금지!)

[반전형] "~는 거짓말이다", "~하면 오히려 망한다"
[시스템형] "[공식명]의 첫 번째 원칙", "[키워드]가 작동하는 방식"
[구체적 방법] "3단계로 ~하는 순서", "10분 만에 ~하는 방법"
[질문형] "왜 ~은 실패하는가", "어떻게 ~할 수 있는가"
[결과형] "이것만 바꿔도 ~가 달라진다"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 부자연스러운 과장:
- "제국을 건설", "왕좌에 오르다", "언더그라운드에서"
- "전설의", "역사를 바꾼", "나의 고백"

❌ AI스러운 표현:
- "~의 이해", "~의 기초", "효과적인", "성공적인"

❌ 같은 비유 반복:
- 모든 챕터에 부동산 비유 쓰지 말 것
- 다양한 비유 사용 (기계, 알고리즘, 시스템, 공식 등)

❌ 숫자 패턴 반복:
- "99%", "1%" 같은 표현은 전체에서 1번만

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 콜론(:) 사용 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ "PART 1. 잘못된 공식: 왜..." (콜론 있음)
✅ "PART 1. 왜 열심히 해도 안 되는가" (콜론 없음)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 형식
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PART 1. [자극적인 제목]
- [소제목 15~30자]
- [소제목 15~30자]
- [소제목 15~30자]

PART 2. [자극적인 제목]
- [소제목]
- [소제목]
- [소제목]

PART 3. [자극적인 제목]
- [소제목]
- [소제목]
- [소제목]

PART 4. [자극적인 제목]
- [소제목]
- [소제목]
- [소제목]

PART 5. [자극적인 제목]
- [소제목]
- [소제목]
- [소제목]

목차만 출력. 콜론(:) 절대 사용 금지."""

        outline_result = ask_ai(outline_prompt, 0.4)

        if outline_result:
            chapters = []
            subtopics = {}
            current_ch = None

            lines = outline_result.split('\n')
            for i, orig_line in enumerate(lines):
                line = orig_line.strip()
                if not line:
                    continue

                # 챕터 감지 (PART, 파트, Chapter, 1., 2. 등 다양한 형식)
                is_chapter = False
                ch_name = None

                # PART 1. 제목 형식
                if re.match(r'^(PART|파트|Part)\s*\d+[\.\s]', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = line
                # Chapter 1. 제목 형식
                elif re.match(r'^(Chapter|챕터)\s*\d+[\.\s]', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = line
                # 마크다운 헤더 형식
                elif re.match(r'^#+\s*(PART|파트|Chapter|챕터|\d+)', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = re.sub(r'^#+\s*', '', line)
                # 1. 제목 형식 (숫자로 시작, 들여쓰기 없음)
                elif re.match(r'^\d+[\.\)]\s', line) and not orig_line.startswith(' '):
                    is_chapter = True
                    ch_name = line
                # 【1부】 형식
                elif re.match(r'^[【\[]?\s*\d+\s*(부|장|편)[】\]]?', line):
                    is_chapter = True
                    ch_name = line

                if is_chapter and ch_name:
                    ch_name = re.sub(r'^[#\*\-\s]+', '', ch_name)
                    ch_name = ch_name.replace('**', '').replace('*', '').replace('#', '').strip()
                    if ch_name and len(ch_name) > 3:
                        current_ch = ch_name
                        if current_ch not in chapters:
                            chapters.append(current_ch)
                            subtopics[current_ch] = []
                elif current_ch:
                    # 소제목 감지
                    is_subtopic = False
                    st_name = None

                    # - 소제목 형식
                    if re.match(r'^[\-\•\·\*\→\▶]\s*', line):
                        is_subtopic = True
                        st_name = re.sub(r'^[\-\•\·\*\→\▶]\s*', '', line)
                    # 1) 소제목, a) 소제목 형식
                    elif re.match(r'^[a-z\d][\)\.\:]\s', line, re.IGNORECASE):
                        is_subtopic = True
                        st_name = re.sub(r'^[a-z\d][\)\.\:]\s*', '', line, flags=re.IGNORECASE)
                    # 들여쓰기된 라인
                    elif orig_line.startswith('  ') or orig_line.startswith('\t'):
                        is_subtopic = True
                        st_name = line.lstrip('- •·*→▶0123456789.):\t ')
                    # 챕터가 아닌 일반 텍스트 (이전이 챕터였고, 현재가 짧은 문장이면 소제목으로 간주)
                    elif len(chapters) > 0 and not re.match(r'^(PART|파트|Part|Chapter|챕터|\d+[\.\)])', line, re.IGNORECASE):
                        if len(line) > 5 and len(line) < 100:
                            is_subtopic = True
                            st_name = line.lstrip('- •·*→▶0123456789.):\t ')

                    if is_subtopic and st_name:
                        st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                        st_name = re.sub(r'^\d+[\.\)\:]\s*', '', st_name)  # 앞 숫자 제거
                        if st_name and len(st_name) > 3 and len(subtopics[current_ch]) < 5:
                            # 챕터 이름과 동일하면 스킵
                            if st_name.lower() != current_ch.lower() and st_name not in subtopics[current_ch]:
                                subtopics[current_ch].append(st_name)

            if chapters:
                st.session_state['outline'] = chapters
                st.session_state['chapters'] = {}
                for ch in chapters:
                    st.session_state['chapters'][ch] = {
                        'subtopics': subtopics.get(ch, []),
                        'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                    }

        # 목차가 생성되지 않았으면 기본 목차 생성
        if not st.session_state.get('outline'):
            progress_placeholder.warning("목차 파싱 실패, 기본 목차 생성 중...")
            default_chapters = [
                "PART 1. 왜 지금인가",
                "PART 2. 진짜 비밀",
                "PART 3. 실전 공식",
                "PART 4. 수익화",
                "PART 5. 다음 단계"
            ]
            default_subtopics = {
                default_chapters[0]: [f"90%가 {topic}에 실패하는 이유", "아무도 말해주지 않는 진실", "지금 시작해야 하는 3가지 이유"],
                default_chapters[1]: ["전문가들이 숨기는 핵심 원칙", f"{topic}의 본질을 꿰뚫는 법", "이것만 알면 절반은 성공"],
                default_chapters[2]: ["바로 써먹는 5단계 공식", "실패 없이 시작하는 체크리스트", "첫 달에 결과 내는 비법"],
                default_chapters[3]: ["월 100만원 만드는 구조", "자동화로 시간 벌기", "확장 전략 A to Z"],
                default_chapters[4]: ["1년 후 당신의 모습", "다음 레벨로 가는 로드맵", "지금 바로 해야 할 첫 번째 행동"]
            }
            st.session_state['outline'] = default_chapters
            st.session_state['chapters'] = {}
            for ch in default_chapters:
                st.session_state['chapters'][ch] = {
                    'subtopics': default_subtopics.get(ch, []),
                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in default_subtopics.get(ch, [])}
                }

        # 저자명 및 인터뷰 데이터 저장
        st.session_state['author_name'] = interview_data.get('author_name', '')
        st.session_state['interview_data'] = interview_data
        st.session_state['topic'] = topic

        progress_placeholder.success("✅ 목차 생성 완료! 목차를 확인하고 수정할 수 있습니다.")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False

def regenerate_single_subtopic(chapter_name, subtopic_index, existing_subtopics):
    """개별 소제목 AI 재생성 - 자청/프드프 스타일"""
    topic = st.session_state.get('topic', '')
    book_concept = st.session_state.get('book_concept', '')

    # 기존 소제목들 (중복 방지용)
    other_subtopics = [s for i, s in enumerate(existing_subtopics) if i != subtopic_index]

    prompt = f"""당신은 자청, 프드프의 편집자입니다.
목차만 보고 결제하게 만드는 소제목 하나를 써주세요.

[책 컨셉]
{book_concept}

[챕터]: {chapter_name}
[주제]: {topic}

[기존 소제목들 - 이것들과 완전히 다르게]
{chr(10).join(f'- {s}' for s in other_subtopics)}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 실제 베스트셀러 소제목 예시
- "유전자를 역행하는 인간"
- "당신의 뇌는 원시인이다"
- "열심히 일하면 부자가 된다는 환상"
- "노력이 배신하는 진짜 이유"
- "돈을 끌어당기는 정체성"

✅ 형식 (하나 선택):
- 충격형: "~는 거짓말이다", "~하면 망한다"
- 고백형: "나는 왜 ~했을까", "~를 깨닫기까지"
- 도발형: "~은 필요 없다", "~만 있으면 된다"
- 질문형: "왜 ~은 실패하는가"

❌ 금지:
- 설명서 표현: "~의 이해", "~하는 방법"
- 유치한 비유: 나침반, 열쇠, 마법
- 기존 소제목과 비슷한 패턴

소제목 하나만 (15~30자, 기호 없이):"""

    result = ask_ai(prompt, 0.9)
    if result:
        return result.strip().strip('"').strip("'").strip('-').strip()
    return None

def regenerate_chapter_subtopics(chapter_name, chapter_index):
    """챕터의 모든 소제목 AI 재생성 - 자청/프드프 스타일"""
    topic = st.session_state.get('topic', '')
    book_concept = st.session_state.get('book_concept', '')
    outline = st.session_state.get('outline', [])

    # 다른 챕터들의 소제목들 (중복 방지)
    other_chapter_subtopics = []
    for ch in outline:
        if ch != chapter_name:
            ch_data = st.session_state['chapters'].get(ch, {})
            other_chapter_subtopics.extend(ch_data.get('subtopics', []))

    # 챕터별 역할 정의
    chapter_roles = {
        0: "착각/각성 - 독자가 몰랐던 불편한 진실을 폭로",
        1: "해체 - 기존 상식과 믿음을 완전히 무너뜨림",
        2: "구조/재구축 - 저자만의 새로운 방법론 제시",
        3: "실전 - 구체적이고 따라할 수 있는 방법",
        4: "도약 - 변화된 미래와 행동 촉구"
    }
    current_role = chapter_roles.get(chapter_index, "핵심 내용 전달")

    prompt = f"""당신은 자청, 프드프의 편집자입니다.
이 챕터의 소제목 3개를 결제하고 싶게 써주세요.

[책 컨셉]
{book_concept}

[주제]: {topic}
[챕터]: {chapter_name}
[이 챕터의 역할]: {current_role}

[다른 챕터 소제목들 - 완전히 다르게 써야 함]
{chr(10).join(f'- {s}' for s in other_chapter_subtopics[:8])}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 실제 베스트셀러 소제목 예시
- "유전자를 역행하는 인간"
- "당신의 뇌는 원시인이다"
- "열심히 일하면 부자가 된다는 환상"
- "저축의 함정, 투자의 진실"
- "멘탈이 무너지면 돈도 무너진다"
- "운을 지배하는 자의 비밀"

✅ 각 소제목은 다른 형식으로:
1번: 충격/반전형 (예: "~는 거짓말이다")
2번: 고백/스토리형 (예: "나는 왜 ~했을까")
3번: 선언/도발형 (예: "~만 있으면 된다")

❌ 금지:
- 설명서 표현: "~의 이해", "~하는 방법", "효과적인"
- 유치한 비유: 나침반, 열쇠, 마법
- 같은 패턴 반복

소제목 3개만 출력 (줄바꿈으로 구분, 기호/번호 없이):"""

    result = ask_ai(prompt, 0.8)
    if result:
        lines = [line.strip().strip('"').strip("'").strip('-').strip() for line in result.strip().split('\n') if line.strip() and len(line.strip()) > 5]
        return lines[:3] if lines else None
    return None

def generate_body_from_outline(interview_data, progress_placeholder):
    """생성된 목차를 기반으로 본문만 생성"""
    try:
        topic = interview_data.get('topic', '')
        book_concept = st.session_state.get('book_concept', '')

        if not st.session_state.get('outline') or not st.session_state.get('chapters'):
            progress_placeholder.error("먼저 목차를 생성해주세요.")
            return False

        # 본문 생성
        total_subtopics = sum(len(st.session_state['chapters'][ch]['subtopics']) for ch in st.session_state['outline'])
        done = 0

        for ch in st.session_state['outline']:
            ch_data = st.session_state['chapters'][ch]
            for sub in ch_data['subtopics']:
                done += 1
                progress_placeholder.info(f"✍️ 본문 작성 중... ({done}/{total_subtopics}) - {sub[:20]}...")

                # 이전 소제목들의 내용 요약 (중복 방지용)
                prev_contents = []
                for prev_ch in st.session_state['outline']:
                    if prev_ch == ch:
                        break
                    prev_ch_data = st.session_state['chapters'].get(prev_ch, {})
                    for prev_sub in prev_ch_data.get('subtopics', []):
                        prev_content = prev_ch_data.get('subtopic_data', {}).get(prev_sub, {}).get('content', '')
                        if prev_content:
                            prev_contents.append(f"- {prev_sub}: {prev_content[:100]}...")

                # 현재 챕터의 이전 소제목들
                current_ch_prev = []
                for prev_sub in ch_data['subtopics']:
                    if prev_sub == sub:
                        break
                    prev_content = ch_data.get('subtopic_data', {}).get(prev_sub, {}).get('content', '')
                    if prev_content:
                        current_ch_prev.append(f"- {prev_sub}: {prev_content[:100]}...")

                prev_summary = "\n".join(prev_contents[-5:] + current_ch_prev) if (prev_contents or current_ch_prev) else "없음"

                # 소제목 인덱스에 따라 다른 시작 스타일 선택
                hook_styles = [
                    "질문으로 시작 (예: '왜 우리는 항상 실패할까요?')",
                    "고백으로 시작 (예: '솔직히 말하면, 저도 처음엔 몰랐습니다.')",
                    "반전 사실로 시작 (예: '대부분의 사람들이 믿는 것과 정반대였습니다.')",
                    "통계/숫자로 시작 (예: '92%가 이 실수를 반복합니다.')",
                    "에피소드로 시작 (예: '어느 날 친구에게서 연락이 왔습니다.')",
                    "선언으로 시작 (예: '결론부터 말하겠습니다. 방법은 하나입니다.')",
                ]
                current_hook_style = hook_styles[done % len(hook_styles)]

                content_prompt = f"""당신은 전세계 베스트셀러 작가들의 기법을 마스터한 작가입니다.

🚨🚨🚨 최우선 규칙 (반드시 지켜라) 🚨🚨🚨
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. 독자/타겟을 직접 부르지 마라. 절대 금지.
   ❌ "김대리님", "30대 직장인 여러분", "주부님들"
   ❌ "초보자 분들", "직장인이라면", "당신은"
   ✅ "저는", "우리는" 정도만 사용

2. 이전 글과 다른 시작으로 시작해라!
   이번 글은 반드시: {current_hook_style}
   ❌ 매번 같은 패턴 금지 (날짜+상황, 고백, 질문 등)
   ❌ "20XX년 X월" 형식의 날짜로 시작 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[집필 정보]
주제: {topic}
챕터: {ch}
현재 작성할 소제목: {sub}
핵심 방법론: {interview_data.get('core_method', '')}

[이 책의 고유 컨셉]
{book_concept}

[이미 작성된 내용 - 중복 금지]
{prev_summary}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 글쓰기 스타일
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[문체] 자청 스타일 합쇼체만 사용:
- "~입니다", "~습니다", "~거든요", "~더라고요", "~잖아요"
- "저는 ~했습니다", "그게 ~였죠", "~한 거죠"

[구성]
- 본문은 하나의 이야기처럼 자연스럽게 흘러가야 함
- 문단과 문단을 자연스럽게 연결
- 반전이나 깨달음 포인트 하나 포함

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 구체적인 실행 방법 (가장 중요!)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

✅ 초보자도 바로 따라할 수 있게:
- 방법을 알려줄 때는 "무엇을 → 어디서 → 어떻게" 순서로
- 예: "네이버에 접속해서 → 검색창에 '주제+정보' 입력 → 연관검색어 5개 메모"
- 예: "엑셀을 열고 → A열에 키워드 입력 → B열에 검색량 입력 → 내림차순 정렬"
- 추상적 설명 금지 ("분석하세요", "파악하세요", "연구하세요")
- 구체적 도구/사이트명 언급 (네이버, 구글, 블랙키위, 키워드마스터 등)

✅ 예시 인물 사용 시:
- "김씨", "이씨", "박씨" 금지
- 자연스러운 가명 사용: "민준이라는 친구", "직장인 수현", "블로거 지우씨"
- 또는 "제 수강생 중 한 분", "저와 함께 시작한 동료"

✅ 고급스러운 표현:
- "생명 주기" → "성장 곡선", "발전 단계"
- "돈 버는" → "수익을 창출하는", "경제적 자유를 얻는"
- "쉽게" → "효율적으로", "체계적으로"
- "열심히" → "꾸준히", "전략적으로"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지 표현
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 유치하거나 어색한 표현:
- "후다닥", "슝슝", "쭉쭉", "팡팡", "뿅뿅"
- "ㅇㅇ조", "ㅋㅋ", "ㅎㅎ", 인터넷 신조어
- "짜잔", "두둥", "쨘", 효과음
- "레알", "진짜루", "완전", "대박" 등 구어체

❌ 유치한 비유:
- "마법", "연금술", "황금", "열쇠", "보물", "나침반", "지도"
- "저수지", "날개", "로켓", "엔진", "머신", "파이프라인"
- "~이 보이는", "~을 여는", "~의 비밀"

❌ AI스러운 표현:
- "중요합니다", "따라서", "결론적으로"
- "~하는 것이 좋습니다", "~해야 합니다"

❌ 반말/높임 금지:
- "~다", "~해", "~해라", "~하자" (X)
- "~하셔야 합니다", "~하시면" (X)

❌ 형식 금지:
- 소제목/번호 나열 ("1.", "2.", "첫째", "둘째")
- 마크다운/이모지
- 글머리 기호 (-, •, *)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📏 분량: 1500~1800자
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

'{sub}' 본문을 작성하세요.
- 이번 글 시작: {current_hook_style}
- 자연스러운 흐름으로 1500~1800자
- 반전/깨달음 하나 필수
- 이전 글과 완전히 다른 톤으로 시작"""

                content = ask_ai(content_prompt, 0.7)
                if content:
                    content = clean_content(content)  # 이모티콘/마크다운 제거
                    ch_data['subtopic_data'][sub]['content'] = content

        # 완료 처리
        st.session_state['interview_completed'] = True
        progress_placeholder.success("✅ 본문 생성 완료!")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False

def go_next():
    if st.session_state['current_page'] < 6:
        st.session_state['current_page'] += 1

def go_prev():
    if st.session_state['current_page'] > 0:
        st.session_state['current_page'] -= 1

def auto_generate_all(topic, progress_placeholder):
    """주제만 입력하면 목차+본문까지 자동 생성"""
    try:
        # 1. 타겟 자동 생성
        progress_placeholder.info("🎯 1/4 타겟 분석 중...")
        target_result = suggest_targets(topic)
        targets = parse_json(target_result)
        if targets and targets.get('targets'):
            first_target = targets['targets'][0]
            persona = f"{first_target.get('name', '')} - {first_target.get('description', '')}"
            st.session_state['target_persona'] = persona

            # 페인포인트 분석
            pain_result = analyze_pains_deep(topic, persona)
            pain_data = parse_json(pain_result)
            if pain_data:
                st.session_state['pains'] = pain_data.get('pains', [])

        # 2. 목차 자동 생성
        progress_placeholder.info("📋 2/4 목차 생성 중...")
        outline_result = generate_outline(
            topic,
            st.session_state.get('target_persona', ''),
            st.session_state.get('pains', [])
        )

        # 목차 텍스트 파싱 (PAGE 4와 동일한 방식)
        if outline_result:
            chapters = []
            subtopics = {}
            current_ch = None

            for line in outline_result.split('\n'):
                orig_line = line
                line = line.strip()
                if not line:
                    continue

                # 챕터 감지 (PART, 1., 2. 등)
                is_chapter = False
                ch_name = None

                if re.match(r'^(PART|파트)\s*\d+', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = line
                elif re.match(r'^\d+[\.\)]\s', line):
                    is_chapter = True
                    ch_name = line
                elif re.match(r'^#+\s*(PART|파트|\d+)', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = re.sub(r'^#+\s*', '', line)

                if is_chapter and ch_name:
                    ch_name = ch_name.replace('**', '').replace('*', '').replace('#', '').strip()
                    if ch_name and len(ch_name) > 3:
                        current_ch = ch_name
                        if current_ch not in chapters:
                            chapters.append(current_ch)
                            subtopics[current_ch] = []
                elif current_ch:
                    # 소제목 감지
                    is_subtopic = False
                    st_name = None

                    if line.startswith('-') or line.startswith('•') or line.startswith('·'):
                        is_subtopic = True
                        st_name = line.strip().lstrip('-•· ')
                    elif re.match(r'^\s+[\da-z][\)\.]', orig_line):
                        is_subtopic = True
                        st_name = re.sub(r'^[\s\da-z\)\.\-]+', '', line).strip()

                    if is_subtopic and st_name:
                        st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                        if st_name and len(st_name) > 3 and not re.match(r'^(PART|파트|Chapter|챕터)', st_name, re.IGNORECASE):
                            subtopics[current_ch].append(st_name)

            if chapters:
                st.session_state['outline'] = chapters
                st.session_state['chapters'] = {}
                for ch in chapters:
                    st.session_state['chapters'][ch] = {
                        'subtopics': subtopics.get(ch, []),
                        'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                    }

        # 3. 본문 자동 생성
        progress_placeholder.info("✍️ 3/4 본문 작성 중...")
        if st.session_state.get('outline') and st.session_state.get('chapters'):
            total_subtopics = sum(len(st.session_state['chapters'][ch]['subtopics']) for ch in st.session_state['outline'])
            done = 0

            for ch in st.session_state['outline']:
                ch_data = st.session_state['chapters'][ch]
                for sub in ch_data['subtopics']:
                    done += 1
                    progress_placeholder.info(f"✍️ 본문 작성 중... ({done}/{total_subtopics})")

                    content = generate_content_premium(sub, ch, [], [], topic, st.session_state.get('target_persona', ''))
                    if content:
                        ch_data['subtopic_data'][sub]['content'] = content
                        ch_data['subtopic_data'][sub]['formatted'] = format_content_html(content)

        # 4. 완료
        progress_placeholder.success("✅ 완료! 본문 페이지로 이동합니다...")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False


# ==========================================
# AI 함수들
# ==========================================
def analyze_market_deep(topic):
    prompt = f"""주제: {topic}

이 주제로 전자책 시장을 분석해주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요.

JSON:
{{
    "verdict": "강력 추천/추천/보류/비추천 중 하나",
    "verdict_reason": "판정 이유 한국어로",
    "total_score": 85,
    "search_data": {{
        "naver_monthly": "네이버 월간 검색량 예시: 12,000회",
        "google_monthly": "구글 월간 검색량 예시: 8,500회",
        "naver_blog_posts": "블로그 게시물 수",
        "youtube_videos": "유튜브 영상 수",
        "search_trend": "상승 또는 유지 또는 하락"
    }},
    "market_size": {{
        "score": 85,
        "level": "매우 큼/큼/보통/작음 중 하나",
        "analysis": "분석 2문장 한국어로"
    }},
    "competition": {{
        "score": 70,
        "level": "치열함/보통/낮음 중 하나",
        "your_opportunity": "차별화 기회 한국어로"
    }},
    "profit": {{
        "score": 80,
        "price_range": "권장 가격대",
        "monthly_revenue": "예상 월 수익"
    }},
    "popular_ebooks": [
        {{
            "title": "이 주제 관련 인기 전자책 제목",
            "platform": "크몽/탈잉/클래스101/리디북스/yes24 중 하나",
            "url": "해당 전자책 실제 URL (예: https://kmong.com/xxx)",
            "price": "가격"
        }},
        {{
            "title": "두번째 인기 전자책",
            "platform": "플랫폼명",
            "url": "URL",
            "price": "가격"
        }},
        {{
            "title": "세번째 인기 전자책",
            "platform": "플랫폼명",
            "url": "URL",
            "price": "가격"
        }}
    ],
    "recommendation": "최종 권장 2문장 한국어로"
}}"""
    return ask_ai(prompt, 0.5)


def suggest_targets(topic):
    prompt = f"""주제: {topic}

이 주제의 전자책을 가장 많이 구매할 것 같은 핵심 타겟 3개만 추천해주세요.
가장 적합하고 구매 가능성이 높은 타겟만 엄선해서 3개만 알려주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요.

JSON:
{{
    "personas": [
        {{
            "name": "타겟 이름 (구체적으로)",
            "demographics": "연령대, 직업",
            "needs": "이 타겟이 이 책을 사는 이유",
            "pain_points": ["핵심 고민1", "고민2", "고민3", "고민4", "고민5"]
        }}
    ]
}}"""
    return ask_ai(prompt, 0.7)


def analyze_pains_deep(topic, persona):
    prompt = f"""주제: {topic}
타겟: {persona}

이 타겟의 고민을 아주 깊이 분석해주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요. 외국어 사용 금지.

JSON:
{{
    "surface_pains": {{
        "pains": ["표면적 고민1", "고민2", "고민3", "고민4", "고민5"],
        "description": "표면적 고민 설명 3문장"
    }},
    "hidden_pains": {{
        "pains": ["숨겨진 진짜 고민1", "고민2", "고민3", "고민4"],
        "description": "숨겨진 고민 설명 3문장"
    }},
    "emotional_pains": {{
        "pains": ["감정적 고통1", "고통2", "고통3"],
        "description": "감정적 고통 설명 2문장"
    }},
    "failed_attempts": {{
        "attempts": ["시도했지만 실패한 것1", "것2", "것3"],
        "why_failed": "실패 이유 2문장"
    }},
    "dream_outcome": {{
        "ideal_result": "이상적인 결과",
        "timeline": "원하는 기간",
        "what_changes": "달라지는 것 2문장"
    }},
    "buying_triggers": {{
        "triggers": ["구매 요인1", "요인2", "요인3"],
        "objections": ["망설임 이유1", "이유2"]
    }},
    "marketing_hook": "마케팅 훅 한 문장"
}}"""
    return ask_ai(prompt, 0.6)


def analyze_competitor_reviews(topic):
    prompt = f"""주제: {topic}

이 주제 관련 전자책/도서의 부정적 리뷰를 분석해주세요.

[매우 중요]
- 모든 답변은 반드시 한국어로만 작성하세요.
- 영어, 러시아어 등 외국어 절대 사용 금지
- 한글과 숫자만 사용하세요.

JSON:
{{
    "analysis_scope": {{
        "books_analyzed": "287권",
        "reviews_analyzed": "3,842개",
        "negative_reviews": "892개 (23%)",
        "platforms": ["크몽", "예스24", "알라딘", "교보문고"]
    }},
    "negative_patterns": [
        {{
            "pattern": "불만 패턴 한국어로",
            "frequency": "67%",
            "example_reviews": ["실제 리뷰 예시 한국어로", "리뷰2"],
            "reader_emotion": "독자 감정 한국어로",
            "hidden_need": "숨겨진 니즈 한국어로",
            "solution": "해결책 한국어로"
        }},
        {{
            "pattern": "두 번째 불만",
            "frequency": "54%",
            "example_reviews": ["리뷰1", "리뷰2"],
            "reader_emotion": "감정",
            "hidden_need": "니즈",
            "solution": "해결책"
        }},
        {{
            "pattern": "세 번째 불만",
            "frequency": "41%",
            "example_reviews": ["리뷰1", "리뷰2"],
            "reader_emotion": "감정",
            "hidden_need": "니즈",
            "solution": "해결책"
        }}
    ],
    "hidden_needs_summary": {{
        "needs": ["숨겨진 니즈1", "니즈2", "니즈3"],
        "insight": "핵심 인사이트 2문장"
    }},
    "concept_suggestions": [
        {{
            "concept": "차별화 컨셉1 한국어로",
            "why_works": "이유 한국어로",
            "unique_point": "차별점 한국어로"
        }},
        {{
            "concept": "컨셉2",
            "why_works": "이유",
            "unique_point": "차별점"
        }}
    ],
    "success_formula": {{
        "must_have": ["필수1", "필수2", "필수3"],
        "must_avoid": ["금지1", "금지2"],
        "differentiation": "차별화 전략 한국어로 2문장"
    }}
}}"""
    return ask_ai(prompt, 0.6)


def generate_titles_bestseller(topic, persona, pains):
    prompt = f"""당신은 교보문고 베스트셀러 TOP 20 제목만 분석하는 전문가입니다.

주제: {topic}
독자 고민: {pains}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🏆 교보문고 TOP 20 제목 패턴 분석
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[패턴 1: 한 단어 임팩트]
• 역행자 → 동사의 명사화, 강렬한 의미
• 초격차 → 신조어 창조
• 불변 → 한자어 한 단어
• 그릿 → 영어 한 단어

[패턴 2: 명사+의+명사 (새로운 개념)]
• 돈의 속성 → 익숙한 것에 낯선 단어 결합
• 부의 추월차선 → 은유적 표현
• 생각의 지도 → 추상적 조합

[패턴 3: 형용사+명사]
• 단단한 삶 → 고급스러운 형용사
• 고요한 용기 → 역설적 조합
• 아주 작은 습관의 힘 → 구체적 수식

[패턴 4: 영문 느낌]
• 언스크립티드 → 영어 그대로
• 더 해빙 → 영어+한글

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⛔ 절대 금지 (유치한 제목 = 판매 실패)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[유튜브/블로그 냄새]
X "~하는 법", "~하는 방법", "~가이드", "~노하우"
X "월 1000만원", "100일 만에" (숫자 과시)
X "부자되는", "성공하는", "돈버는" (뻔한 동사)

[인터넷 광고 냄새]
X "비밀", "비법", "공식", "정석", "바이블"
X "마법", "연금술", "황금", "보물", "열쇠"
X "머니", "캐시", "머신", "시스템", "파이프라인"
X "터보", "부스터", "로켓"

[게임/판타지 냄새]
X "레벨업", "스킬", "공략", "정복", "마스터"
X "무기", "전투", "퀘스트"

[과장/저렴함]
X "완벽한", "궁극의", "최고의", "기적의"
X "30대를 위한", "직장인을 위한" (타겟 명시)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 제목 검증 기준
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. 교보문고 베스트셀러 코너에 진열해도 품격이 유지되는가?
2. "역행자" 옆에 놓아도 어색하지 않은가?
3. 유튜브 썸네일이나 블로그 제목 같지 않은가?
4. 1~3단어로 강렬한가?
5. "이게 뭐지?" 궁금증이 생기는가?

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

'{topic}' 주제로 교보문고 품격의 제목 5개 생성.
각 제목은 서로 다른 패턴으로.

JSON:
{{
    "titles": [
        {{"title": "제목 (1~3단어)", "subtitle": "부제 (15자 이내)", "concept": "컨셉 한줄"}},
        {{"title": "제목", "subtitle": "부제", "concept": "컨셉"}},
        {{"title": "제목", "subtitle": "부제", "concept": "컨셉"}},
        {{"title": "제목", "subtitle": "부제", "concept": "컨셉"}},
        {{"title": "제목", "subtitle": "부제", "concept": "컨셉"}}
    ]
}}"""
    return ask_ai(prompt, 0.75)


def analyze_text_content(text, source=""):
    prompt = f"""출처: {source}
내용: {text[:5000]}

분석:

JSON:
{{
    "title": "주제",
    "key_points": ["핵심1", "핵심2", "핵심3", "핵심4", "핵심5"],
    "insights": ["인사이트1", "인사이트2", "인사이트3"],
    "action_items": ["실행1", "실행2", "실행3"],
    "ebook_ideas": ["아이디어1", "아이디어2"],
    "summary": "요약 3문장"
}}"""
    return ask_ai(prompt, 0.5)


def summarize_all_knowledge(items, topic):
    """전체 학습 내용 통합 요약"""
    all_points = []
    all_tips = []
    all_ideas = []

    for item in items:
        if isinstance(item, dict):
            all_points.extend(item.get('key_points', []))
            all_tips.extend(item.get('actionable_tips', item.get('action_items', [])))
            all_ideas.extend(item.get('ebook_applications', item.get('ebook_ideas', [])))

    prompt = f"""전자책 주제: {topic}

학습한 모든 정보를 통합 분석해주세요.

수집된 핵심 포인트들:
{chr(10).join([f"• {p}" for p in all_points[:25]])}

실행 팁들:
{chr(10).join([f"• {t}" for t in all_tips[:15]])}

전자책 활용 아이디어:
{chr(10).join([f"• {i}" for i in all_ideas[:10]])}

JSON:
{{
    "integrated_summary": "전체 학습 내용 통합 요약 5문장",
    "core_insights": [
        "핵심 인사이트 1",
        "인사이트 2",
        "인사이트 3",
        "인사이트 4",
        "인사이트 5"
    ],
    "action_plan": [
        "즉시 실행할 것 1",
        "실행 2",
        "실행 3"
    ],
    "ebook_structure": [
        "추천 목차 1장",
        "2장",
        "3장",
        "4장"
    ],
    "unique_angle": "이 전자책만의 차별화된 관점",
    "study_plan": {{
        "week1": "1주차: 무엇을 할지",
        "week2": "2주차: 무엇을 할지",
        "week3": "3주차: 무엇을 할지",
        "week4": "4주차: 무엇을 할지"
    }},
    "expert_tips": [
        "전문가 팁 1",
        "팁 2",
        "팁 3"
    ]
}}"""
    return ask_ai(prompt, 0.6)


def generate_outline(topic, persona, pains, gaps=None):
    """자청/프드프 스타일 자극적 목차 생성"""
    prompt = f"""당신은 자청입니다. 목차만 보고 결제하게 만드세요.

[주제]: {topic}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📚 자청/프드프 실제 목차 (이대로 써라)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[자청 - 역행자]
PART 1. 자의식 해체
- 왜 당신의 노력은 물거품이 되는가
- 유전자를 역행하는 법

PART 2. 정체성
- 스스로를 속여야 이긴다
- 운은 컨트롤할 수 있다

[프드프 스타일]
PART 1. 착각
- 열심히 하면 된다는 거짓말
- 99%가 실패하는 진짜 이유

PART 2. 구조
- 돈이 들어오는 시스템
- 1%만 아는 수익 공식

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 자극적 소제목 공식 (이 중에서 골라 써라)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

• "99%가 모르는 ~"
• "왜 ~하면 망하는가"
• "~는 거짓말이다"
• "~의 진짜 이유"
• "이것만 알면 상위 1%"
• "딱 하나만 바꿨더니"

[챕터명] - 2~4글자, 명사형
착각 / 거짓말 / 함정 / 구조 / 공식 / 실행 / 전환 / 도약

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 유치한 표현 (하나라도 있으면 탈락)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

나침반/지도/열쇠/보물/황금/마법/연금술/비밀무기
첫걸음/가이드/완벽/핵심전략/기초/중급/고급
로켓/터보/엔진/머신/파이프라인/저수지/샘물
~의 모든 것/~하는 방법/효과적인/성공적인
"~이 보이는 나침반", "~을 여는 열쇠", "~의 비밀"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 콜론(:) 사용 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ "PART 1. 착각: 열심히..." (콜론 있음)
✅ "PART 1. 착각" (콜론 없음)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

## PART 1. [2~4글자 명사]
- [자극적 소제목]
- [자극적 소제목]
- [자극적 소제목]

## PART 2. [2~4글자 명사]
- [자극적 소제목]
- [자극적 소제목]
- [자극적 소제목]

## PART 3. [2~4글자 명사]
- [자극적 소제목]
- [자극적 소제목]
- [자극적 소제목]

## PART 4. [2~4글자 명사]
- [자극적 소제목]
- [자극적 소제목]
- [자극적 소제목]

## PART 5. [2~4글자 명사]
- [자극적 소제목]
- [자극적 소제목]
- [자극적 소제목]

목차만 출력. 콜론(:) 절대 사용 금지."""
    return ask_ai(prompt, 0.4)


def generate_content_premium(subtopic, chapter, questions, answers, topic, persona):
    """짧지만 밀도 높은 몰입형 글쓰기"""
    qa_pairs = ""
    for i, (q, a) in enumerate(zip(questions, answers), 1):
        if a.strip():
            qa_pairs += f"\n질문{i}: {q}\n답변{i}: {a}\n"

    prompt = f"""'{subtopic}'에 대해 글을 씁니다.

[주제]: {topic}
[챕터]: {chapter}
[참고 내용]
{qa_pairs}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
핵심 원칙: 짧게. 임팩트 있게.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1500~1800자. 이 안에서 끝내라.
길면 지루하다. 짧고 강렬하게.
핵심만 남기고 다 쳐내라.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 글 구조: 반전이 있어야 한다
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[도입] 후킹 (1~2문단)
- "대부분 여기서 틀립니다." / "솔직히 말할게요."

[전개] 상식 뒤집기 (2문단)
- "저도 그렇게 생각했습니다. 근데 반대였습니다."

[반전] 핵심 인사이트 (2~3문단)
- "근데 진짜 중요한 건 따로 있었습니다."
- 숫자로 증명

[마무리] 액션 (1문단)
- 당장 할 수 있는 것 하나

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 전문성 = 인사이더 정보
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

읽고 나서 "이 사람 진짜 해본 사람이네" 느끼게 해라.

이렇게 써라:
- "대부분 A라고 생각하는데, 실제로는 B입니다"
- "책에선 안 나오는데, 현장에서는..."
- "97%가 모르는 게 있습니다"
- "저도 3번 실패하고 나서야 알았습니다"

구체적인 숫자를 써라:
- ✗ "많이 벌었습니다" → ✓ "월 340만원"
- ✗ "오래 걸렸습니다" → ✓ "47일"
- ✗ "효과가 좋습니다" → ✓ "전환율 3.7배"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚨 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 유치한 비유 (이거 쓰면 글 망함)
- "마법", "연금술", "황금", "열쇠", "보물", "나침반", "지도"
- "저수지", "날개", "로켓", "엔진", "머신", "파이프라인"
- "~이 보이는", "~을 여는", "~의 세계로", "~의 비밀"
→ 비유 쓰지 마라. 사실만 담담하게.

❌ 설명형 글쓰기
- "OOO란 무엇인가요?"
- "~에 대해 알아보겠습니다"
→ 설명하지 말고 보여줘라

❌ AI 표현
- "중요합니다", "따라서", "결론적으로"
- "~하는 것이 좋습니다"

❌ 타겟 직접 언급 (절대 금지!)
- "김대리님", "30대 직장인 여러분" 금지
- "주부 여러분", "초보자 분들" 금지
- "당신"도 남발 금지
- "저", "우리" 정도만 써라

❌ 뻔한 조언
- "꾸준히 하세요", "포기하지 마세요"

❌ 소제목/번호
- "1. / 2. / 3." 나열 금지
- 소설처럼 자연스럽게 연결

❌ 마크다운/이모지
- **굵게**, - 글머리, 🔥 금지

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 예시 (이 길이로 써라)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

"대부분 A가 정답이라고 생각합니다. 저도 그랬습니다.

근데 A를 3개월 했을 때 결과는 0이었습니다.

B로 바꿨습니다. 2주 만에 47만원이 찍혔습니다.

핵심은 간단합니다. [구체적 방법]. 업계에선 이걸 'OOO'라고 부릅니다.

지금 바로 해보세요."

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📏 분량: 1500~1800자 (절대 넘기지 마라)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 짧은 문장. 2~3문장이 한 문단.
- 문체: ~입니다 / ~거든요 / ~더라고요
- 도표는 꼭 필요할 때만"""
    return ask_ai(prompt, 0.75)


def format_content_html(content):
    """본문을 HTML 형식으로 변환 (강조 표시 적용)"""
    if not content:
        return ""
    # 「」 → 주황색 볼드
    formatted = re.sub(r'「([^」]+)」', r'<b style="color:#e67e22;">\1</b>', content)
    # ★ → 주황색 볼드 문장
    formatted = re.sub(r'★\s*(.+?)(?=\n|$)', r'<p style="color:#e67e22;font-weight:700;margin:20px 0;font-size:17px;">★ \1</p>', formatted)
    # 문단 구분 (빈 줄) → 문단 간격
    formatted = formatted.replace('\n\n', '</p><p style="color:#000000;margin:25px 0;line-height:2.0;font-size:17px;">')
    # 단일 줄바꿈 제거 (문단 내 연결)
    formatted = formatted.replace('\n', ' ')
    formatted = f'<p style="color:#000000;margin:25px 0;line-height:2.0;font-size:17px;">{formatted}</p>'
    return formatted


def generate_questions(subtopic, chapter, topic):
    prompt = f"""'{topic}' 전자책 '{chapter}' 챕터의 '{subtopic}' 작성용 질문 3개:

Q1: [질문]
Q2: [질문]
Q3: [질문]"""
    return ask_ai(prompt, 0.7)


# ==========================================
# 메인 UI
# ==========================================
# 비디오 배경 헤더
video_path = "/Users/hyunwoo/Desktop/title_bg.mp4"
header_video_b64 = get_video_base64(video_path)

if header_video_b64:
    st.markdown(f"""
    <style>
    @keyframes writeyGlow {{
        0%, 100% {{ text-shadow: 0 0 20px rgba(212,175,55,0.5), 0 0 40px rgba(212,175,55,0.3), 0 0 60px rgba(212,175,55,0.2); }}
        50% {{ text-shadow: 0 0 30px rgba(212,175,55,0.8), 0 0 60px rgba(212,175,55,0.5), 0 0 90px rgba(212,175,55,0.3); }}
    }}
    @keyframes gradientMove {{
        0% {{ background-position: 0% 50%; }}
        50% {{ background-position: 100% 50%; }}
        100% {{ background-position: 0% 50%; }}
    }}
    @keyframes subtitleFade {{
        0%, 100% {{ opacity: 0.7; }}
        50% {{ opacity: 1; }}
    }}
    .writey-title {{
        font-family: 'Playfair Display', 'Cormorant Garamond', serif !important;
        font-size: 90px !important;
        font-weight: 700 !important;
        font-style: italic;
        background: linear-gradient(135deg, #fff 0%, #d4af37 25%, #fff 50%, #d4af37 75%, #fff 100%);
        background-size: 300% 300%;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        animation: gradientMove 4s ease infinite, writeyGlow 3s ease-in-out infinite;
        margin: 0;
        letter-spacing: 8px;
    }}
    .writey-subtitle {{
        font-family: 'Cinzel', serif !important;
        color: rgba(212,175,55,0.9) !important;
        font-size: 15px !important;
        letter-spacing: 8px !important;
        margin-bottom: 15px;
        font-weight: 500;
        text-transform: uppercase;
    }}
    .writey-tagline {{
        color: rgba(255,255,255,0.85) !important;
        font-size: 20px !important;
        margin-top: 20px;
        font-weight: 300;
        letter-spacing: 3px;
        animation: subtitleFade 4s ease-in-out infinite;
    }}
    </style>
    <div style="position:relative;border-radius:24px;overflow:hidden;margin-bottom:35px;box-shadow:0 15px 50px rgba(0,0,0,0.5), 0 0 100px rgba(212,175,55,0.1);">
        <video autoplay muted loop playsinline style="width:100%;height:320px;object-fit:cover;filter:brightness(0.3) saturate(1.2);">
            <source src="data:video/mp4;base64,{header_video_b64}" type="video/mp4">
        </video>
        <div style="position:absolute;top:0;left:0;right:0;bottom:0;background:linear-gradient(180deg, rgba(0,0,0,0.2) 0%, rgba(0,0,0,0.4) 100%);"></div>
        <div style="position:absolute;top:0;left:0;right:0;bottom:0;display:flex;flex-direction:column;justify-content:center;align-items:center;text-align:center;">
            <div class="writey-subtitle">✦ CASHMAKER ✦</div>
            <h1 class="writey-title">Writey</h1>
            <p class="writey-tagline">아이디어부터 출판까지, AI 원스톱 전자책 제작</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <style>
    @keyframes writeyGlowFallback {{
        0%, 100% {{ text-shadow: 0 0 20px rgba(212,175,55,0.5); }}
        50% {{ text-shadow: 0 0 40px rgba(212,175,55,0.8); }}
    }}
    </style>
    <div style="text-align:center;padding:60px 20px;margin-bottom:30px;background:linear-gradient(180deg, rgba(20,20,20,0.9) 0%, rgba(10,10,10,0.95) 100%);border-radius:20px;border:1px solid rgba(212,175,55,0.2);">
        <div style="color:rgba(212,175,55,0.9);font-size:14px;letter-spacing:6px;margin-bottom:15px;font-weight:500;">✦ CASHMAKER ✦</div>
        <h1 style="font-family:'Playfair Display',serif;font-size:72px;font-weight:700;font-style:italic;color:#d4af37;margin:0;letter-spacing:6px;animation:writeyGlowFallback 3s ease-in-out infinite;">Writey</h1>
        <p style="color:rgba(255,255,255,0.7);font-size:18px;margin-top:20px;letter-spacing:2px;">아이디어부터 출판까지, AI 원스톱 전자책 제작</p>
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 인터뷰 모드 (interview_completed가 False일 때)
# ==========================================
if not st.session_state.get('interview_completed', False):
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">INTERVIEW</span>
        <h2>나만의 전자책 만들기</h2>
        <p>몇 가지 질문에 답하면 AI가 전자책을 완성해드립니다</p>
    </div>
    """, unsafe_allow_html=True)

    # 인터뷰 진행 상태
    if 'interview_step' not in st.session_state:
        st.session_state['interview_step'] = 1

    step = st.session_state['interview_step']
    total_steps = 6

    # 진행률 표시
    st.progress(step / total_steps)
    st.caption(f"질문 {step} / {total_steps}")

    st.markdown("---")

    # 인터뷰 데이터 임시 저장
    if 'temp_interview' not in st.session_state:
        st.session_state['temp_interview'] = {}

    # ========== STEP 1: 기본 정보 ==========
    if step == 1:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(212,175,55,0.1) 0%, rgba(183,110,121,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">👋 먼저 당신에 대해 알려주세요</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">전자책의 저자로서 기본 정보를 입력해주세요</p>
        </div>
        """, unsafe_allow_html=True)

        with st.form(key="step1_form"):
            author_name = st.text_input("저자명 (필명 가능)", value=st.session_state['temp_interview'].get('author_name', ''), placeholder="예: 김성장, 머니메이커 등")
            field = st.text_input("당신의 전문 분야는?", value=st.session_state['temp_interview'].get('field', ''), placeholder="예: 주식투자, 블로그 수익화, 다이어트, 영어회화 등")

            exp_options = ["선택하세요", "1년 미만", "1~2년", "3~5년", "5~10년", "10년 이상"]
            saved_exp = st.session_state['temp_interview'].get('experience_years', '선택하세요')
            exp_index = exp_options.index(saved_exp) if saved_exp in exp_options else 0
            experience = st.selectbox("이 분야 경험은?", exp_options, index=exp_index)

            col1, col2 = st.columns([1, 1])
            with col2:
                submitted = st.form_submit_button("다음 →", use_container_width=True, type="primary")

            if submitted:
                if not author_name.strip() or not field.strip() or experience == "선택하세요":
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['author_name'] = author_name.strip()
                    st.session_state['temp_interview']['field'] = field.strip()
                    st.session_state['temp_interview']['experience_years'] = experience
                    st.session_state['interview_step'] = 2
                    st.rerun()

    # ========== STEP 2: 주제와 노하우 ==========
    elif step == 2:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(212,175,55,0.1) 0%, rgba(183,110,121,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">📚 어떤 내용을 담을까요?</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">당신만의 핵심 노하우를 알려주세요</p>
        </div>
        """, unsafe_allow_html=True)

        col_prev, col_next = st.columns([1, 1])
        with col_prev:
            if st.button("← 이전", key="interview_prev_2", use_container_width=True):
                st.session_state['interview_step'] = 1
                st.rerun()

        with st.form(key="step2_form"):
            topic = st.text_input("전자책 주제", value=st.session_state['temp_interview'].get('topic', ''), placeholder="예: 월 100만원 배당 투자, 하루 1시간 블로그로 월 300 벌기")
            core_method = st.text_area("당신만의 핵심 방법/노하우는?", value=st.session_state['temp_interview'].get('core_method', ''), height=120, placeholder="예: 저는 고배당 ETF를 활용해서 안정적으로 수익을 내는 방법을 알려드립니다. 핵심은 분산투자와 복리의 마법입니다...")

            col1, col2 = st.columns([1, 1])
            with col2:
                submitted = st.form_submit_button("다음 →", use_container_width=True, type="primary")

            if submitted:
                if not topic.strip() or not core_method.strip():
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['topic'] = topic.strip()
                    st.session_state['temp_interview']['core_method'] = core_method.strip()
                    st.session_state['interview_step'] = 3
                    st.rerun()

    # ========== STEP 3: 타겟 독자 (AI 추천) ==========
    elif step == 3:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(212,175,55,0.1) 0%, rgba(183,110,121,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">🎯 누구를 위한 책인가요?</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">AI가 시장 데이터를 분석해 최적의 타겟을 추천해드립니다</p>
        </div>
        """, unsafe_allow_html=True)

        topic = st.session_state['temp_interview'].get('topic', '')

        # AI 타겟 분석 (캐시)
        if 'ai_target_suggestions' not in st.session_state or st.session_state.get('ai_target_topic') != topic:
            if st.button("🔍 AI 타겟 분석 시작", key="analyze_target", use_container_width=True, type="primary"):
                with st.spinner("시장 데이터 분석 중..."):
                    result = suggest_targets(topic)
                    parsed = parse_json(result)
                    if parsed and parsed.get('personas'):
                        st.session_state['ai_target_suggestions'] = parsed['personas']
                        st.session_state['ai_target_topic'] = topic
                        st.rerun()
                    else:
                        st.error("분석 실패. 다시 시도해주세요.")

        # AI 추천 결과 표시
        if st.session_state.get('ai_target_suggestions'):
            st.markdown("### 📊 AI 추천 타겟")
            personas = st.session_state['ai_target_suggestions']

            selected_idx = st.session_state.get('selected_target_idx', 0)

            for idx, persona in enumerate(personas[:3]):
                is_selected = (idx == selected_idx)
                border_color = "var(--gold)" if is_selected else "var(--line)"
                bg_color = "rgba(212,175,55,0.1)" if is_selected else "rgba(20,20,20,0.5)"

                pain_list = persona.get('pain_points', [])[:3]
                pains_text = " / ".join(pain_list) if pain_list else "고민 분석 중..."

                st.markdown(f"""
                <div style="background:{bg_color};border:1px solid {border_color};border-radius:10px;padding:15px;margin-bottom:10px;">
                    <div style="font-weight:bold;color:var(--gold);margin-bottom:5px;">{persona.get('name', '타겟')}</div>
                    <div style="font-size:13px;color:var(--text2);margin-bottom:8px;">{persona.get('demographics', '')}</div>
                    <div style="font-size:12px;color:var(--text);opacity:0.8;">💭 {pains_text}</div>
                </div>
                """, unsafe_allow_html=True)

                if st.button(f"✓ 이 타겟 선택" if not is_selected else "✓ 선택됨", key=f"select_target_{idx}", use_container_width=True, disabled=is_selected):
                    st.session_state['selected_target_idx'] = idx
                    st.rerun()

            st.markdown("---")

            # 선택된 타겟 정보 자동 입력
            selected_persona = personas[selected_idx] if selected_idx < len(personas) else personas[0]
            default_reader = f"{selected_persona.get('name', '')} ({selected_persona.get('demographics', '')})"
            default_problem = " ".join(selected_persona.get('pain_points', [])[:3])

            st.markdown("##### 선택된 타겟 정보 (수정 가능)")
            target_reader = st.text_input("타겟 독자", value=st.session_state['temp_interview'].get('target_reader', '') or default_reader, key="target_reader_input")
            target_problem = st.text_area("이 독자들의 가장 큰 고민/문제", value=st.session_state['temp_interview'].get('target_problem', '') or default_problem, height=80, key="target_problem_input")

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("← 이전", key="interview_prev_3", use_container_width=True):
                    st.session_state['interview_step'] = 2
                    st.rerun()
            with col2:
                if st.button("다음 →", key="interview_next_3", use_container_width=True, type="primary"):
                    if not target_reader or not target_problem:
                        st.error("타겟 독자와 고민을 입력해주세요")
                    else:
                        st.session_state['temp_interview']['target_reader'] = target_reader
                        st.session_state['temp_interview']['target_problem'] = target_problem
                        st.session_state['interview_step'] = 4
                        st.rerun()
        else:
            # AI 분석 전 직접 입력 옵션
            st.markdown("---")
            st.markdown("##### 또는 직접 입력")
            target_reader = st.text_input("타겟 독자", value=st.session_state['temp_interview'].get('target_reader', ''), placeholder="예: 30대 직장인, 투자 초보자")
            target_problem = st.text_area("이 독자들의 가장 큰 고민/문제는?", value=st.session_state['temp_interview'].get('target_problem', ''), height=80, placeholder="예: 월급만으로는 부족하고, 어디서부터 시작해야 할지 모르겠다...")

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("← 이전", key="interview_prev_3_manual", use_container_width=True):
                    st.session_state['interview_step'] = 2
                    st.rerun()
            with col2:
                if st.button("다음 →", key="interview_next_3_manual", use_container_width=True, type="primary"):
                    if not target_reader or not target_problem:
                        st.error("타겟 독자와 고민을 입력해주세요")
                    else:
                        st.session_state['temp_interview']['target_reader'] = target_reader
                        st.session_state['temp_interview']['target_problem'] = target_problem
                        st.session_state['interview_step'] = 4
                        st.rerun()

    # ========== STEP 4: 스토리 & 경력 ==========
    elif step == 4:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(212,175,55,0.1) 0%, rgba(183,110,121,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">💪 당신의 이야기를 들려주세요</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">독자들이 공감할 수 있는 진솔한 경험담과 경력</p>
        </div>
        """, unsafe_allow_html=True)

        struggle_story = st.text_area("처음 시작할 때 겪었던 어려움/실패는?", value=st.session_state['temp_interview'].get('struggle_story', ''), height=100, placeholder="예: 처음에는 주식으로 500만원을 잃었습니다. 유튜브 정보만 믿고 투자했다가 큰 손실을 봤죠...")
        breakthrough = st.text_area("어떻게 극복하고 성과를 냈나요?", value=st.session_state['temp_interview'].get('breakthrough', ''), height=100, placeholder="예: 그 후 기본서 10권을 정독하고, 나만의 원칙을 세웠습니다. 1년 후 손실을 모두 만회하고 수익을 내기 시작했습니다...")

        st.markdown("---")
        st.markdown("##### 📌 작가 경력/경험 (선택)")
        author_career = st.text_area("관련 경력이나 자격, 성과가 있다면?", value=st.session_state['temp_interview'].get('author_career', ''), height=100, placeholder="예: 금융회사 7년 근무, 투자 관련 유튜브 구독자 5만명, 월 수익 3천만원 달성, CFA 자격증 보유, 강의 경력 3년...")

        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_4", use_container_width=True):
                st.session_state['interview_step'] = 3
                st.rerun()
        with col2:
            if st.button("다음 →", key="interview_next_4", use_container_width=True, type="primary"):
                if not struggle_story or not breakthrough:
                    st.error("어려움/실패와 극복 스토리는 필수입니다")
                else:
                    st.session_state['temp_interview']['struggle_story'] = struggle_story
                    st.session_state['temp_interview']['breakthrough'] = breakthrough
                    st.session_state['temp_interview']['author_career'] = author_career
                    st.session_state['interview_step'] = 5
                    st.rerun()

    # ========== STEP 5: 마무리 ==========
    elif step == 5:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(212,175,55,0.1) 0%, rgba(183,110,121,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">✨ 마지막으로!</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">독자에게 전하고 싶은 메시지</p>
        </div>
        """, unsafe_allow_html=True)

        why_write = st.text_area("왜 이 책을 쓰려고 하나요?", value=st.session_state['temp_interview'].get('why_write', ''), height=80, placeholder="예: 저처럼 헤매는 사람들이 시행착오 없이 바로 성과를 낼 수 있도록 도와주고 싶습니다...")
        final_message = st.text_area("독자에게 마지막으로 전하고 싶은 말", value=st.session_state['temp_interview'].get('final_message', ''), height=80, placeholder="예: 누구나 할 수 있습니다. 포기하지 않으면 반드시 성공합니다...")

        # 입력 내용 미리보기
        st.markdown("---")
        st.markdown("### 📋 입력 내용 확인")

        preview_data = st.session_state['temp_interview']
        st.markdown(f"""
        <div style="background:rgba(20,20,20,0.8);padding:20px;border-radius:10px;border:1px solid var(--line);">
            <p><b>저자:</b> {preview_data.get('author_name', '')}</p>
            <p><b>분야:</b> {preview_data.get('field', '')} ({preview_data.get('experience_years', '')})</p>
            <p><b>주제:</b> {preview_data.get('topic', '')}</p>
            <p><b>타겟:</b> {preview_data.get('target_reader', '')}</p>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_5", use_container_width=True):
                st.session_state['interview_step'] = 4
                st.rerun()
        with col2:
            if st.button("📋 목차 생성하기", key="interview_generate_outline", use_container_width=True, type="primary"):
                if not get_api_key():
                    st.error("사이드바에서 API 키를 먼저 입력해주세요")
                elif not why_write or not final_message:
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['why_write'] = why_write
                    st.session_state['temp_interview']['final_message'] = final_message

                    # 목차만 먼저 생성
                    progress_box = st.empty()
                    interview_data = st.session_state['temp_interview']
                    success = generate_outline_only(interview_data, progress_box)

                    if success:
                        import time
                        time.sleep(1)
                        st.session_state['interview_step'] = 6  # 목차 확인 단계로 이동
                        st.rerun()

    # ========== STEP 6: 목차 확인 및 본문 생성 ==========
    elif step == 6:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(212,175,55,0.1) 0%, rgba(183,110,121,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">📋 목차 확인 및 수정</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">생성된 목차를 확인하고, 직접 수정하거나 AI로 재생성할 수 있습니다</p>
        </div>
        """, unsafe_allow_html=True)

        # 제목 표시
        book_title = st.session_state.get('book_title', '')
        subtitle = st.session_state.get('subtitle', '')
        book_concept = st.session_state.get('book_concept', '')

        if book_title:
            st.markdown(f"""
            <div style="background:rgba(30,30,30,0.9);padding:25px;border-radius:15px;border:2px solid var(--gold);margin-bottom:20px;text-align:center;">
                <h2 style="color:var(--gold);margin:0 0 10px 0;font-size:32px;">{book_title}</h2>
                <p style="color:var(--text2);margin:0;font-size:18px;">{subtitle}</p>
            </div>
            """, unsafe_allow_html=True)

        # 컨셉 표시
        if book_concept:
            with st.expander("💡 이 책의 고유 컨셉 보기", expanded=False):
                st.markdown(f"""
                <div style="background:rgba(212,175,55,0.1);padding:20px;border-radius:10px;border-left:3px solid var(--gold);">
                    {book_concept.replace(chr(10), '<br>')}
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")

        # 목차 표시 및 편집
        outline = st.session_state.get('outline', [])
        chapters = st.session_state.get('chapters', {})

        if outline:
            st.markdown("### 📖 목차 구성")
            st.caption("각 챕터와 소제목을 직접 수정하거나, 🔄 버튼으로 AI가 새로 생성합니다")

            for i, ch in enumerate(outline):
                ch_data = chapters.get(ch, {})
                subtopics = ch_data.get('subtopics', [])

                # 챕터 헤더
                st.markdown(f"""
                <div style="background:linear-gradient(90deg, rgba(212,175,55,0.2) 0%, rgba(30,30,30,0.9) 100%);
                            padding:15px 20px;border-radius:10px;margin:20px 0 10px 0;
                            border-left:4px solid var(--gold);">
                    <span style="color:var(--gold);font-weight:bold;font-size:18px;">PART {i+1}</span>
                </div>
                """, unsafe_allow_html=True)

                # 챕터 제목 편집
                new_ch_name = st.text_input(
                    f"챕터 {i+1} 제목",
                    value=ch,
                    key=f"ch_edit_{i}",
                    label_visibility="collapsed"
                )

                # 챕터 이름 변경 적용
                if new_ch_name != ch and new_ch_name.strip():
                    # 목차에서 이름 변경
                    st.session_state['outline'][i] = new_ch_name.strip()
                    # chapters 딕셔너리에서도 키 변경
                    st.session_state['chapters'][new_ch_name.strip()] = st.session_state['chapters'].pop(ch)
                    st.rerun()

                # 소제목들
                for j, sub in enumerate(subtopics):
                    col1, col2 = st.columns([0.5, 5.5])
                    with col1:
                        st.markdown(f"<div style='color:var(--text2);padding-top:8px;'>•</div>", unsafe_allow_html=True)
                    with col2:
                        new_sub = st.text_input(
                            f"소제목 {j+1}",
                            value=sub,
                            key=f"sub_edit_{i}_{j}",
                            label_visibility="collapsed"
                        )
                        # 소제목 변경 적용
                        if new_sub != sub and new_sub.strip():
                            st.session_state['chapters'][ch]['subtopics'][j] = new_sub.strip()
                            # subtopic_data도 업데이트
                            old_data = st.session_state['chapters'][ch]['subtopic_data'].pop(sub, {'questions': [], 'answers': [], 'content': ''})
                            st.session_state['chapters'][ch]['subtopic_data'][new_sub.strip()] = old_data
                            st.rerun()

            st.markdown("---")

        # 하단 버튼
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_6", use_container_width=True):
                st.session_state['interview_step'] = 5
                st.rerun()
        with col2:
            if st.button("✍️ 본문 생성하기", key="generate_body", use_container_width=True, type="primary"):
                progress_box = st.empty()
                interview_data = st.session_state.get('interview_data', st.session_state['temp_interview'])
                success = generate_body_from_outline(interview_data, progress_box)

                if success:
                    import time
                    time.sleep(1)
                    st.session_state['current_page'] = 7  # 최종 출력 페이지로 이동
                    st.rerun()

    # 기존 방식 사용 옵션
    st.markdown("---")
    st.markdown("<div style='text-align:center;'>", unsafe_allow_html=True)
    if st.button("📝 전자책 상세 작성 (전문가용)", key="skip_interview"):
        st.session_state['interview_completed'] = True
        st.session_state['current_page'] = 0
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

    st.stop()

# ==========================================
# 여기서부터 기존 페이지 로직 (인터뷰 완료 후)
# ==========================================

# 페이지 네비게이션 (간소화: 4단계)
simple_pages = ["주제", "목차", "본문", "완성"]
page_mapping = [0, 4, 5, 7]  # 실제 페이지 인덱스
current = st.session_state['current_page']

# 현재 페이지가 간소화된 네비게이션의 어디에 해당하는지
def get_simple_index(current_page):
    if current_page <= 0:
        return 0
    elif current_page <= 4:
        return 1
    elif current_page <= 5:
        return 2
    else:
        return 3

simple_current = get_simple_index(current)

# 프리미엄 네비게이션 바 (4단계)
st.markdown('<div class="premium-nav-container">', unsafe_allow_html=True)
cols = st.columns(4)
for i, (col, page) in enumerate(zip(cols, simple_pages)):
    with col:
        if i == simple_current:
            st.markdown(f'<div class="nav-item active">{i+1}. {page}</div>', unsafe_allow_html=True)
        else:
            if st.button(f"{i+1}. {page}", key=f"nav_{i}", use_container_width=True):
                st.session_state['current_page'] = page_mapping[i]
                st.rerun()
st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

# API 키가 없으면 환영 화면 표시
if not get_api_key():
    st.markdown("""
    <div style="background:linear-gradient(135deg, rgba(212,175,55,0.2) 0%, rgba(30,30,30,0.98) 100%);
                border:3px solid rgba(212,175,55,0.6);border-radius:25px;padding:50px 40px;text-align:center;margin:20px 0;">
        <p style="font-size:60px;margin:0 0 20px 0;">👋</p>
        <h2 style="color:var(--gold);font-size:42px;margin-bottom:20px;font-weight:700;">환영합니다!</h2>
        <p style="color:var(--text);font-size:24px;margin-bottom:10px;line-height:1.8;">
            AI가 전자책을 대신 써주는 프로그램입니다
        </p>
        <p style="color:var(--text2);font-size:20px;">
            처음 한 번만 설정하면 바로 사용할 수 있어요
        </p>
    </div>
    """, unsafe_allow_html=True)

    # 큰 안내 박스
    st.markdown("""
    <div style="background:#1a1a2e;border:3px solid #e74c3c;padding:30px;border-radius:20px;margin:30px 0;">
        <p style="font-size:28px;margin:0;line-height:1.6;color:#fff;text-align:center;">
            🔑 <b style="color:#e74c3c;">첫 번째 할 일</b><br><br>
            <span style="font-size:24px;">👈 왼쪽에 <span style="color:#d4af37;font-weight:700;">"API 키"</span>를 넣어야 해요</span>
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <p style="text-align:center;font-size:32px;color:var(--gold);margin:40px 0 30px 0;font-weight:700;">
        📖 딱 3단계만 하면 끝!
    </p>
    """, unsafe_allow_html=True)

    # STEP 1 - Anthropic 가입
    st.markdown("""
    <div style="background:linear-gradient(135deg, #7c3aed 0%, #5b21b6 100%);padding:30px;border-radius:20px;margin:20px 0;">
        <p style="font-size:28px;margin:0 0 15px 0;color:#fff;font-weight:800;text-align:center;">
            1️⃣ Anthropic 회원가입
        </p>
        <div style="background:rgba(255,255,255,0.15);padding:15px;border-radius:10px;">
            <p style="font-size:16px;margin:0;color:#fff;line-height:1.8;">
                1. 아래 버튼 클릭<br>
                2. <b>"Sign up"</b> 클릭<br>
                3. Google 계정으로 가입 (가장 쉬움)
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.link_button("👆 Anthropic 가입하기", "https://console.anthropic.com/", use_container_width=True, type="primary")

    st.markdown("<p style='height:20px;'></p>", unsafe_allow_html=True)

    # STEP 2 - 결제 등록
    st.markdown("""
    <div style="background:linear-gradient(135deg, #e74c3c 0%, #c0392b 100%);padding:30px;border-radius:20px;margin:20px 0;">
        <p style="font-size:28px;margin:0 0 15px 0;color:#fff;font-weight:800;text-align:center;">
            2️⃣ 결제 수단 & 크레딧 충전
        </p>
        <div style="background:rgba(255,255,255,0.15);padding:15px;border-radius:10px;">
            <p style="font-size:16px;margin:0;color:#fff;line-height:1.8;">
                1. 로그인 후 왼쪽 메뉴 <b>"Settings"</b> 클릭<br>
                2. <b>"Billing"</b> 클릭<br>
                3. <b>"Add payment method"</b>로 카드 등록<br>
                4. <b>"Add credits"</b>로 $5~10 충전
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.link_button("👆 Billing 페이지 바로가기", "https://console.anthropic.com/settings/billing", use_container_width=True, type="primary")

    st.markdown("<p style='height:20px;'></p>", unsafe_allow_html=True)

    # STEP 3 - 키 받기
    st.markdown("""
    <div style="background:linear-gradient(135deg, #27ae60 0%, #2ecc71 100%);padding:30px;border-radius:20px;margin:20px 0;">
        <p style="font-size:28px;margin:0 0 15px 0;color:#fff;font-weight:800;text-align:center;">
            3️⃣ API 키 발급
        </p>
        <div style="background:rgba(255,255,255,0.15);padding:15px;border-radius:10px;">
            <p style="font-size:16px;margin:0;color:#fff;line-height:1.8;">
                1. 왼쪽 메뉴 <b>"API Keys"</b> 클릭<br>
                2. <b>"Create Key"</b> 버튼 클릭<br>
                3. 이름 입력 (예: ebook) → <b>"Create Key"</b><br>
                4. 생성된 키 <b>복사</b> (sk-ant-api03-... 형식)
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.link_button("👆 API Keys 페이지 바로가기", "https://console.anthropic.com/settings/keys", use_container_width=True, type="primary")

    # 마무리
    st.markdown("""
    <div style="background:linear-gradient(135deg, #d4af37 0%, #b8860b 100%);padding:30px;border-radius:20px;margin:40px 0;text-align:center;">
        <p style="font-size:28px;margin:0 0 10px 0;color:#000;font-weight:800;">
            👈 복사한 키를 왼쪽 사이드바에 붙여넣기
        </p>
        <p style="font-size:16px;margin:0;color:#000;">
            💰 비용: 전자책 1권 약 200~500원 (Claude Sonnet 4 기준)
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.warning("⚠️ API 키는 생성 시 한 번만 보여줍니다. 꼭 복사해두세요!")

    st.markdown("---")

    # 도움말
    st.markdown("""
    <p style="text-align:center;font-size:20px;color:var(--text2);margin:20px 0;">
        😕 어려우시면 유튜브 영상을 보세요
    </p>
    """, unsafe_allow_html=True)

    st.link_button("📺 Claude API 키 발급 방법 (유튜브)", "https://www.youtube.com/results?search_query=anthropic+claude+api+key+발급", use_container_width=True)

    st.markdown("""
    <div style="background:rgba(100,100,100,0.2);padding:20px;border-radius:15px;margin:30px 0;text-align:center;">
        <p style="font-size:18px;margin:0;color:var(--text2);">
            💡 <b>팁:</b> 키는 한 번만 넣으면 저장돼요. 다음부터는 바로 시작!
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.stop()  # API 키 없으면 여기서 멈춤

# ==========================================
# PAGE 0: 주제 & 시장분석
# ==========================================
if current == 0:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 01</span>
        <h2>주제 선정 & 시장 분석</h2>
        <p>AI가 전자책의 성공 가능성을 분석합니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        st.markdown("### 주제 입력")

        topic = st.text_input("어떤 주제로 전자책을 쓸까요?", value=st.session_state['topic'], placeholder="예: 주식 배당으로 월 100만원", key="p0_topic", label_visibility="collapsed")
        if topic != st.session_state['topic']:
            st.session_state['topic'] = topic
            st.session_state['score_details'] = None

        # 빠른 제작 버튼 (자동 모드)
        st.markdown("""
        <div style="background:linear-gradient(135deg, #d4af37 0%, #b8860b 100%);padding:20px;border-radius:15px;margin:20px 0;text-align:center;">
            <p style="font-size:14px;margin:0 0 5px 0;color:#000;opacity:0.8;">⚡ 클릭 한 번으로</p>
            <p style="font-size:20px;margin:0;color:#000;font-weight:800;">목차 + 본문 자동 완성</p>
        </div>
        """, unsafe_allow_html=True)

        if st.button("🚀 빠른 제작 시작", use_container_width=True, key="p0_auto", type="primary"):
            if not topic:
                st.error("주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("사이드바에서 API 키를 입력해주세요")
            else:
                progress_box = st.empty()
                success = auto_generate_all(topic, progress_box)
                if success:
                    import time
                    time.sleep(1)
                    st.session_state['current_page'] = 5  # 본문 페이지로 이동
                    st.rerun()

        st.markdown("---")
        st.caption("또는 시장 분석부터 단계별로 진행:")

        if st.button("📊 시장 분석 먼저 하기", use_container_width=True, key="p0_analyze"):
            if not topic:
                st.error("주제를 입력해주세요")
            elif not get_api_key():
                st.error("사이드바에서 API 키를 입력해주세요")
            else:
                with st.spinner("AI가 시장을 분석하고 있습니다..."):
                    result = analyze_market_deep(topic)
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['score_details'] = parsed
                        st.rerun()

    with col2:
        if st.session_state.get('score_details'):
            d = st.session_state['score_details']
            score = d.get('total_score', 0)
            verdict = d.get('verdict', '')
            v_class = "verdict-go" if "추천" in verdict else ("verdict-wait" if "보류" in verdict else "verdict-no")

            st.markdown(f"""
            <div class="score-card">
                <div class="score-number">{score}</div>
                <div style="font-size:14px;color:var(--text-dim);margin-top:8px;">종합 점수</div>
                <div style="margin-top:24px;"><span class="{v_class}">{verdict}</span></div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown(f"""
            <div class="result-card" style="margin-top:20px;">
                <div style="font-size:13px;color:var(--text-dim);margin-bottom:8px;">AI 분석 요약</div>
                <div style="font-size:15px;color:var(--text-bright);line-height:1.7;">{d.get('verdict_reason', '')}</div>
            </div>
            """, unsafe_allow_html=True)

            sd = d.get('search_data', {})
            if sd:
                st.markdown(f"""
                <div class="data-card" style="margin-top:16px;">
                    <b>검색 데이터</b><br><br>
                    <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;">
                        <div>• 네이버: <b>{sd.get('naver_monthly', 'N/A')}</b></div>
                        <div>• 구글: <b>{sd.get('google_monthly', 'N/A')}</b></div>
                        <div>• 블로그: <b>{sd.get('naver_blog_posts', 'N/A')}</b></div>
                        <div>• 유튜브: <b>{sd.get('youtube_videos', 'N/A')}</b></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

            ms = d.get('market_size', {})
            comp = d.get('competition', {})

            c1, c2 = st.columns(2)
            with c1:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{ms.get("level", "")}</div><div class="stat-label">시장 규모 ({ms.get("score", 0)}점)</div></div>', unsafe_allow_html=True)
            with c2:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{comp.get("level", "")}</div><div class="stat-label">경쟁 강도 ({comp.get("score", 0)}점)</div></div>', unsafe_allow_html=True)

            if comp.get('your_opportunity'):
                st.success(f"**차별화 기회:** {comp.get('your_opportunity', '')}")

            # 경쟁 도서 검색 - 주제 키워드로 직접 검색
            current_topic = st.session_state.get('topic', '')
            if current_topic:
                st.markdown("""
                <div style="margin-top:35px;">
                    <div style="display:flex;align-items:center;gap:12px;margin-bottom:25px;">
                        <div style="width:50px;height:50px;background:linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 100%);border-radius:12px;display:flex;align-items:center;justify-content:center;">
                            <span style="font-size:26px;">🔍</span>
                        </div>
                        <div>
                            <h4 style="color:var(--gold);margin:0;font-size:22px;font-weight:600;">경쟁 도서 직접 확인하기</h4>
                            <p style="color:var(--text2);margin:4px 0 0 0;font-size:14px;">각 플랫폼에서 이 주제의 책들을 살펴보세요</p>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # 플랫폼별 검색 URL 매핑
                platforms = [
                    {
                        'name': '크몽 전자책',
                        'icon': '📘',
                        'url': 'https://kmong.com/search?c=ebook&q=',
                        'desc': '전자책/PDF 마켓',
                        'gradient': 'linear-gradient(135deg, #667eea 0%, #764ba2 100%)'
                    },
                    {
                        'name': '리디북스',
                        'icon': '📗',
                        'url': 'https://ridibooks.com/search?q=',
                        'desc': '국내 최대 전자책',
                        'gradient': 'linear-gradient(135deg, #11998e 0%, #38ef7d 100%)'
                    },
                    {
                        'name': 'YES24',
                        'icon': '📙',
                        'url': 'https://www.yes24.com/Product/Search?domain=BOOK&query=',
                        'desc': '종합 서점',
                        'gradient': 'linear-gradient(135deg, #f093fb 0%, #f5576c 100%)'
                    },
                    {
                        'name': '교보문고',
                        'icon': '📕',
                        'url': 'https://search.kyobobook.co.kr/search?keyword=',
                        'desc': '국내 대표 서점',
                        'gradient': 'linear-gradient(135deg, #4facfe 0%, #00f2fe 100%)'
                    },
                    {
                        'name': '클래스101',
                        'icon': '🎓',
                        'url': 'https://class101.net/search?query=',
                        'desc': '온라인 클래스',
                        'gradient': 'linear-gradient(135deg, #fa709a 0%, #fee140 100%)'
                    },
                    {
                        'name': '탈잉',
                        'icon': '👨‍🏫',
                        'url': 'https://taling.me/search?query=',
                        'desc': '재능 마켓',
                        'gradient': 'linear-gradient(135deg, #a8edea 0%, #fed6e3 100%)'
                    }
                ]

                search_query = urllib.parse.quote(current_topic)

                cols = st.columns(3)
                for idx, platform in enumerate(platforms):
                    with cols[idx % 3]:
                        search_url = platform['url'] + search_query
                        st.markdown(f"""
                        <a href="{search_url}" target="_blank" style="text-decoration:none;display:block;margin-bottom:15px;">
                            <div style="background:rgba(25,25,25,0.9);border:1px solid rgba(212,175,55,0.3);border-radius:16px;overflow:hidden;transition:all 0.3s ease;">
                                <div style="height:80px;background:{platform['gradient']};display:flex;align-items:center;justify-content:center;">
                                    <span style="font-size:40px;">{platform['icon']}</span>
                                </div>
                                <div style="padding:18px;text-align:center;">
                                    <div style="font-size:17px;color:var(--text);font-weight:700;margin-bottom:6px;">
                                        {platform['name']}
                                    </div>
                                    <div style="font-size:13px;color:var(--text2);margin-bottom:12px;">
                                        {platform['desc']}
                                    </div>
                                    <div style="background:linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 100%);color:var(--dark);padding:10px 16px;border-radius:8px;font-size:13px;font-weight:700;">
                                        🔍 "{current_topic[:15]}{'...' if len(current_topic) > 15 else ''}" 검색
                                    </div>
                                </div>
                            </div>
                        </a>
                        """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="empty-state">
                <div class="empty-state-text">
                    주제를 입력하고 <b>AI 시장 분석</b>을 시작하세요<br>
                    검색량, 경쟁 강도, 수익 가능성을 분석합니다
                </div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        if st.button("다음 단계로 타겟 설정", key="p0_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 1: 타겟 & 컨셉
# ==========================================
elif current == 1:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 02</span>
        <h2>타겟 설정 & 제목 생성</h2>
        <p>구매할 사람을 정하고 끌리는 제목을 만듭니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 타겟 선정")

        if st.button("AI 타겟 추천", key="p1_target"):
            if st.session_state['topic'] and get_api_key():
                with st.spinner("분석 중..."):
                    result = suggest_targets(st.session_state['topic'])
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['suggested_targets'] = parsed
                        st.rerun()

        if st.session_state.get('suggested_targets'):
            personas = st.session_state['suggested_targets'].get('personas', [])[:3]

            for i, p in enumerate(personas):
                target_name = p.get('name', '')
                target_demo = p.get('demographics', '')
                target_needs = p.get('needs', '')
                target_pains = p.get('pain_points', [])

                st.markdown(f"""<div class="data-card">
                    <b>{html.escape(str(target_name))}</b><br>
                    <small>{html.escape(str(target_demo))}</small><br>
                    <small style="color:var(--gold);">{html.escape(str(target_needs))}</small>
                </div>""", unsafe_allow_html=True)

                if st.button(f"이 타겟 선택", key=f"sel_target_{i}", use_container_width=True):
                    selected_target = f"{target_name} - {target_demo}"
                    st.session_state['target_persona'] = selected_target
                    st.session_state['p1_persona'] = selected_target
                    st.session_state['pain_points'] = ", ".join(target_pains[:5])
                    st.session_state['suggested_targets'] = None
                    st.rerun()

        st.markdown("---")
        st.markdown("### 선택된 타겟")
        persona = st.text_area("타겟:", value=st.session_state.get('target_persona', ''), height=60, key="p1_persona", placeholder="AI 추천에서 선택하거나 직접 입력")
        st.session_state['target_persona'] = persona

        if st.button("고민 심층 분석", key="p1_analyze", use_container_width=True):
            if not persona:
                st.error("타겟을 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("심층 분석 중..."):
                    r = analyze_pains_deep(st.session_state['topic'], persona)
                    parsed = parse_json(r)
                    if parsed:
                        st.session_state['analyzed_pains'] = parsed
                        surface = parsed.get('surface_pains', {}).get('pains', [])
                        hidden = parsed.get('hidden_pains', {}).get('pains', [])
                        st.session_state['pain_points'] = ", ".join((surface + hidden)[:6])
                        st.rerun()
                    else:
                        st.error("분석 실패. 다시 시도해주세요")

        if st.session_state.get('analyzed_pains'):
            p = st.session_state['analyzed_pains']
            st.markdown("**표면적 고민**")
            for pain in p.get('surface_pains', {}).get('pains', []):
                st.write(f"• {pain}")
            st.markdown("**숨겨진 진짜 고민**")
            for pain in p.get('hidden_pains', {}).get('pains', []):
                st.write(f"• {pain}")
            if p.get('marketing_hook'):
                st.info(f"**마케팅 훅:** {p.get('marketing_hook', '')}")

    with col2:
        st.markdown("### 베스트셀러급 제목 생성")

        # 선택된 제목이 있으면 상단에 확정 표시
        if st.session_state.get('book_title'):
            st.markdown(f"""
            <div style="background:linear-gradient(135deg,#10b981,#059669);padding:16px 20px;border-radius:12px;margin-bottom:20px;">
                <div style="color:white;font-size:12px;margin-bottom:6px;">✓ 확정된 제목</div>
                <div style="color:white;font-size:20px;font-weight:700;">{html.escape(st.session_state.get('book_title', ''))}</div>
                <div style="color:rgba(255,255,255,0.85);font-size:14px;margin-top:4px;">{html.escape(st.session_state.get('subtitle', ''))}</div>
            </div>
            """, unsafe_allow_html=True)

        pain_points = st.text_area("독자의 고민:", value=st.session_state['pain_points'], height=60, key="p1_pains")
        st.session_state['pain_points'] = pain_points

        if st.button("베스트셀러 제목 생성", key="p1_title"):
            if st.session_state['topic']:
                with st.spinner("베스트셀러 패턴 분석 중..."):
                    r = generate_titles_bestseller(st.session_state['topic'], st.session_state['target_persona'], st.session_state['pain_points'])
                    parsed = parse_json(r)
                    if parsed:
                        st.session_state['generated_titles'] = parsed
                        st.rerun()

        if st.session_state.get('generated_titles'):
            titles_list = st.session_state['generated_titles'].get('titles', [])[:5]
            for i, t in enumerate(titles_list):
                title_val = t.get('title', '')
                subtitle_val = t.get('subtitle', '')
                concept_val = t.get('concept', '')

                st.markdown(f"""
                <div class="title-card">
                    <div class="title-main">{html.escape(title_val)}</div>
                    <div class="title-sub">{html.escape(subtitle_val)}</div>
                    <div style="font-size:11px;color:var(--gold);margin-top:12px;letter-spacing:2px;">{html.escape(concept_val)}</div>
                </div>
                """, unsafe_allow_html=True)

                if st.button(f"✓ 이 제목으로 확정", key=f"sel_title_{i}", use_container_width=True):
                    st.session_state['book_title'] = title_val
                    st.session_state['subtitle'] = subtitle_val
                    st.toast(f"'{title_val}' 제목이 확정되었습니다!")
                    st.rerun()

        # 직접 입력 옵션
        st.markdown("---")
        st.markdown("#### 또는 직접 입력")
        manual_title = st.text_input("제목 입력", key="manual_title_v3")
        manual_subtitle = st.text_input("부제 입력", key="manual_subtitle_v3")
        if st.button("✓ 직접 입력한 제목으로 확정", key="manual_confirm_v3", use_container_width=True):
            if manual_title:
                st.session_state['book_title'] = manual_title
                st.session_state['subtitle'] = manual_subtitle if manual_subtitle else ''
                st.toast(f"'{manual_title}' 제목이 확정되었습니다!")
                st.rerun()

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p1_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 경쟁분석", key="p1_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 2: 경쟁도서 분석
# ==========================================
elif current == 2:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 03</span>
        <h2>경쟁 도서 분석</h2>
        <p>기존 도서의 부정 리뷰를 분석해서 숨은 니즈를 찾습니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 부정 리뷰 분석")

        if st.button("경쟁 도서 분석하기", use_container_width=True, key="p2_analyze"):
            if not st.session_state['topic']:
                st.error("주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("경쟁 도서 분석 중..."):
                    result = analyze_competitor_reviews(st.session_state['topic'])
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['review_analysis'] = parsed
                        concepts = parsed.get('concept_suggestions', [])
                        st.session_state['market_gaps'] = [c.get('concept', '') for c in concepts]
                        st.rerun()

        if st.session_state.get('review_analysis'):
            a = st.session_state['review_analysis']
            scope = a.get('analysis_scope', {})
            col_s1, col_s2 = st.columns(2)
            with col_s1:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{scope.get("books_analyzed", "N/A")}</div><div class="stat-label">분석 도서</div></div>', unsafe_allow_html=True)
            with col_s2:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{scope.get("negative_reviews", "N/A")}</div><div class="stat-label">부정 리뷰</div></div>', unsafe_allow_html=True)

    with col2:
        st.markdown("### 분석 결과")

        if st.session_state.get('review_analysis'):
            a = st.session_state['review_analysis']

            patterns = a.get('negative_patterns', [])
            if patterns:
                st.markdown("#### 독자 불만 패턴")
                for i, p in enumerate(patterns[:3], 1):
                    st.markdown(f"""<div class="data-card">
                        <b>{i}. {p.get('pattern', '')} ({p.get('frequency', '')})</b>
                    </div>""", unsafe_allow_html=True)
                    for rev in p.get('example_reviews', []):
                        st.caption(f'"{rev}"')
                    st.info(f"**숨겨진 니즈:** {p.get('hidden_need', '')}")
                    st.success(f"**해결책:** {p.get('solution', '')}")

            concepts = a.get('concept_suggestions', [])
            if concepts:
                st.markdown("#### 차별화 컨셉")
                for c in concepts[:2]:
                    st.markdown(f"""
                    <div class="info-card">
                        <b>「{html.escape(c.get('concept', ''))}」</b><br>
                        <span style="color:rgba(255,255,255,0.7);">{html.escape(c.get('why_works', ''))}</span>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.markdown('<div style="text-align:center;padding:60px;background:rgba(255,255,255,0.03);border-radius:16px;border:1px solid rgba(212,175,55,0.15);"><p style="color:rgba(255,255,255,0.5);">분석 버튼을 눌러주세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p2_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 학습", key="p2_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 3: 학습 & 리서치
# ==========================================
elif current == 3:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 04</span>
        <h2>학습 & 리서치</h2>
        <p>베스트셀러 분석, 트렌드 파악, 핵심 인사이트를 수집합니다</p>
    </div>
    """, unsafe_allow_html=True)

    # 탭으로 구분
    tab1, tab2, tab3 = st.tabs(["레퍼런스 추천", "트렌드 분석", "경쟁서 분석"])

    # ========== 탭1: 레퍼런스 추천 & 아이디어 ==========
    with tab1:
        topic = st.session_state.get('topic', '')

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("### 레퍼런스 자동 추천")
            st.markdown('<p style="color:var(--text2);font-size:13px;">주제에 맞는 참고 자료를 AI가 자동으로 추천합니다</p>', unsafe_allow_html=True)

            if not topic:
                st.warning("먼저 시장분석 페이지에서 주제를 입력해주세요")
            else:
                st.markdown(f'<p style="color:var(--accent);font-size:14px;margin:10px 0;">현재 주제: <b>{html.escape(topic)}</b></p>', unsafe_allow_html=True)

                ref_category = st.selectbox("추천 카테고리", ["베스트셀러 도서", "핵심 개념/이론", "성공 사례", "전문가 인사이트"], key="ref_cat")

                if st.button("레퍼런스 추천받기", use_container_width=True, key="auto_ref_btn"):
                    if not get_api_key():
                        st.error("API 키를 입력해주세요")
                    else:
                        with st.spinner("관련 레퍼런스 심층 분석 중..."):
                            prompt = f"""'{topic}' 주제로 전자책을 쓰려고 합니다.
'{ref_category}' 카테고리에서 참고할 만한 자료 3개를 추천해주세요.

중요: 마치 이 책/자료를 직접 읽은 것처럼 아주 상세하게 설명해주세요.

각 추천 자료에 대해 다음을 포함해주세요:
1. 제목과 저자
2. 책/자료의 핵심 메시지 (10문장 이상으로 상세히)
3. 주요 챕터/섹션별 핵심 내용
4. 저자의 핵심 주장과 근거
5. 실제 사례나 스토리
6. 전자책에 활용할 수 있는 구체적 인사이트

중요: 책의 모든 주요 챕터를 빠짐없이 요약해주세요. 일부만 하지 말고 전체 목차를 다 포함해주세요.

JSON 형식으로 응답:
{{
    "recommendations": [
        {{
            "title": "자료 제목",
            "author": "저자/출처",
            "core_message": "이 책의 핵심 메시지와 주장을 10문장 이상으로 상세하게 설명",
            "chapters": [
                {{"name": "1장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "2장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "3장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "4장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "5장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "6장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "7장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "8장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}}
            ],
            "key_arguments": ["저자의 핵심 주장1과 근거", "핵심 주장2와 근거", "핵심 주장3과 근거"],
            "real_examples": ["책에 나온 실제 사례/스토리 1", "사례 2", "사례 3"],
            "key_insights": ["전자책에 활용할 인사이트 1", "인사이트 2", "인사이트 3", "인사이트 4", "인사이트 5"],
            "application": "내 전자책에 구체적으로 활용하는 방법 (3문장 이상)"
        }}
    ]
}}"""
                            result = ask_ai(prompt, 0.8)
                            parsed = parse_json(result)
                            if parsed and parsed.get('recommendations'):
                                st.session_state['recommended_refs'] = parsed['recommendations']
                                st.rerun()
                            else:
                                st.error("추천 생성 실패")

                # 추천된 레퍼런스 표시
                if st.session_state.get('recommended_refs'):
                    st.markdown("---")
                    st.markdown("#### 추천 레퍼런스")

                    for i, ref in enumerate(st.session_state['recommended_refs']):
                        st.markdown(f"""<div class="data-card">
                            <b>{html.escape(str(ref.get('title', '')))}</b>
                            <br><small style="color:var(--text2);">{html.escape(str(ref.get('author', '')))}</small>
                        </div>""", unsafe_allow_html=True)

                        # 핵심 메시지
                        if ref.get('core_message'):
                            st.markdown("**핵심 메시지**")
                            st.markdown(f'<p style="color:#e0e0e0;line-height:1.8;">{html.escape(str(ref.get("core_message", "")))}</p>', unsafe_allow_html=True)

                        # 챕터 요약
                        if ref.get('chapters'):
                            st.markdown("**챕터별 요약**")
                            for ch in ref.get('chapters', []):
                                st.markdown(f"""<div style="background:rgba(255,255,255,0.05);padding:12px 16px;margin:8px 0;border-left:3px solid var(--gold);">
                                    <b style="color:var(--gold);">{html.escape(str(ch.get('name', '')))}</b><br>
                                    <span style="color:#e0e0e0;">{html.escape(str(ch.get('summary', '')))}</span>
                                </div>""", unsafe_allow_html=True)

                        # 핵심 주장
                        if ref.get('key_arguments'):
                            st.markdown("**저자의 핵심 주장**")
                            for arg in ref.get('key_arguments', []):
                                st.info(arg)

                        # 실제 사례
                        if ref.get('real_examples'):
                            st.markdown("**실제 사례**")
                            for ex in ref.get('real_examples', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(ex))}</p>', unsafe_allow_html=True)

                        # 활용 인사이트
                        if ref.get('key_insights'):
                            st.markdown("**활용 인사이트**")
                            for insight in ref.get('key_insights', []):
                                st.success(insight)

                        # 적용 방법
                        if ref.get('application'):
                            st.markdown("**내 책에 적용하는 방법**")
                            st.markdown(f'<p style="color:#e0e0e0;line-height:1.8;">{html.escape(str(ref.get("application", "")))}</p>', unsafe_allow_html=True)

                        if st.button("이 자료 저장하기", key=f"save_ref_{i}", use_container_width=True):
                            ref_item = {
                                'title': ref.get('title', ''),
                                'source': ref.get('author', ''),
                                'core_message': ref.get('core_message', ''),
                                'chapters': ref.get('chapters', []),
                                'key_arguments': ref.get('key_arguments', []),
                                'real_examples': ref.get('real_examples', []),
                                'key_insights': ref.get('key_insights', []),
                                'application': ref.get('application', ''),
                                'type': 'recommended',
                                'added_at': datetime.now().strftime('%Y-%m-%d %H:%M')
                            }
                            st.session_state['knowledge_hub'].append(ref_item)
                            st.success("저장 완료")
                            st.rerun()

                        st.markdown("---")

        with col2:
            st.markdown("### 저장된 자료 & 아이디어 도출")
            hub = st.session_state.get('knowledge_hub', [])

            if hub:
                st.caption(f"총 {len(hub)}개 자료 저장됨")

                for i, item in enumerate(hub):
                    title = item.get('title', item.get('main_topic', item.get('source', f'자료 {i+1}')))

                    st.markdown(f"""<div class="data-card">
                        <b>{html.escape(str(title))}</b><br>
                        <small>{html.escape(str(item.get('source', '')))} | {item.get('added_at', '')}</small>
                    </div>""", unsafe_allow_html=True)

                    # 핵심 메시지 전체 표시
                    if item.get('core_message'):
                        st.write(item['core_message'])

                    # 핵심 인사이트 표시
                    if item.get('key_insights'):
                        st.markdown("**핵심 인사이트:**")
                        for insight in item.get('key_insights', [])[:3]:
                            st.success(insight)

                    # 적용 방법 표시
                    if item.get('application'):
                        st.info(f"적용법: {item['application']}")

                    col_a, col_b = st.columns([1, 1])
                    with col_a:
                        if st.button("상세보기", key=f"view_ref_{i}"):
                            st.session_state[f'show_detail_{i}'] = not st.session_state.get(f'show_detail_{i}', False)
                            st.rerun()
                    with col_b:
                        if st.button("삭제", key=f"del_ref_{i}"):
                            st.session_state['knowledge_hub'].pop(i)
                            st.rerun()

                    # 상세 보기 토글
                    if st.session_state.get(f'show_detail_{i}', False):
                        if item.get('chapters'):
                            st.markdown("**챕터 요약:**")
                            for ch in item.get('chapters', []):
                                st.markdown(f"""<div style="background:rgba(255,255,255,0.05);padding:10px 14px;margin:6px 0;border-left:2px solid var(--gold);">
                                    <b style="color:var(--gold);">{html.escape(str(ch.get('name', '')))}</b><br>
                                    <span style="color:#e0e0e0;font-size:14px;">{html.escape(str(ch.get('summary', '')))}</span>
                                </div>""", unsafe_allow_html=True)
                        if item.get('key_arguments'):
                            st.markdown("**핵심 주장:**")
                            for arg in item.get('key_arguments', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(arg))}</p>', unsafe_allow_html=True)
                        if item.get('real_examples'):
                            st.markdown("**실제 사례:**")
                            for ex in item.get('real_examples', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(ex))}</p>', unsafe_allow_html=True)

                    st.markdown("---")

                st.markdown("---")
                st.markdown("#### 아이디어 도출")
                st.markdown('<p style="color:var(--text2);font-size:13px;">수집된 자료를 바탕으로 전자책 아이디어를 생성합니다</p>', unsafe_allow_html=True)

                if st.button("아이디어 생성하기", use_container_width=True, key="ideate_btn"):
                    if not get_api_key():
                        st.error("API 키를 입력해주세요")
                    else:
                        with st.spinner("아이디어 생성 중..."):
                            hub_summary = ""
                            for item in hub[:5]:
                                hub_summary += f"\n[{item.get('title', '')}]\n"
                                if item.get('core_message'):
                                    hub_summary += f"핵심 메시지: {item.get('core_message', '')}\n"
                                if item.get('key_insights'):
                                    hub_summary += f"인사이트: {', '.join(item.get('key_insights', []))}\n"
                                if item.get('application'):
                                    hub_summary += f"적용법: {item.get('application', '')}\n"

                            prompt = f"""다음 수집된 자료들을 철저히 분석하여 '{topic}' 주제의 전자책 아이디어를 도출해주세요:

수집된 자료:
{hub_summary}

위 자료들의 공통점, 차이점, 빈틈을 분석하고 다음을 포함해서 아이디어를 생성해주세요:
1. 기존 책들과 확실히 다른 차별화된 콘셉트
2. 독자의 문제를 해결하는 독특한 관점
3. 구체적인 목차/콘텐츠 구성 아이디어
4. 타겟 독자에게 강하게 어필할 포인트

JSON 형식으로 응답:
{{
    "main_concept": "핵심 콘셉트 한 문장 (경쟁작과 어떻게 다른지 명확히)",
    "unique_angles": ["독특한 관점 1 (왜 이 관점이 효과적인지 설명)", "관점 2", "관점 3"],
    "content_ideas": ["챕터 아이디어 1", "챕터 아이디어 2", "챕터 아이디어 3", "챕터 아이디어 4", "챕터 아이디어 5"],
    "appeal_points": ["어필 포인트 1", "포인트 2", "포인트 3"],
    "title_suggestions": ["제목 제안 1 (부제 포함)", "제목 제안 2 (부제 포함)", "제목 제안 3 (부제 포함)"],
    "differentiation": "경쟁작 대비 구체적인 차별화 전략 (3문장 이상)"
}}"""
                            result = ask_ai(prompt, 0.9)
                            parsed = parse_json(result)
                            if parsed:
                                st.session_state['generated_ideas'] = parsed
                                st.rerun()
                            else:
                                st.error("아이디어 생성 실패")

                # 생성된 아이디어 표시
                if st.session_state.get('generated_ideas'):
                    ideas = st.session_state['generated_ideas']

                    st.markdown(f"""<div class="summary-hub">
                        <b>핵심 콘셉트</b><br>
                        {html.escape(str(ideas.get('main_concept', '')))}
                    </div>""", unsafe_allow_html=True)

                    if ideas.get('unique_angles'):
                        st.markdown("**독특한 관점**")
                        for angle in ideas.get('unique_angles', []):
                            st.info(angle)

                    if ideas.get('title_suggestions'):
                        st.markdown("**제목 제안**")
                        for title in ideas.get('title_suggestions', []):
                            st.success(title)

                    if ideas.get('content_ideas'):
                        st.markdown("**콘텐츠 아이디어**")
                        for idea in ideas.get('content_ideas', []):
                            st.write(f"- {idea}")

                    if ideas.get('differentiation'):
                        st.markdown(f"""<div class="data-card">
                            <b>차별화 전략</b><br>
                            <small>{html.escape(str(ideas.get('differentiation', '')))}</small>
                        </div>""", unsafe_allow_html=True)
            else:
                st.markdown('<div style="text-align:center;padding:60px 20px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">추천받은 레퍼런스를 저장하면<br>아이디어를 도출할 수 있습니다</p></div>', unsafe_allow_html=True)

    # ========== 탭2: 트렌드 분석 ==========
    with tab2:
        st.markdown("### 시장 트렌드 분석")
        st.markdown('<p style="color:var(--text2);">현재 인기 있는 전자책 트렌드와 키워드를 파악합니다</p>', unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("#### 트렌드 키워드 분석")
            trend_topic = st.text_input("분석할 분야", key="trend_topic", placeholder="예: 재테크, 자기계발, 다이어트...")

            if st.button("트렌드 분석", use_container_width=True, key="trend_btn"):
                if not trend_topic:
                    st.error("분야를 입력하세요")
                elif not get_api_key():
                    st.error("API 키를 입력해주세요")
                else:
                    with st.spinner("트렌드 분석 중..."):
                        prompt = f"""'{trend_topic}' 분야의 전자책 시장 트렌드를 분석해주세요.

JSON 형식으로 응답:
{{
    "hot_keywords": ["인기 키워드 1", "키워드 2", "키워드 3", "키워드 4", "키워드 5"],
    "rising_topics": ["떠오르는 주제 1", "주제 2", "주제 3"],
    "reader_needs": ["독자가 원하는 것 1", "원하는 것 2", "원하는 것 3"],
    "content_gaps": ["시장에서 부족한 콘텐츠 1", "부족한 콘텐츠 2"],
    "recommended_angles": ["추천 접근 방식 1", "접근 방식 2", "접근 방식 3"],
    "avoid": ["피해야 할 것 1", "피해야 할 것 2"]
}}"""
                        result = ask_ai(prompt, 0.8)
                        parsed = parse_json(result)
                        if parsed:
                            st.session_state['trend_analysis'] = parsed
                            st.rerun()

            if st.session_state.get('trend_analysis'):
                ta = st.session_state['trend_analysis']
                if ta.get('hot_keywords'):
                    st.write("**인기 키워드**")
                    st.write(" | ".join(ta.get('hot_keywords', [])))
                if ta.get('rising_topics'):
                    st.write("**떠오르는 주제**")
                    for t in ta.get('rising_topics', []):
                        st.write(f"- {t}")

        with col2:
            st.markdown("#### 독자 니즈")
            if st.session_state.get('trend_analysis'):
                ta = st.session_state['trend_analysis']
                if ta.get('reader_needs'):
                    st.write("**독자가 원하는 것**")
                    for n in ta.get('reader_needs', []):
                        st.info(n)
                if ta.get('content_gaps'):
                    st.write("**시장 빈틈**")
                    for g in ta.get('content_gaps', []):
                        st.success(g)
                if ta.get('recommended_angles'):
                    st.write("**추천 접근법**")
                    for r in ta.get('recommended_angles', []):
                        st.write(f"- {r}")
            else:
                st.markdown('<div style="text-align:center;padding:60px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">분야를 입력하고<br>트렌드 분석을 시작하세요</p></div>', unsafe_allow_html=True)

    # ========== 탭3: 경쟁서 분석 ==========
    with tab3:
        st.markdown("### 경쟁 도서 분석")
        st.markdown('<p style="color:var(--text2);">경쟁 전자책의 목차, 리뷰, 강점을 분석합니다</p>', unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("#### 경쟁서 정보 입력")
            comp_title = st.text_input("책 제목", key="comp_title", placeholder="예: 돈의 심리학")
            comp_toc = st.text_area("목차 (복사/붙여넣기)", height=150, key="comp_toc", placeholder="1장. 제목\n2장. 제목\n...")
            comp_reviews = st.text_area("대표 리뷰 (선택)", height=100, key="comp_reviews", placeholder="인상적인 리뷰를 붙여넣으세요...")

            if st.button("경쟁서 분석", use_container_width=True, key="comp_btn"):
                if not comp_title or not comp_toc:
                    st.error("제목과 목차를 입력하세요")
                elif not get_api_key():
                    st.error("API 키를 입력해주세요")
                else:
                    with st.spinner("분석 중..."):
                        prompt = f"""다음 경쟁 도서를 분석해주세요:

제목: {comp_title}
목차:
{comp_toc}

리뷰: {comp_reviews if comp_reviews else '없음'}

JSON 형식으로 응답:
{{
    "book_summary": "이 책의 핵심 콘셉트",
    "target_audience": "예상 타겟 독자",
    "strengths": ["강점 1", "강점 2", "강점 3"],
    "weaknesses": ["약점/빈틈 1", "약점 2"],
    "unique_selling_point": "이 책만의 차별점",
    "improvement_opportunities": ["내 책에서 더 잘할 수 있는 것 1", "기회 2", "기회 3"],
    "key_chapters": ["핵심 챕터 1", "챕터 2"],
    "content_structure": "콘텐츠 구성 방식"
}}"""
                        result = ask_ai(prompt, 0.7)
                        parsed = parse_json(result)
                        if parsed:
                            parsed['title'] = comp_title
                            parsed['added_at'] = datetime.now().strftime('%Y-%m-%d %H:%M')
                            if 'competitor_analysis' not in st.session_state:
                                st.session_state['competitor_analysis'] = []
                            st.session_state['competitor_analysis'].append(parsed)
                            st.success("분석 완료")
                            st.rerun()

        with col2:
            st.markdown("#### 분석 결과")
            comps = st.session_state.get('competitor_analysis', [])

            if comps:
                for i, comp in enumerate(comps):
                    st.markdown(f"""<div class="data-card">
                        <b>{html.escape(str(comp.get('title', f'경쟁서 {i+1}')))}</b>
                        <br><small>{comp.get('added_at', '')}</small>
                    </div>""", unsafe_allow_html=True)
                    st.caption(comp.get('book_summary', ''))

                    if comp.get('improvement_opportunities'):
                        for o in comp.get('improvement_opportunities', [])[:2]:
                            st.success(f"차별화: {o}")

                    if st.button("삭제", key=f"del_comp_{i}"):
                        st.session_state['competitor_analysis'].pop(i)
                        st.rerun()
            else:
                st.markdown('<div style="text-align:center;padding:60px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">경쟁 도서 정보를 입력하고<br>분석해보세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p3_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음", key="p3_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 4: 목차 설계
# ==========================================
elif current == 4:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 05</span>
        <h2>목차 설계</h2>
        <p>독자의 호기심을 자극하는 목차를 만듭니다</p>
    </div>
    """, unsafe_allow_html=True)

    if st.session_state.get('market_gaps'):
        st.success(f"{len(st.session_state['market_gaps'])}개 차별화 포인트 반영")

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 목차 생성")

        st.markdown("""
        <div class="info-card">
            <b>🔥 목차 작성 팁</b><br><br>
            • 설명하지 말고 <b>궁금하게</b><br>
            • 구체적 <b>숫자 + 결과</b> 보여주기<br>
            • <b>실패담/고백</b>으로 공감 얻기<br>
            • "99%가 모르는" <b>비밀</b> 암시<br>
            • <b>반전</b>이 있을 것 같은 느낌<br><br>
            <span style="color:var(--gold);">❌ "시간관리의 중요성"</span><br>
            <span style="color:#50c878;">✓ "20대에 이걸 몰라서 5년 날렸다"</span>
        </div>
        """, unsafe_allow_html=True)

        if st.button("목차 생성하기", use_container_width=True, key="p4_outline_btn"):
            if not st.session_state.get('topic'):
                st.error("주제를 입력하세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("목차 생성 중..."):
                    result = generate_outline(
                        st.session_state['topic'],
                        st.session_state.get('target_persona', ''),
                        st.session_state.get('pain_points', ''),
                        st.session_state.get('market_gaps', [])
                    )

                    if result:
                        lines = result.split('\n')
                        chapters = []
                        current_ch = None
                        subtopics = {}

                        for line in lines:
                            orig_line = line
                            line = line.strip()
                            if not line:
                                continue

                            # 마크다운 정리 (먼저 정리한 후 검사)
                            clean_line = re.sub(r'^[#\*\s]+', '', line).strip()
                            clean_line = clean_line.replace('**', '').replace('*', '').strip()

                            # PART 또는 챕터 형식 인식 (더 유연하게)
                            is_chapter = False

                            # PART 형식 (다양한 변형)
                            if re.search(r'PART\s*\d+', clean_line, re.IGNORECASE):
                                is_chapter = True
                            # 파트 형식 (한글)
                            elif re.search(r'파트\s*\d+', clean_line):
                                is_chapter = True
                            # Chapter 형식
                            elif re.search(r'(Chapter|챕터)\s*\d+', clean_line, re.IGNORECASE):
                                is_chapter = True
                            # "1. 제목" 형식
                            elif re.match(r'^\d+[\.\)]\s*.+', clean_line) and not clean_line.startswith('-'):
                                is_chapter = True
                            # 숫자로 시작하는 제목 (예: "1 첫번째 파트")
                            elif re.match(r'^\d+\s+[가-힣A-Za-z]', clean_line):
                                is_chapter = True

                            if is_chapter:
                                name = clean_line
                                if name and len(name) > 3:
                                    current_ch = name
                                    chapters.append(current_ch)
                                    subtopics[current_ch] = []

                            # 소제목 - 다양한 형식 지원
                            elif current_ch:
                                is_subtopic = False
                                st_name = ""

                                # "-" 또는 "•" 또는 "·" 로 시작
                                if re.match(r'^\s*[\-\•\·]\s*', line):
                                    is_subtopic = True
                                    st_name = re.sub(r'^[\s\-\•\·]+', '', line).strip()
                                # 들여쓰기 된 내용
                                elif orig_line.startswith('  ') or orig_line.startswith('\t'):
                                    if not any(x in line.upper() for x in ['PART', 'CHAPTER', '파트']):
                                        is_subtopic = True
                                        st_name = line.strip().lstrip('-•· ')
                                # "  1)" 또는 "  a)" 형식
                                elif re.match(r'^\s+[\da-z][\)\.]', orig_line):
                                    is_subtopic = True
                                    st_name = re.sub(r'^[\s\da-z\)\.\-]+', '', line).strip()

                                if is_subtopic:
                                    st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                                    # 소제목이 충분히 길고 유효한 경우만 추가
                                    if st_name and len(st_name) > 3 and not re.match(r'^(PART|파트|Chapter|챕터)', st_name, re.IGNORECASE):
                                        subtopics[current_ch].append(st_name)

                        if chapters:
                            st.session_state['outline'] = chapters
                            st.session_state['chapters'] = {}
                            for ch in chapters:
                                st.session_state['chapters'][ch] = {
                                    'subtopics': subtopics.get(ch, []),
                                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                                }
                            st.success(f"{len(chapters)}개 챕터 생성!")
                            st.rerun()
                        else:
                            st.error("목차 생성 실패. 다시 시도해주세요.")
                    else:
                        st.error("AI 응답 없음. 다시 시도해주세요.")

    with col2:
        st.markdown("### 현재 목차")

        if st.session_state.get('outline'):
            # 수정 모드 토글
            if 'edit_outline_mode' not in st.session_state:
                st.session_state['edit_outline_mode'] = False

            col_view, col_edit = st.columns([1, 1])
            with col_view:
                if st.button("👁 보기 모드", use_container_width=True, disabled=not st.session_state['edit_outline_mode']):
                    st.session_state['edit_outline_mode'] = False
                    st.rerun()
            with col_edit:
                if st.button("✏️ 수정 모드", use_container_width=True, disabled=st.session_state['edit_outline_mode']):
                    st.session_state['edit_outline_mode'] = True
                    st.rerun()

            st.markdown("---")

            if st.session_state['edit_outline_mode']:
                # 수정 모드
                st.markdown('<p style="color:var(--gold);font-size:14px;">📝 제목을 직접 수정할 수 있습니다</p>', unsafe_allow_html=True)

                updated_outline = []
                updated_chapters = {}

                for ch_idx, ch in enumerate(st.session_state['outline']):
                    # 챕터 제목 수정
                    new_ch_title = st.text_input(
                        f"PART {ch_idx + 1}",
                        value=ch,
                        key=f"edit_ch_{ch_idx}"
                    )
                    updated_outline.append(new_ch_title)
                    updated_chapters[new_ch_title] = {'subtopics': [], 'subtopic_data': {}}

                    # 소제목 수정
                    subtopics = st.session_state['chapters'].get(ch, {}).get('subtopics', [])
                    new_subtopics = []
                    for st_idx, st_name in enumerate(subtopics):
                        new_st = st.text_input(
                            f"  └ 소제목 {st_idx + 1}",
                            value=st_name,
                            key=f"edit_st_{ch_idx}_{st_idx}",
                            label_visibility="collapsed"
                        )
                        if new_st.strip():
                            new_subtopics.append(new_st)
                            # 기존 데이터 유지
                            old_data = st.session_state['chapters'].get(ch, {}).get('subtopic_data', {}).get(st_name, {'questions': [], 'answers': [], 'content': ''})
                            updated_chapters[new_ch_title]['subtopic_data'][new_st] = old_data

                    updated_chapters[new_ch_title]['subtopics'] = new_subtopics
                    st.markdown("---")

                # 저장 버튼
                if st.button("💾 수정 내용 저장", use_container_width=True, type="primary"):
                    st.session_state['outline'] = updated_outline
                    st.session_state['chapters'] = updated_chapters
                    st.session_state['edit_outline_mode'] = False
                    st.success("목차가 수정되었습니다!")
                    st.rerun()

            else:
                # 보기 모드 - 예쁘게 표시
                for ch_idx, ch in enumerate(st.session_state['outline']):
                    st.markdown(f"""
                    <div style="background:linear-gradient(135deg, rgba(212,175,55,0.15) 0%, rgba(212,175,55,0.05) 100%);
                                padding:16px 20px;border-radius:12px;margin-bottom:8px;border-left:4px solid var(--gold);">
                        <span style="color:var(--gold);font-size:13px;font-weight:600;">PART {ch_idx + 1}</span>
                        <p style="color:var(--text);font-size:17px;font-weight:600;margin:8px 0 0 0;">{ch}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    subtopics = st.session_state['chapters'].get(ch, {}).get('subtopics', [])
                    for st_idx, st_name in enumerate(subtopics):
                        st.markdown(f"""
                        <div style="padding:10px 20px 10px 35px;color:var(--text);font-size:15px;">
                            <span style="color:var(--gold);margin-right:8px;">•</span>{st_name}
                        </div>
                        """, unsafe_allow_html=True)

                    st.markdown("<div style='margin-bottom:20px;'></div>", unsafe_allow_html=True)

        else:
            st.markdown('<div style="text-align:center;padding:60px;background:rgba(255,255,255,0.03);border-radius:16px;border:1px solid rgba(212,175,55,0.15);"><p style="color:rgba(255,255,255,0.5);">목차를 생성해주세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p4_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 본문", key="p4_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 5: 본문 작성
# ==========================================
elif current == 5:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 06</span>
        <h2>본문 작성</h2>
        <p>AI가 각 챕터의 콘텐츠를 작성합니다</p>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.get('outline'):
        st.warning("먼저 목차를 설계하세요")
    else:
        col_sel1, col_sel2 = st.columns([1, 1])
        with col_sel1:
            selected_ch = st.selectbox("챕터", st.session_state['outline'], key="p5_chapter")

        # 선택된 챕터가 있고 chapters에 존재하는지 확인
        if selected_ch and selected_ch in st.session_state.get('chapters', {}):
            ch_data = st.session_state['chapters'][selected_ch]
            subtopics_list = ch_data.get('subtopics', [])

            # 소제목이 있는 경우에만 선택박스 표시
            selected_st = None
            if subtopics_list:
                with col_sel2:
                    selected_st = st.selectbox("소제목", subtopics_list, key="p5_subtopic")

            # 진행률 표시
            completed = sum(1 for s in subtopics_list if ch_data.get('subtopic_data', {}).get(s, {}).get('content'))
            total = len(subtopics_list)
            if total > 0:
                st.progress(completed / total)
                st.caption(f"{completed}/{total} 완료")

            # 소제목이 선택된 경우에만 편집 UI 표시
            if selected_st:
                # subtopic_data 초기화 확인
                if 'subtopic_data' not in ch_data:
                    ch_data['subtopic_data'] = {}
                if selected_st not in ch_data['subtopic_data']:
                    ch_data['subtopic_data'][selected_st] = {'questions': [], 'answers': [], 'content': ''}

                st_data = ch_data['subtopic_data'][selected_st]

                col1, col2 = st.columns([1, 1])

                # 버튼 키를 위한 고유 식별자
                st_key = f"{selected_ch}_{selected_st}".replace(" ", "_")

                with col1:
                    st.markdown("### 인터뷰")
                    if st.button("질문 생성", key=f"gen_q_{st_key}"):
                        if not get_api_key():
                            st.error("API 키를 입력해주세요")
                        else:
                            with st.spinner("생성 중..."):
                                q_text = generate_questions(selected_st, selected_ch, st.session_state['topic'])
                                if q_text:
                                    questions = re.findall(r'Q\d+:\s*(.+)', q_text)
                                    if not questions:
                                        questions = [q.strip() for q in q_text.split('\n') if '?' in q][:3]
                                    if questions:
                                        st_data['questions'] = questions
                                        st_data['answers'] = [''] * len(questions)
                                        st.rerun()
                                    else:
                                        st.error("질문 생성에 실패했습니다")

                    if st_data.get('questions'):
                        for i, q in enumerate(st_data['questions']):
                            st.markdown(f"**Q{i+1}.** {q}")
                            # answers 리스트 크기 확인
                            while len(st_data.get('answers', [])) <= i:
                                st_data['answers'].append('')
                            st_data['answers'][i] = st.text_area(f"A{i+1}", value=st_data['answers'][i], height=80, key=f"ans_{st_key}_{i}", label_visibility="collapsed")

                with col2:
                    st.markdown("### 본문")
                    has_ans = st_data.get('questions') and any(a.strip() for a in st_data.get('answers', []))

                    if has_ans:
                        if st.button("본문 생성", key=f"gen_content_{st_key}", use_container_width=True, type="primary"):
                            if not get_api_key():
                                st.error("API 키를 입력해주세요")
                            else:
                                with st.spinner("본문 작성 중... (1~2분 소요)"):
                                    content = generate_content_premium(selected_st, selected_ch, st_data['questions'], st_data['answers'], st.session_state['topic'], st.session_state['target_persona'])
                                    if content:
                                        st_data['content'] = content
                                        st.success("본문 생성 완료!")
                                        st.rerun()
                                    else:
                                        st.error("본문 생성에 실패했습니다. 다시 시도해주세요.")
                    else:
                        st.info("왼쪽에서 질문에 답변을 입력하면 본문을 생성할 수 있습니다")

                    # 본문 표시
                    current_content = st_data.get('content', '')
                    if current_content:
                        # HTML 형식으로 변환하여 표시
                        formatted_html = format_content_html(current_content)
                        st.markdown(f"""
                        <style>
                        .content-preview-box {{
                            background:#ffffff !important;
                            padding:25px 30px;
                            border-radius:12px;
                            border:1px solid rgba(212,175,55,0.3);
                            margin:15px 0;
                            font-family:'S-CoreDream', sans-serif !important;
                            font-size:17px;
                            max-height:500px;
                            overflow-y:auto;
                        }}
                        .content-preview-box,
                        .content-preview-box p,
                        .content-preview-box span,
                        .content-preview-box div {{
                            color:#000000 !important;
                            -webkit-text-fill-color:#000000 !important;
                        }}
                        .content-preview-box b[style*="color:#e67e22"],
                        .content-preview-box p[style*="color:#e67e22"] {{
                            color:#e67e22 !important;
                            -webkit-text-fill-color:#e67e22 !important;
                        }}
                        </style>
                        <div class="content-preview-box">
                            {formatted_html}
                        </div>
                        """, unsafe_allow_html=True)
                        st.caption(f"📝 {len(current_content.replace(' ', '').replace(chr(10), '')):,}자")

                        # 이미지 추가 기능
                        st.markdown("---")
                        st.markdown("**📷 이미지 추가**")
                        uploaded_img = st.file_uploader("이미지 업로드", type=['png', 'jpg', 'jpeg'], key=f"img_{st_key}", label_visibility="collapsed")
                        if uploaded_img:
                            # 이미지 저장
                            if 'images' not in st_data:
                                st_data['images'] = []
                            img_b64 = base64.b64encode(uploaded_img.read()).decode()
                            st_data['images'].append({'name': uploaded_img.name, 'data': img_b64})
                            st.success(f"이미지 '{uploaded_img.name}' 추가됨!")
                            st.rerun()

                        # 추가된 이미지 표시
                        if st_data.get('images'):
                            st.caption(f"추가된 이미지: {len(st_data['images'])}개")
                            for idx, img in enumerate(st_data['images']):
                                col_img, col_del = st.columns([4, 1])
                                with col_img:
                                    st.image(f"data:image/png;base64,{img['data']}", caption=img['name'], width=200)
                                with col_del:
                                    if st.button("삭제", key=f"del_img_{st_key}_{idx}"):
                                        st_data['images'].pop(idx)
                                        st.rerun()

                        # 수정 기능
                        st.markdown("---")
                        with st.expander("✏️ 본문 직접 수정"):
                            st.caption("「중요단어」 → 주황색 강조 | ★ 문장 → 핵심 강조")
                            edited = st.text_area("본문 편집", value=current_content, height=400, key=f"content_{st_key}", label_visibility="collapsed")
                            if edited != current_content:
                                st_data['content'] = edited
                                st.rerun()
                    else:
                        st.markdown('<div style="text-align:center;padding:80px 20px;background:rgba(255,255,255,0.03);border-radius:12px;border:1px dashed rgba(212,175,55,0.3);"><p style="color:var(--text2);font-size:16px;">본문이 아직 없습니다<br>질문에 답변 후 "본문 생성" 버튼을 누르세요</p></div>', unsafe_allow_html=True)
            else:
                st.info("이 챕터에는 소제목이 없습니다. 목차를 다시 생성해주세요.")

        st.markdown("---")
        st.markdown("### 전체 본문")
        full_content = get_full_content()
        if full_content:
            char_count = len(full_content.replace(' ', '').replace('\n', ''))
            st.success(f"총 {char_count:,}자 | 약 {char_count//500}페이지")

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p5_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 출력", key="p5_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 6: 표지 디자인
# ==========================================
elif current == 6:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 07</span>
        <h2>표지 디자인</h2>
        <p>전문 디자인 툴로 고품질 표지를 만드세요</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 표지 정보 정리")

        # 이전 페이지에서 설정한 제목/부제 자동 연동
        saved_title = st.session_state.get('book_title', '')
        saved_subtitle = st.session_state.get('subtitle', '')

        cover_title = st.text_input("표지 제목", value=saved_title, key="cover_title", placeholder="예: 돈의 속성")
        cover_subtitle = st.text_input("부제목", value=saved_subtitle, key="cover_subtitle", placeholder="예: 당신이 모르는 부의 법칙")
        cover_author = st.text_input("저자명", key="cover_author", placeholder="예: 홍길동")

        st.markdown("---")
        st.markdown("### AI 표지 스타일 추천")

        if st.button("내 주제에 맞는 표지 스타일 추천받기", use_container_width=True, key="ai_cover_suggest"):
            topic = st.session_state.get('topic', '')
            if not topic:
                st.error("시장분석 페이지에서 주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("베스트셀러 표지 분석 중..."):
                    prompt = f"""'{topic}' 주제의 전자책 표지 디자인을 추천해주세요.

이 분야의 실제 베스트셀러 책 표지를 분석해서 추천해주세요.

JSON 형식으로 응답:
{{
    "recommended_style": "추천 스타일명",
    "color_scheme": "추천 색상 조합 (예: 검정 배경 + 금색 텍스트)",
    "design_concept": "디자인 콘셉트 설명 (2문장)",
    "typography_tip": "타이포그래피 팁 (폰트 스타일, 크기 등)",
    "reference_books": ["참고할 베스트셀러 표지 1", "표지 2", "표지 3"],
    "canva_search_keyword": "Canva에서 검색할 키워드 (영문)"
}}"""
                    result = ask_ai(prompt, 0.7)
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['cover_suggestion'] = parsed
                        st.rerun()

        if st.session_state.get('cover_suggestion'):
            sug = st.session_state['cover_suggestion']
            st.markdown(f"""<div class="data-card">
                <b>추천 스타일: {html.escape(str(sug.get('recommended_style', '')))}</b><br>
                <small>색상: {html.escape(str(sug.get('color_scheme', '')))}</small>
            </div>""", unsafe_allow_html=True)
            st.write(sug.get('design_concept', ''))
            if sug.get('typography_tip'):
                st.info(f"💡 타이포그래피 팁: {sug.get('typography_tip', '')}")
            if sug.get('reference_books'):
                st.markdown("**참고 베스트셀러:**")
                for book in sug.get('reference_books', []):
                    st.caption(f"- {book}")
            if sug.get('canva_search_keyword'):
                st.session_state['canva_keyword'] = sug.get('canva_search_keyword', '')

    with col2:
        st.markdown("### Canva로 표지 만들기")

        st.markdown("""
        <div class="data-card">
            <p style="font-size:16px;margin-bottom:15px;">
                <b>Canva</b>는 전문 디자이너 수준의 표지를 무료로 만들 수 있는 온라인 툴입니다.
            </p>
            <p style="color:var(--text2);font-size:14px;">
                ✓ 수천 개의 프로 템플릿<br>
                ✓ 드래그 앤 드롭 편집<br>
                ✓ 무료 이미지/아이콘<br>
                ✓ 한글 폰트 지원
            </p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")

        # Canva 검색 키워드 설정
        canva_keyword = st.session_state.get('canva_keyword', 'book cover')

        # Canva 책 표지 템플릿 링크
        canva_url = f"https://www.canva.com/templates/?query={canva_keyword}%20book%20cover"

        st.markdown(f"""
        <a href="{canva_url}" target="_blank" style="
            display:block;
            background:linear-gradient(135deg,#7c3aed,#6366f1);
            color:white;
            padding:18px 24px;
            border-radius:12px;
            text-decoration:none;
            text-align:center;
            font-size:18px;
            font-weight:600;
            margin-bottom:15px;
            transition:transform 0.2s;
        ">
            🎨 Canva에서 표지 만들기
        </a>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### 사용 방법")
        st.markdown("""
        1. **Canva 열기** - 위 버튼 클릭 (무료 가입)
        2. **템플릿 선택** - 마음에 드는 디자인 클릭
        3. **텍스트 수정** - 제목, 부제목, 저자명 입력
        4. **다운로드** - PNG 또는 PDF로 저장
        """)

        st.markdown("---")
        st.markdown("### 추천 검색어")

        search_keywords = [
            "ebook cover", "book cover minimalist",
            "book cover gold", "book cover business",
            "korean book cover", "self help book cover"
        ]

        cols = st.columns(2)
        for i, kw in enumerate(search_keywords):
            with cols[i % 2]:
                if st.button(kw, key=f"canva_kw_{i}", use_container_width=True):
                    st.session_state['canva_keyword'] = kw
                    st.rerun()

        st.markdown("---")

        # 복사할 텍스트
        if cover_title or cover_subtitle or cover_author:
            st.markdown("### 복사할 텍스트")
            copy_text = f"제목: {cover_title}\n부제목: {cover_subtitle}\n저자: {cover_author}"
            st.code(copy_text, language=None)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p6_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음", key="p6_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 7: 최종 출력
# ==========================================
elif current == 7:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 08</span>
        <h2>최종 출력</h2>
        <p>완성된 전자책을 다운로드하세요</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1.5, 1])

    with col1:
        st.markdown("### 다운로드")

        final_title = st.text_input("제목", value=st.session_state.get('book_title', ''), key="p6_title")
        final_subtitle = st.text_input("부제", value=st.session_state.get('subtitle', ''), key="p6_subtitle")

        full = f"{final_title}\n{final_subtitle}\n\n{'='*50}\n\n"
        for ch in st.session_state.get('outline', []):
            if ch in st.session_state.get('chapters', {}):
                ch_data = st.session_state['chapters'][ch]
                ch_content = ""
                for s in ch_data.get('subtopics', []):
                    c = ch_data.get('subtopic_data', {}).get(s, {}).get('content', '')
                    if c:
                        ch_content += f"\n\n【{s}】\n\n{clean_content(c)}"
                if ch_content:
                    full += f"\n\n{ch}\n{'-'*40}{ch_content}\n"

        st.markdown("**미리보기**")
        st.text_area("전체 내용", value=full, height=300, disabled=True, key="p7_preview")

        # 저자명 가져오기
        author_name = st.session_state.get('author_name', '') or st.session_state.get('interview_data', {}).get('author_name', '')

        # 다운로드 버튼 3개
        st.markdown("### 📥 다운로드")

        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("📄 TXT", full, file_name=f"{final_title or 'ebook'}.txt", use_container_width=True, key="p7_txt")
        with c2:
            # HTML 내보내기 - 특수문자 이스케이프 처리
            escaped_title = html.escape(final_title)
            escaped_content = html.escape(full).replace('\n', '<br>')
            html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{escaped_title}</title>
    <style>
        body {{
            max-width: 800px;
            margin: 0 auto;
            padding: 60px 40px;
            font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 2;
            background: #fafafa;
            color: #333;
        }}
        h1 {{ font-size: 32px; color: #1a1a2e; margin-bottom: 10px; }}
        h2 {{ font-size: 14px; color: #888; font-weight: normal; }}
    </style>
</head>
<body>
{escaped_content}
</body>
</html>"""
            st.download_button("🌐 HTML", html_content, file_name=f"{final_title or 'ebook'}.html", use_container_width=True, key="p7_html")

        with c3:
            # DOCX 다운로드
            if DOCX_AVAILABLE:
                docx_data, docx_error = create_ebook_docx(
                    final_title,
                    final_subtitle,
                    author_name,
                    st.session_state.get('chapters', {}),
                    st.session_state.get('outline', []),
                    st.session_state.get('interview_data', {})
                )
                if docx_data:
                    st.download_button(
                        "📘 WORD",
                        docx_data,
                        file_name=f"{final_title or 'ebook'}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="p7_docx"
                    )
                else:
                    st.button("📘 WORD", disabled=True, use_container_width=True, key="p7_docx_disabled")
                    if docx_error:
                        st.caption(f"⚠️ {docx_error[:30]}")
            else:
                st.button("📘 WORD", disabled=True, use_container_width=True, key="p7_docx_na")
                st.caption("pip install python-docx")

        total = len(full.replace(' ', '').replace('\n', ''))
        if total > 0:
            st.success(f"총 {total:,}자 | 약 {total//500}페이지")

    with col2:
        st.markdown("### 현황")
        total_st = sum(len(ch.get('subtopics', [])) for ch in st.session_state.get('chapters', {}).values())
        done = sum(1 for ch in st.session_state.get('chapters', {}).values() for s in ch.get('subtopic_data', {}).values() if s.get('content'))

        if total_st > 0:
            st.progress(done / total_st)
            st.write(f"**완료:** {done}/{total_st}")

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 2, 1])
    with c1:
        if st.button("이전", key="p7_prev", use_container_width=True):
            go_prev()
            st.rerun()


st.markdown("""
<div style="
    text-align: center;
    padding: 30px 20px;
    margin-top: 50px;
    border-top: 1px solid rgba(212,175,55,0.3);
    color: #ffffff !important;
    font-size: 16px;
    letter-spacing: 2px;
    background: rgba(0,0,0,0.3);
">
    <span style="color: #d4af37;">CASHMAKER</span> | 제작: <span style="color: #ffffff;">남현우 작가</span>
</div>
""", unsafe_allow_html=True)
